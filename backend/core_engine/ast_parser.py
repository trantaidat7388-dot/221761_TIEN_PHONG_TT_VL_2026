import os
import re
import hashlib
import base64
from typing import Dict, Any, List, Optional
from lxml import etree

from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .config import MAP_STYLE, HEADING_PATTERNS, A_NAMESPACE, REL_NAMESPACE, W_NAMESPACE
from docx.oxml.ns import qn
from .utils import loc_ky_tu
from .xu_ly_toan import BoXuLyToan
from .xu_ly_ole_equation import ole_equation_to_latex
from .semantic_parser import du_doan_loai_node
from .word_loader import mo_tai_lieu_word_co_fallback

class WordASTParser:
    """
    Parser converts a Word document (.docx) into an Intermediate Representation (IR).
    The IR is a JSON-serializable dictionary capturing the semantic meaning of the document
    (Metadata + Body Nodes) independently of LaTeX layout.
    """
    def __init__(self, doc_path: str, thu_muc_anh: str = "images", mode: str = "latex"):
        self.doc_path = doc_path
        self.thu_muc_anh = thu_muc_anh
        self.mode = mode  # "latex" or "word2word"
        self.doc = None
        self._temp_word_files: List[str] = []
        self.bo_toan = BoXuLyToan()
        self.dem_anh = 0
        self.total_formulas = 0
        
        # Intermediate Representation
        self.ir: Dict[str, Any] = {
            "metadata": {
                "title": "",
                "authors": [],
                "abstract": "",
                "keywords": [],
                "total_formulas": 0
            },
            "body": [],
            "references": []
        }
        
    def parse(self) -> Dict[str, Any]:
        """Main entry point to parse the document."""
        try:
            self.doc, self._temp_word_files = mo_tai_lieu_word_co_fallback(self.doc_path)

            elements = self._extract_elements_in_order()
            self._build_semantic_tree(elements)
            # Skip citation post-processing in word2word mode to avoid
            # re-indexing [1] -> \cite{ref1} -> [1] round-trip corruption.
            if self.mode != "word2word":
                self._post_process_citations()

            return self.ir
        finally:
            for temp_path in self._temp_word_files:
                try:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                except Exception:
                    pass
        
    def _extract_elements_in_order(self) -> List[tuple]:
        """Flatten the document body including elements inside content controls."""
        elements = []
        body = self.doc.element.body
        
        def traverse(node):
            if not hasattr(node, "tag") or not isinstance(node.tag, str):
                return
            tag = node.tag.split("}")[-1]
            if tag == "p":
                elements.append(("paragraph", Paragraph(node, self.doc)))
            elif tag in ("oMathPara", "oMath"):
                elements.append(("omml", node))
                # Capture nested blocks only from explicit containers where paragraph
                # text is semantically embedded (e.g., text boxes / content controls),
                # avoiding broad recursion that can duplicate paragraphs.
                seen_container_ids = set()
                for descendant in node.iter():
                    dtag = descendant.tag.split("}")[-1] if hasattr(descendant, "tag") and isinstance(descendant.tag, str) else ""
                    if dtag in ("txbxContent", "sdtContent"):
                        did = id(descendant)
                        if did in seen_container_ids:
                            continue
                        seen_container_ids.add(did)
                        for nested in descendant:
                            traverse(nested)
            elif tag == "tbl":
                elements.append(("table", Table(node, self.doc)))
            else:
                for child in node:
                    traverse(child)
                    
        for node in body:
            traverse(node)
            
        return elements
        
    def _get_style_name(self, element: Any) -> str:
        """Safely extract style name from a Paragraph object, handling missing styles or attribute errors."""
        try:
            if hasattr(element, "style") and element.style is not None:
                name = getattr(element.style, "name", "")
                return str(name or "")
        except (AttributeError, KeyError, Exception):
            pass
        return ""

    def _is_abstract_label(self, text: str) -> bool:
        norm = re.sub(r"^[\d\.]+\s*", "", (text or "").strip().upper())
        return norm.startswith("ABSTRACT") or norm.startswith("TÓM TẮT") or norm.startswith("TOM TAT")
        
    def _is_keywords_label(self, text: str) -> bool:
        # ACM style often uses "Additional Keywords and Phrases:"
        return bool(re.match(r'^(?:(?:ADDITIONAL\s+)?KEYWORDS?|TỪ KHÓA|TU KHOA|INDEX TERMS)\b', (text or "").strip(), re.IGNORECASE))

    def _image_ext_from_content_type(self, content_type: str) -> str:
        ct = (content_type or '').lower()
        if 'jpeg' in ct:
            return 'jpg'
        if 'x-emf' in ct or ct.endswith('/emf'):
            return 'emf'
        if 'x-wmf' in ct or ct.endswith('/wmf'):
            return 'wmf'
        if '/' in ct:
            return ct.split('/')[-1]
        return 'png'

    def _save_image_from_relationship(self, rel) -> Optional[str]:
        """Persist image blob and return LaTeX path; convert EMF/WMF to PNG when possible."""
        try:
            img_blob = rel.target_part.blob
            img_ext = self._image_ext_from_content_type(getattr(rel.target_part, 'content_type', ''))

            img_hash = hashlib.md5(img_blob).hexdigest()[:8]
            if self.thu_muc_anh:
                os.makedirs(self.thu_muc_anh, exist_ok=True)
                out_dir = self.thu_muc_anh
            else:
                out_dir = ''

            if img_ext in ('emf', 'wmf'):
                source_name = f"img_{img_hash}.{img_ext}"
                source_path = os.path.join(out_dir, source_name) if out_dir else source_name
                if not os.path.exists(source_path):
                    with open(source_path, "wb") as f:
                        f.write(img_blob)
                try:
                    from PIL import Image
                    with Image.open(source_path) as img:
                        png_name = f"img_{img_hash}.png"
                        png_path = os.path.join(out_dir, png_name) if out_dir else png_name
                        img.convert("RGBA").save(png_path, format="PNG")
                    try:
                        if os.path.exists(source_path):
                            os.remove(source_path)
                    except Exception:
                        pass
                    final_name = png_name
                except Exception:
                    print(f"[WARN] Skip unsupported vector image format: {source_name}")
                    return None
            else:
                final_name = f"img_{img_hash}.{img_ext}"
                final_path = os.path.join(out_dir, final_name) if out_dir else final_name
                if not os.path.exists(final_path):
                    with open(final_path, "wb") as f:
                        f.write(img_blob)

            ten_thu_muc = os.path.basename(self.thu_muc_anh) if self.thu_muc_anh else ''
            return f"{ten_thu_muc}/{final_name}" if ten_thu_muc else final_name
        except Exception:
            return None

    def _includegraphics_options(self, width_expr: str) -> str:
        """Use bounded image sizing to preserve layout when source vector crop info is lossy."""
        return f"width={width_expr},keepaspectratio"
        
    def _is_body_label(self, text: str) -> bool:
        norm = re.sub(r"^[\d\.]+\s*", "", (text or "").strip().upper())
        for kw in ["INTRODUCTION", "GIỚI THIỆU", "GIOI THIEU", "MỞ ĐẦU", "CHAPTER 1", "BACKGROUND"]:
            if norm.startswith(kw):
                return True
        if re.match(r"^[IV]+\.\s+", (text or "").strip()):
            return True
        return False

    def _is_authors_label(self, text: str) -> bool:
        norm = (text or "").strip().upper()
        for kw in ["AUTHORS", "TÁC GIẢ", "TAC GIA"]:
            if kw in norm and len(text) < 15:
                return True
        return False

    def _is_references_label(self, text: str) -> bool:
        norm = re.sub(r"^[\d\.]+\s*", "", (text or "").strip().upper())
        # Avoid false positives like "References and Footnotes" in publisher templates.
        if re.match(r"^REFERENCES\s*[:\.]?$", norm):
            return True
        if re.match(r"^BIBLIOGRAPHY\s*[:\.]?$", norm):
            return True
        if norm.startswith("TÀI LIỆU THAM KHẢO") or norm.startswith("TAI LIEU THAM KHAO"):
            return True
        return False

    def _looks_like_reference_entry(self, text: str) -> bool:
        t = (text or '').strip()
        if not t or len(t) < 12:
            return False
        tl = t.lower()

        # Skip common guide/template lines.
        if re.match(r'^(examples?:|footnotes?|references?\s+and\s+footnotes|books?|periodicals?|reports?|patents?|electronic\s+sources|standards?)\b', tl):
            return False
        if tl.startswith("reference numbers are set"):
            return False
        if tl.startswith("other than books"):
            return False
        if tl.startswith("for papers published"):
            return False
        if "first check if you have an existing account" in tl:
            return False
        if "ieee.org/publications_standards/publications/" in tl:
            return False

        # Typical bibliography signals.
        if re.search(r'\b(19|20)\d{2}\b', t):
            return True
        if re.search(r'\b(doi|arxiv|vol\.|pp\.|patent|thesis|dissertation|tech\.\s*rep\.|[Oo]nline\.|available:)\b', t):
            return True
        if re.search(r'https?://|www\.', t):
            return True
        # Catch common start patterns for bibliography entries like [1], 1., etc.
        if re.match(r'^\[?\d+\]?\s+\.?[A-Z]', t):
            return True
        if re.match(r'^[A-Z][a-z]+,\s+[A-Z]\.', t): # Author Name, I.
            return True

        # Author-like leading pattern: initials/names followed by comma.
        if re.match(r'^([A-Z]\.\s*){1,4}[A-Za-z\-\']+\s*,', t):
            return True

        return False

    def _get_para_text_with_br(self, p) -> str:
        """Helper to reliably extract text from a paragraph, including OMML math equations.
        Previously, iterating over p.runs ignored all math nodes."""
        return self._parse_paragraph(p).get("text", "").strip()

    # ====== HEURISTIC: Table/Image detection (ported from legacy xu_ly_bang.py) ======

    def _la_bang_chua_anh(self, table: Table) -> bool:
        """Phát hiện bảng chứa chủ yếu ảnh (figure layout).
        Ported from BoXuLyBang.la_bang_chua_anh()."""
        try:
            so_cell_co_anh = 0
            so_cell_co_text_dai = 0
            tong_cell = 0
            cells_da_kiem = set()

            for hang in table.rows:
                for cell in hang.cells:
                    cell_id = id(cell._tc)
                    if cell_id in cells_da_kiem:
                        continue
                    cells_da_kiem.add(cell_id)
                    tong_cell += 1
                    cell_text = (cell.text or "").strip()
                    co_anh = False

                    for para in cell.paragraphs:
                        for run in para.runs:
                            blips = run._element.findall(f'.//{{{A_NAMESPACE}}}blip')
                            if blips:
                                co_anh = True
                                break
                        drawings = para._element.findall(f'.//{{{A_NAMESPACE}}}blip')
                        if drawings:
                            co_anh = True

                    if co_anh:
                        so_cell_co_anh += 1

                    if re.match(r'^[\(\[]*[a-zA-Z0-9][\)\]]*\.?$', cell_text):
                        pass
                    elif re.match(r'^(Hình|Figure|Fig\.?|Bảng|Table)\s*\d+(\.\d+)*', cell_text, re.IGNORECASE):
                        pass
                    elif len(cell_text) > 20:
                        so_cell_co_text_dai += 1

            if tong_cell == 0:
                return False
            if so_cell_co_anh >= 1:
                if so_cell_co_text_dai <= 1:
                    return True
                if so_cell_co_anh / tong_cell >= 0.3:
                    return True
        except Exception as e:
            print(f'[WARNING] la_bang_chua_anh (AST): {e}')
        return False

    def _trich_xuat_anh_tu_bang(self, table: Table) -> List[str]:
        """Trích xuất ảnh từ bảng figure-layout, lưu vào thu_muc_anh.
        Ported from BoXuLyBang.trich_xuat_anh_tu_bang()."""
        danh_sach_anh = []
        seen_names = set()
        for hang in table.rows:
            for cell in hang.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        blips = run._element.findall(f'.//{{{A_NAMESPACE}}}blip')
                        for blip in blips:
                            embed = blip.get(f'{{{REL_NAMESPACE}}}embed')
                            if not embed:
                                continue
                            try:
                                rel = para.part.rels.get(embed)
                                if not rel:
                                    continue
                                latex_path = self._save_image_from_relationship(rel)
                                if not latex_path:
                                    continue
                                img_name = os.path.basename(latex_path)
                                if img_name not in seen_names:
                                    seen_names.add(img_name)
                                    danh_sach_anh.append(img_name)
                            except Exception:
                                pass
        return danh_sach_anh

    def _tim_caption_con_trong_bang(self, table: Table) -> List[str]:
        """Tìm caption con (a), (b)... trong các cell của bảng."""
        danh_sach = []
        try:
            for hang in table.rows:
                for cell in hang.cells:
                    text = (cell.text or "").strip()
                    match = re.match(r'^\(([a-z])\)(.*)$', text)
                    if match:
                        nhan = match.group(1)
                        mo_ta = match.group(2).strip()
                        caption = f"({nhan})"
                        if mo_ta:
                            caption += f" {mo_ta}"
                        danh_sach.append(loc_ky_tu(caption))
        except Exception:
            pass
        return danh_sach

    def _bat_caption_bang(self, elements: List[tuple], idx: int, used_nodes: set) -> str:
        """Bắt caption thật của bảng từ paragraph ngay phía TRÊN (idx - 1).
        Ported from ChuyenDoiWordSangLatex.bat_caption_bang()."""
        try:
            idx_prev = idx - 1
            if idx_prev < 0 or idx_prev >= len(elements):
                return None

            loai_prev, para_prev = elements[idx_prev]
            if loai_prev != 'paragraph':
                return None

            text_prev = (para_prev.text or "").strip()
            if not text_prev:
                return None

            # Two-line IEEE table caption support:
            #   TABLE II
            #   AVERAGE OF ...
            # Prefer the descriptive line and consume the label line.
            idx_prev2 = idx - 2
            text_prev2 = ""
            if idx_prev2 >= 0:
                loai_prev2, para_prev2 = elements[idx_prev2]
                if loai_prev2 == 'paragraph':
                    text_prev2 = (para_prev2.text or "").strip()

            is_prev_table_line = bool(re.match(r'^(BẢNG|BANG|TABLE)\b', text_prev, re.IGNORECASE))
            is_prev2_table_label = bool(re.match(r'^(BẢNG|BANG|TABLE)\s*[IVXLCDM\d]+\s*$', text_prev2, re.IGNORECASE))

            if is_prev_table_line:
                used_nodes.add(idx_prev)
                if is_prev2_table_label:
                    used_nodes.add(idx_prev2)
                caption_text = self._chuan_hoa_ten_caption(text_prev, kind='table')
                return caption_text

            # Fallback: previous line is plain caption title and line before is TABLE label.
            if text_prev and is_prev2_table_label:
                used_nodes.add(idx_prev)
                used_nodes.add(idx_prev2)
                return self._chuan_hoa_ten_caption(text_prev, kind='table')
        except Exception as e:
            print(f"[WARNING] _bat_caption_bang: {e}")
        return None

    def _bat_caption_hinh_theo_style(self, elements: List[tuple], idx: int, used_nodes: set) -> str:
        """Fallback caption extraction for templates using dedicated caption styles.

        Supports patterns like "Example of a figure caption. (figure caption)"
        even when they do not start with "Figure/Fig".
        """
        def _extract_text(raw: str) -> str:
            txt = loc_ky_tu((raw or '').strip())
            if not txt:
                return ""
            txt = re.sub(r'\(\s*figure\s+caption\s*\)', '', txt, flags=re.IGNORECASE).strip()
            txt = re.sub(r'^(example\s+of\s+a\s+)?figure\s+caption\s*[:\-–]\s*', '', txt, flags=re.IGNORECASE).strip()
            txt = re.sub(r'\s{2,}', ' ', txt)
            return txt

        # Prefer nearby look-ahead first.
        for buoc in range(1, 16):
            j = idx + buoc
            if j >= len(elements):
                break
            loai, phan_tu = elements[j]
            if loai != 'paragraph':
                continue
            text = (phan_tu.text or '').strip()
            if not text:
                continue
            style_name = self._get_style_name(phan_tu).lower()
            if 'heading 1' in style_name:
                break
            is_caption_style = ('figure caption' in style_name) or ('caption' == style_name)
            is_caption_marker = 'figure caption' in text.lower()
            if is_caption_style or is_caption_marker:
                used_nodes.add(j)
                cleaned = _extract_text(text)
                if cleaned:
                    return cleaned

        # Then short look-behind.
        for buoc in range(1, 8):
            j = idx - buoc
            if j < 0:
                break
            loai, phan_tu = elements[j]
            if loai != 'paragraph':
                continue
            text = (phan_tu.text or '').strip()
            if not text:
                continue
            style_name = self._get_style_name(phan_tu).lower()
            is_caption_style = ('figure caption' in style_name) or ('caption' == style_name)
            is_caption_marker = 'figure caption' in text.lower()
            if is_caption_style or is_caption_marker:
                used_nodes.add(j)
                cleaned = _extract_text(text)
                if cleaned:
                    return cleaned

        return None

    def _bat_caption_hinh(self, elements: List[tuple], idx: int, used_nodes: set) -> str:
        """Bắt caption thật của hình từ paragraph phía DƯỚI (tìm tối đa 5 đoạn).
        Ported from ChuyenDoiWordSangLatex.bat_caption_hinh()."""
        try:
            for buoc in range(1, 6):
                idx_sau = idx + buoc
                if idx_sau >= len(elements):
                    break
                loai, phan_tu = elements[idx_sau]
                if loai == 'table':
                    break
                if loai != 'paragraph':
                    continue
                text = (phan_tu.text or "").strip()
                if not text:
                    continue
                # FIX 1: Support decimal chapter numbers like "Figure 3.1" and IEEE "Fig. 1."
                if re.match(r'^(HÌNH|HINH|ẢNH|ANH|FIGURE|FIG|PIG)(?:\.|\b)', text, re.IGNORECASE):
                    used_nodes.add(idx_sau)
                    caption_text = self._chuan_hoa_ten_caption(text, kind='figure')
                    return caption_text
                # Dừng nếu gặp section heading mới
                style_name = self._get_style_name(phan_tu)
                if style_name and 'Heading' in style_name:
                    break
        except Exception as e:
            print(f"[WARNING] _bat_caption_hinh: {e}")
        return None

    def _chuan_hoa_ten_caption(self, text: str, kind: str) -> str:
        """Return pure caption content without leading label/number.

        This ensures display style (Table/Fig naming) is delegated to the
        LaTeX template/class instead of duplicating labels from Word input.
        """
        caption_text = loc_ky_tu((text or '').strip())
        if not caption_text:
            return caption_text

        # Strip Word field-code artifacts like: "TableSEQ Table * ARABIC1"
        caption_text = re.sub(
            r"^(?:TableSEQ|Table\s*SEQ)\s*Table\s*\\?\*\s*ARABIC\s*\d*\s*",
            "",
            caption_text,
            flags=re.IGNORECASE,
        )
        caption_text = re.sub(
            r"^SEQ\s*Table\s*\\?\*\s*ARABIC\s*\d*\s*",
            "",
            caption_text,
            flags=re.IGNORECASE,
        )

        if kind == 'table':
            # Examples stripped: "Table 1:", "TABLE 3.1 -", "BANG 2.", "TABLE I" (IEEE)
            pattern = r'^(Bảng|BANG|Bang|Table|TABLE)\.?\s*[\dIIVX]+(\.\d+)*\s*[:\.\-–—]?\s*'
        else:
            # Examples stripped: "Figure 2:", "Fig. 3.1", "HINH 1 -", "ẢNH 5", "Pig. 3"
            pattern = r'^(Hình|HINH|Hình|Ảnh|ANH|ẢNH|Figure|FIGURE|Fig\.?|Pig\.?)\s*\d+(\.\d+)*\s*[:\.\-–—]?\s*'

        caption_text = re.sub(pattern, '', caption_text, flags=re.IGNORECASE).strip()
        return caption_text

    def _is_title_paragraph(self, p: Paragraph, idx: int) -> bool:
        """Heuristic for title: usually bold, large, or specific style 'Title'."""
        text = (p.text or "").strip()
        if not text or len(text) < 3: return False
        if "Short Title" in text or "ACM Reference Format" in text: return False
        
        style_name = self._get_style_name(p)
        style_lc = style_name.lower()
        if "title" in style_lc or "header" in style_lc:
            return True
        
        # Heuristics: Center aligned + Bold or Large font + Bold
        aligned_center = False
        try:
            if p.paragraph_format.alignment == 1: aligned_center = True
        except: pass
        
        runs = p.runs
        all_bold = all(r.bold for r in runs if (r.text or "").strip()) if runs else False
        large_font = any(r.font.size and r.font.size.pt >= 14 for r in runs) if runs else False
        
        if (aligned_center and all_bold) or (large_font and all_bold):
            return True
            
        return False

    def _build_semantic_tree(self, elements: List[tuple]):
        """State machine to classify elements into Metadata vs Body nodes."""
        print(f"[*] _build_semantic_tree: Processing {len(elements)} elements")
        state = "pre_title"

        def _la_style_heading(style_name: str) -> bool:
            s = (style_name or "").strip().lower()
            if not s:
                return False
            if s.startswith("heading") or s.startswith("head"):
                return True
            return s in {"heading1", "heading2", "heading3", "heading4", "headings", "referencehead"}
        
        # Temp buffers for metadata
        title_buf = []
        authors_buf = []
        abstract_buf = []
        keywords_buf = []
        seen_figure_paths = set()
        
        # Set of element indices already used as captions (skip in body)
        used_nodes = set()
        
        # Pre-scan: mark caption paragraphs (Table captions above, Figure captions below)
        for idx, (etype, element) in enumerate(elements):
            if etype == 'table':
                # Table caption: look 1 paragraph ABOVE
                if idx > 0:
                    prev_type, prev_el = elements[idx - 1]
                    if prev_type == 'paragraph':
                        text = (prev_el.text or "").strip()
                        # Strictly match short table labels like "TABLE I" or "Table 1"
                        if text and len(text) < 100 and re.match(r'^(BẢNG|BANG|TABLE)\s+[IVXLCDM\d]+\s*[:.\-]?\s*$', text, re.IGNORECASE):
                            used_nodes.add(idx - 1)
                        else:
                            # Two-line table caption pattern:
                            # line N-2: "TABLE I"
                            # line N-1: "DATASET FEATURES ..."
                            if idx > 1:
                                prev2_type, prev2_el = elements[idx - 2]
                                if prev2_type == 'paragraph':
                                    text2 = (prev2_el.text or "").strip()
                                    if text2 and re.match(r'^(BẢNG|BANG|TABLE)\s+[IVXLCDM\d]+\s*[:.\-]?\s*$', text2, re.IGNORECASE):
                                        used_nodes.add(idx - 2)
                                        if text:
                                            used_nodes.add(idx - 1)
            if etype == 'paragraph':
                # Check if this paragraph has images (look for drawings/blips or vml/imagedata)
                has_img = bool(element._p.findall(f'.//{{{A_NAMESPACE}}}blip')) or bool(element._p.findall(r'.//{urn:schemas-microsoft-com:vml}imagedata'))
                if has_img:
                    # Figure caption: look 1-5 paragraphs BELOW
                    # FIX 1: Match decimal figure numbers like "Figure 3.1"
                    for step in range(1, 6):
                        idx_after = idx + step
                        if idx_after >= len(elements):
                            break
                        a_type, a_el = elements[idx_after]
                        if a_type == 'table':
                            break
                        if a_type != 'paragraph':
                            continue
                        a_text = (a_el.text or "").strip()
                        if not a_text:
                            continue
                        # Strictly match figure labels like "Fig. 1." or "Figure 2."
                        # If the paragraph is very long, it's likely body text starting with "Figure X..."
                        if len(a_text) < 300 and re.match(r'^(HÌNH|HINH|ẢNH|ANH|FIGURE|FIG|PIG)\s+[IVXLCDM\d\.]+\s*[:.\-]?\s*', a_text, re.IGNORECASE):
                            used_nodes.add(idx_after)
                            break
                        style_name = self._get_style_name(a_el)
                        if style_name and 'Heading' in style_name:
                            break
        
        for idx, (etype, element) in enumerate(elements):
            # Skip paragraphs already consumed as captions
            if idx in used_nodes:
                continue
                
            text = ""
            is_bold = False
            prediction = "PARAGRAPH"
            style_name = ""
            style_cmd = ""
            
            if etype == "omml":
                try:
                    omml_tag = element.tag.split('}')[-1] if hasattr(element, "tag") else ""
                    is_block = omml_tag == "oMathPara"
                    self.total_formulas += 1
                    try:
                        omml_str = etree.tostring(element, encoding='unicode')
                        omml_b64 = base64.b64encode(omml_str.encode('utf-8')).decode('utf-8')
                        if is_block:
                            node_text = f"\\begin{{equation}}\n«OMML:{omml_b64}»\n\\end{{equation}}\n"
                        else:
                            node_text = f" «OMML:{omml_b64}» "
                    except Exception:
                        latex_math = self.bo_toan.omml_element_to_latex(element)
                        if is_block:
                            node_text = f"\\begin{{equation}}\n{latex_math}\n\\end{{equation}}\n"
                        else:
                            node_text = f" ${latex_math}$ "

                    self.ir["body"].append({"type": "paragraph", "text": node_text, "has_math": True})
                except Exception as e:
                    print(f"[WARNING] Lỗi parse OMML element: {e}")
                continue

            if etype == "paragraph":
                text = (element.text or "").strip()
                has_img_para = bool(element._p.findall(f'.//{{{A_NAMESPACE}}}blip')) or bool(element._p.findall(r'.//{urn:schemas-microsoft-com:vml}imagedata'))
                if not text and state == "pre_title" and (not has_img_para):
                    continue # Skip empty leading lines (but keep image-only paragraphs)
                if text == "Short Title": continue
                
                style_name = self._get_style_name(element)
                style_cmd = MAP_STYLE.get(style_name, "")
                
                # Font features
                for r in element.runs:
                    if (r.text or "").strip() and r.bold:
                        is_bold = True
                        break
                prediction = du_doan_loai_node(text, idx, is_bold)

                # State Transitions
                if state == "pre_title":
                    if style_cmd == r"\title" or self._is_title_paragraph(element, idx):
                        state = "title"
                    elif self._is_abstract_label(text) or prediction == "ABSTRACT" or style_name == "Abstract":
                        state = "abstract"
                    elif (
                        self._is_authors_label(text)
                        or style_cmd == r"\author"
                        or prediction == "AUTHOR"
                        or style_name in ("Authors", "Author", "AuthorsBlock")
                        or style_name.lower() in ("authors", "author", "authorsblock")
                    ):
                        state = "authors"
                    elif self._is_body_label(text) or prediction == "HEADING" or _la_style_heading(style_name):
                        state = "body"
                
                # State-specific transition checks (to exit current state)
                if state == "title":
                    if style_name == "Subtitle" or style_name == "subtitle":
                        state = "title"
                    elif self._is_abstract_label(text) or prediction == "ABSTRACT" or style_name == "Abstract":
                        state = "abstract"
                    elif self._is_body_label(text) or prediction == "HEADING" or _la_style_heading(style_name):
                        state = "body"
                    elif (len(text) > 250 or prediction == "AUTHOR" or style_name in ("Authors", "AuthorsBlock", "Author") or self._is_authors_label(text)):
                        state = "authors"
                        
                elif state == "authors":
                    if (self._is_abstract_label(text) or prediction == "ABSTRACT" or style_name == "Abstract" or
                        self._is_keywords_label(text) or prediction == "KEYWORDS" or style_name in ("KeyWords", "Keywords", "CCSCONCEPTS") or
                        self._is_body_label(text) or prediction == "HEADING" or _la_style_heading(style_name) or
                        (style_name in ("BodyText", "Body Text", "Normal") and len(text) > 150)):
                        
                        if self._is_abstract_label(text) or prediction == "ABSTRACT" or style_name == "Abstract":
                            state = "abstract"
                        elif self._is_keywords_label(text) or prediction == "KEYWORDS" or style_name in ("KeyWords", "Keywords", "CCSCONCEPTS"):
                            state = "keywords"
                        else:
                            state = "body"

                elif state == "abstract":
                    if (self._is_keywords_label(text) or prediction == "KEYWORDS" or style_name in ("KeyWords", "Keywords", "CCSCONCEPTS") or
                        self._is_body_label(text) or prediction == "HEADING" or _la_style_heading(style_name)):
                        
                        if self._is_keywords_label(text) or prediction == "KEYWORDS" or style_name in ("KeyWords", "Keywords", "CCSCONCEPTS"):
                            state = "keywords"
                        else:
                            state = "body"
                    
                    # Safety: if abstract contains a table, it likely means the abstract block is over (IEEE style)
                    if etype == "table":
                        state = "body"

                elif state == "keywords":
                    keyword_overflow = (
                        len(text) > 140 and
                        (len(text.split()) > 20 or '.' in text)
                    )

                    if keyword_overflow:
                        # Some IEEE docs merge the first body paragraph into keywords line.
                        # Stop metadata capture and keep this paragraph as body content.
                        state = "body"
                        node = self._parse_paragraph(element)
                        self.ir["body"].append(node)
                        continue

                    if (self._is_body_label(text) or prediction == "HEADING" or _la_style_heading(style_name)):
                        state = "body"

                # Global: detect references section (can transition from any state)
                if (self._is_references_label(text) or
                    style_name.lower() in ("referenceitem", "references", "bibliography")):
                    state = "references"
                    # If this is just the label, don't add it as a paragraph yet
                    # (it will be skipped in the 'references' block below)
                    if self._is_references_label(text):
                        continue

                # Action: Add to buffer or ir
                if state == "title":
                    title_buf.append(loc_ky_tu(text))
                elif state == "authors":
                    # Strip boilerplate
                    if not any(x in text for x in ("Submission Template", "Reference Format", "Short Title")):
                        if "\n" in text:
                            for line in text.splitlines():
                                line_clean = line.strip()
                                if not line_clean:
                                    continue
                                # IEEE Word template often stores author blocks as:
                                # "line 1: ...", "line 2: ..."; strip this wrapper.
                                line_clean = re.sub(r'^line\s*\d+\s*:\s*', '', line_clean, flags=re.IGNORECASE)
                                if line_clean:
                                    authors_buf.append(loc_ky_tu(line_clean))
                        else:
                            cleaned = re.sub(r'^line\s*\d+\s*:\s*', '', text, flags=re.IGNORECASE).strip()
                            if cleaned:
                                authors_buf.append(loc_ky_tu(cleaned))
                elif state == "abstract":
                    # FIX 3: Smart split — paragraph may contain BOTH "Abstract." and "Keywords:"
                    combined_match = re.search(
                        r'(?:abstract|t[oó]m\s+t[aắ]t)[\.\u2014\u2013\-:]*\s*(.+?)\s*\b(?:keywords?|index\s+terms?|t[uừ]\s+kh[oó]a)\b\s*[:\.\-–—]+\s*(.+)',
                        text, re.IGNORECASE | re.DOTALL
                    )
                    if combined_match:
                        abs_text = combined_match.group(1).strip()
                        kw_text  = combined_match.group(2).strip()
                        if abs_text:
                            abstract_buf.append(loc_ky_tu(abs_text))
                        if kw_text:
                            keywords_buf.append(loc_ky_tu(kw_text))
                        # Jump state forward so subsequent parsing picks up body correctly
                        state = "keywords"
                    else:
                        # Handle Springer "Abstract.", IEEE "Abstract—" (em-dash), and Vietnamese "Tóm Tắt."
                        clean_text = re.sub(r"^(?:abstract|t[oó]m\s+t[aắ]t)\s*[\.\u2013\u2014\-:]*\s*", "", text, flags=re.IGNORECASE).strip()
                        if clean_text:
                            abstract_buf.append(loc_ky_tu(clean_text))
                elif state == "keywords":
                    # Remove label — handle both en-dash (–) and em-dash (—) separators
                    clean_text = re.sub(r"^(Additional Keywords and Phrases\s*[:\-–—]|Keywords?\s*[:\-–—]|Index Terms\s*[:\-–—]|Từ khóa\s*[:\-–—])\s*", "", text, flags=re.IGNORECASE).strip()
                    if clean_text:
                        keywords_buf.append(loc_ky_tu(clean_text))
                elif state == "body" or (has_img_para and state != "references"):
                    # Suppress orphan table labels that are part of captions and should not become body text.
                    if re.match(r'^\s*(?:BẢNG|BANG|TABLE)\s+[IVXLCDM\d]+\s*[:.\-]?\s*$', text, re.IGNORECASE):
                        continue
                    # Suppress Word SEQ field-code artifacts that show up as standalone caption lines.
                    if re.match(
                        r'^\s*(?:BẢNG|BANG|TABLE)\s*SEQ\s*TABLE\s*\\?\*\s*ARABIC\s*\d*\s*.*$',
                        text,
                        re.IGNORECASE,
                    ):
                        continue
                    # Suppress inline table caption paragraphs when the next element is a table.
                    idx_next = idx + 1
                    if idx_next < len(elements):
                        loai_next, _ = elements[idx_next]
                        if loai_next == "table" and re.match(
                            r'^\s*(?:BẢNG|BANG|TABLE)\s*[IVXLCDM\d]+\b.*$',
                            text,
                            re.IGNORECASE,
                        ):
                            continue

                    try:
                        node = self._parse_paragraph(element)
                        node_text = node.get('text', '')
                        # Post-process: nếu paragraph chứa standalone figure, tìm caption phía dưới
                        if '\\includegraphics' in node_text and '\\begin{figure' in node_text:
                            img_match = re.search(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}', node_text)
                            img_path = img_match.group(1).strip() if img_match else ""
                            if img_path and img_path in seen_figure_paths:
                                continue
                            caption = self._bat_caption_hinh(elements, idx, used_nodes)
                            if not caption:
                                caption = self._bat_caption_hinh_theo_style(elements, idx, used_nodes)
                            if caption:
                                node['text'] = node_text.replace('\\caption{}', f'\\caption{{{caption}}}')
                            if img_path:
                                seen_figure_paths.add(img_path)
                        self.ir["body"].append(node)
                    except Exception as e:
                        print(f"[WARNING] Lỗi parse paragraph {idx}: {e}")
                        # Fallback: add raw text if parsing failed
                        self.ir["body"].append({"type": "paragraph", "text": loc_ky_tu(text) if self.mode != "word2word" else text, "style": style_name})
                elif state == "references":
                    # Allow returning to body if we see a new heading (e.g. Appendix)
                    # BUT only if it's NOT the references label itself!
                    if (_la_style_heading(style_name) or prediction == "HEADING") and not self._is_references_label(text):
                        state = "body"
                        node = self._parse_paragraph(element)
                        self.ir["body"].append(node)
                        continue

                    # Skip the header label itself (e.g., "Tài Liệu Tham Khảo", "References")
                    if (not self._is_references_label(text)):
                        if self._looks_like_reference_entry(text):
                            node = self._parse_paragraph(element)
                            node_text = node.get("text", "")
                            if node_text.startswith("\\url{") and len(self.ir["references"]) > 0:
                                self.ir["references"][-1]["text"] += " " + node_text
                            else:
                                self.ir["references"].append(node)
                        else:
                            # If it doesn't look like a reference but we are in reference state,
                            # it might be a continuation or just a non-standard reference.
                            # We keep it as a body paragraph to avoid losing content.
                            node = self._parse_paragraph(element)
                            self.ir["body"].append(node)
                
            elif etype == "table":
                # Bypass: Nếu bảng đứng trước abstract, có thể đây là Author Table đặc thù của IEEE
                if state in ("pre_title", "title", "authors") and not abstract_buf:
                    try:
                        table_lines = []
                        table_lines = []
                        seen_cells = set()
                        for row in element.rows:
                            for cell in row.cells:
                                cell_id = id(cell._tc)
                                if cell_id in seen_cells:
                                    continue
                                seen_cells.add(cell_id)
                                cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()]).strip()
                                if cell_text:
                                    table_lines.extend([line.strip() for line in cell_text.split('\n') if line.strip()])
                                        
                        table_text = "\n".join(table_lines)
                        if len(table_text) < 1500 and not self._la_bang_chua_anh(element):
                            state = "authors"
                            # Ensure table-based author content is sanitized
                            authors_buf.extend([loc_ky_tu(ln) for ln in table_lines])
                            # Đánh dấu bảng này đã được dùng cho authors để tránh tự xuất hiện lại trong body (nếu được xử lý tiếp)
                            used_nodes.add(idx)
                            continue
                    except Exception as e:
                        print(f"[WARN] Error parsing potential IEEE author table: {e}")
                        
                # Tables force body UNLESS we are already in references
                if state != "references":
                    state = "body"

                eq_node = self._detect_equation_table(element)
                if eq_node:
                    self.ir["body"].append(eq_node)
                else:
                    danh_sach_anh = self._trich_xuat_anh_tu_bang(element)
                    if danh_sach_anh:
                        # Group all images from the table into a SINGLE figure block
                        # so the IEEE renderer treats them as one figure with one caption.
                        caption_chinh = self._bat_caption_hinh(elements, idx, used_nodes)
                        if not caption_chinh:
                            caption_chinh = self._bat_caption_hinh_theo_style(elements, idx, used_nodes)
                        ten_thu_muc = os.path.basename(self.thu_muc_anh)
                        self.dem_anh += 1
                        fig_tex = "\\begin{figure}[H]\n\\centering\n"
                        for ten_anh in danh_sach_anh:
                            img_path = f"{ten_thu_muc}/{ten_anh}"
                            if img_path in seen_figure_paths:
                                continue
                            seen_figure_paths.add(img_path)
                            fig_tex += (
                                f"  \\includegraphics[width=\\columnwidth,height=0.4\\textheight,"
                                f"keepaspectratio]{{{img_path}}}\n"
                            )
                        fig_tex += f"  \\caption{{{caption_chinh or ''}}}\n"
                        fig_tex += f"  \\label{{fig:img_{self.dem_anh}}}\n"
                        fig_tex += "\\end{figure}\n\n"
                        self.ir["body"].append({"type": "paragraph", "text": fig_tex})
                    else:
                        # Regular data table — with caption from look-behind
                        caption = self._bat_caption_bang(elements, idx, used_nodes)
                        table_node = self._parse_table(element)
                        if caption:
                            table_node["caption"] = caption
                        
                        if state == "references":
                            # If we are in references state and see a table, maybe the refs are in the table!
                            # Convert table cell text into reference entries if they look like it
                            for row in table_node.get("data", []): # Corrected from "rows" to "data"
                                for cell in row:
                                    cell_text = cell.get("text", "")
                                    if cell_text and self._looks_like_reference_entry(cell_text):
                                        self.ir["references"].append({"type": "paragraph", "text": cell_text})
                            # Also keep it in body as a table just in case
                            self.ir["body"].append(table_node)
                        else:
                            self.ir["body"].append(table_node)

        extracted_title = " ".join(title_buf).strip()
        
        # Fallback: Nếu Parser heuristic không tìm thấy Title, bốc ngay Paragraph đầu tiên trong body làm Title
        if not extracted_title and len(self.ir["body"]) > 0:
            for i, p_node in enumerate(self.ir["body"]):
                p_text = p_node.get("text", "")
                is_figure_para = "\\begin{figure" in p_text or "\\includegraphics" in p_text
                if p_node.get("type") == "paragraph" and p_text.strip() and (not is_figure_para):
                    extracted_title = p_node.get("text").strip()
                    self.ir["body"].pop(i)
                    break
        
        if not authors_buf and len(self.ir["body"]) > 0:
            author_candidates = []
            while len(self.ir["body"]) > 0:
                nxt = self.ir["body"][0]
                nxt_text = nxt.get("text", "")
                is_figure_para = "\\begin{figure" in nxt_text or "\\includegraphics" in nxt_text
                if nxt.get("type") == "paragraph" and (not is_figure_para) and len(nxt_text.split()) < 15:
                    author_candidates.append(nxt.get("text"))
                    self.ir["body"].pop(0)
                else:
                    break
            authors_buf.extend(author_candidates)
                     
        # Final metadata assignment
        self.ir["metadata"]["title"] = extracted_title
        parsed_authors = self._parse_authors(authors_raw=authors_buf)
        self.ir["metadata"]["authors"] = parsed_authors
        self.ir["metadata"]["author_block"] = ""  # Will be generated by renderer based on template class
        self.ir["metadata"]["abstract"] = "\n\n".join(abstract_buf).strip()
        kw_candidates = [k.strip() for k in " ".join(keywords_buf).replace(";", ",").split(",") if k.strip()]
        kw_list = []
        for kw in kw_candidates:
            clean_kw = re.sub(r"^(Additional Keywords and Phrases\s*[:\-–—]|Keywords?\s*[:\-–—]|Index Terms\s*[:\-–—]|Từ khóa\s*[:\-–—])\s*", "", kw, flags=re.IGNORECASE).strip()
            clean_kw = re.sub(r"^[\-–—,;:.\s]+", "", clean_kw)
            clean_kw = re.sub(r"[\-–—,;:.\s]+$", "", clean_kw)
            clean_kw = clean_kw.strip("\"'`“”‘’")
            if not clean_kw:
                continue
            if re.search(r"\b(first\s+section|a\s+subsection|references?|acknowledg|table\s+\d|fig\.?\s*\d)\b", clean_kw, re.IGNORECASE):
                continue
            if len(clean_kw) > 80:
                continue
            if len(clean_kw.split()) > 7:
                continue
            kw_list.append(clean_kw)

        # De-duplicate while preserving order.
        seen_kw = set()
        kw_list = [k for k in kw_list if not (k.lower() in seen_kw or seen_kw.add(k.lower()))]
        self.ir["metadata"]["keywords"] = kw_list
        self.ir["metadata"]["keywords_str"] = ", ".join(kw_list)
        self.ir["metadata"]["total_formulas"] = self.total_formulas

        # Post-process: merge consecutive figure blocks with the exact same caption
        # This handles the case where IEEE document has images in separate paragraphs 
        # but they belong to the same figure group with a single shared caption below them.
        merged_body = []
        for node in self.ir["body"]:
            if not merged_body:
                merged_body.append(node)
                continue
                
            prev = merged_body[-1]
            if prev.get("type") == "paragraph" and node.get("type") == "paragraph":
                prev_text = prev.get("text", "")
                curr_text = node.get("text", "")
                
                # Check if both are figures
                if "\\begin{figure" in prev_text and "\\begin{figure" in curr_text:
                    # Extract captions
                    prev_cap_match = re.search(r"\\caption\{([^}]*)\}", prev_text)
                    curr_cap_match = re.search(r"\\caption\{([^}]*)\}", curr_text)
                    
                    if prev_cap_match and curr_cap_match:
                        prev_cap = prev_cap_match.group(1).strip()
                        curr_cap = curr_cap_match.group(1).strip()
                        
                        # If captions are identical and non-empty
                        if prev_cap and curr_cap and prev_cap == curr_cap:
                            # Merge them!
                            # Extract \includegraphics from curr_text
                            curr_paths = re.findall(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", curr_text)
                            
                            # Insert these paths into prev_text before \caption
                            for path in curr_paths:
                                img_line = f"  \\includegraphics[width=\\columnwidth,height=0.4\\textheight,keepaspectratio]{{{path}}}\n"
                                prev_text = prev_text.replace(f"\\caption{{{prev_cap}}}", f"{img_line}  \\caption{{{prev_cap}}}")
                            
                            prev["text"] = prev_text
                            continue
                            
            merged_body.append(node)
            
        self.ir["body"] = merged_body

    def _extract_author_with_superscripts(self, p) -> str:
        """Extract author text preserving superscript markers as \\textsuperscript{}."""
        result = ""
        for run in p.runs:
            txt = run.text
            if not txt:
                continue
            if run.font.superscript:
                result += f"\\textsuperscript{{{loc_ky_tu(txt)}}}"
            else:
                result += loc_ky_tu(txt)
        return result.strip() if result.strip() else loc_ky_tu(p.text).strip()

    def _detect_equation_table(self, t) -> dict:
        """Detect if a Word table is actually a layout table for an equation.
        Pattern: 1 row, 2-3 columns, last column contains equation number like (1).
        """
        try:
            rows = t.rows
            if len(rows) != 1:
                return None
            cells = rows[0].cells
            if len(cells) < 2 or len(cells) > 3:
                return None
            
            # Check if last cell is equation number like (1), (2), (1a), (A1), etc.
            last_text = (cells[-1].text or "").strip()
            if not re.match(r'^\(([A-Za-z0-9\.\-\*]+)\)$', last_text):
                return None
            
            # The formula usually occupies all cells except the rightmost equation number.
            formula_parts = [(c.text or "").strip() for c in cells[:-1] if (c.text or "").strip()]
            formula_text = " ".join(formula_parts)
            
            # Check if first cell(s) contain math-like content (=, ½, fractions, symbols)
            math_indicators = ['=', '½', '∑', '∏', '∫', '+', '−', '×', '÷', 'frac', '\\', 
                             'Accuracy', 'Precision', 'Recall', 'Sensitivity', 'Specificity']
            has_math = any(ind in formula_text for ind in math_indicators)
            if not has_math:
                return None
            
            # It's an equation table! Convert to equation node
            # Try to extract OMML math if present
            ns_m = "http://schemas.openxmlformats.org/officeDocument/2006/math"
            tbl_xml = t._tbl
            omml_math = tbl_xml.find(f".//{{{ns_m}}}oMathPara")
            if omml_math is None:
                omml_math = tbl_xml.find(f".//{{{ns_m}}}oMath")
            
            if omml_math is not None:
                latex_math = self.bo_toan.omml_element_to_latex(omml_math)
                # Capture raw OMML XML
                try:
                    omml_str = etree.tostring(omml_math, encoding='unicode')
                    omml_b64 = base64.b64encode(omml_str.encode('utf-8')).decode('utf-8')
                    eq_text = f"\\begin{{equation}}\n«OMML:{omml_b64}»\n\\tag{{{last_text.strip('()')}}}\n\\end{{equation}}"
                except Exception:
                    eq_text = f"\\begin{{equation}}\n{latex_math}\n\\tag{{{last_text.strip('()')}}}\n\\end{{equation}}"
            else:
                # No OMML, use the plain text and try to make it look like an equation
                # Replace ½ with \\frac{1}{2}, etc.
                formula_text = formula_text.replace('½', '\\frac{1}{2}')
                formula_text = formula_text.replace('×', '\\times')
                formula_text = formula_text.replace('−', '-')
                eq_text = f"\\begin{{equation}}\n{formula_text}\n\\tag{{{last_text.strip('()')}}}\n\\end{{equation}}"
            
            self.total_formulas += 1
            return {"type": "paragraph", "text": eq_text, "has_math": True}
        except Exception:
            return None

    def _post_process_citations(self):
        """Scans all body paragraphs and replaces bracket citations like [1] or [1, 2] with \\cite{ref1}."""
        citation_pattern = re.compile(r'\[([0-9\s,\-]+)\]')
        
        def replace_cite(match):
            inner = match.group(1).replace(" ", "")
            parts = inner.split(",")
            ref_keys = []
            for part in parts:
                if "-" in part:
                    bounds = part.split("-")
                    if len(bounds) == 2 and bounds[0].isdigit() and bounds[1].isdigit():
                        start, end = int(bounds[0]), int(bounds[1])
                        for i in range(start, end + 1):
                            ref_keys.append(f"ref{i}")
                    else:
                        ref_keys.append(f"ref{part}")
                elif part.isdigit():
                    ref_keys.append(f"ref{part}")
                else:
                    return match.group(0) # Not a valid citation format
            
            if ref_keys:
                return f"\\cite{{{','.join(ref_keys)}}}"
            return match.group(0)

        for node in self.ir["body"]:
            if node["type"] in ("paragraph", "section"):
                node["text"] = citation_pattern.sub(replace_cite, node["text"])
            elif node["type"] == "table":
                if "data" in node:
                    for row in node["data"]:
                        for cell in row:
                            if "text" in cell:
                                cell["text"] = citation_pattern.sub(replace_cite, cell["text"])

    def _parse_authors(self, authors_raw: List[str]) -> List[Dict]:
        """Convert a flat list of author strings into structured dicts.
        Hỗ trợ: tác giả trên nhiều dòng, hoặc nhiều tác giả trên 1 dòng (phân tách bằng dấu phẩy / and).
        Handles superscript number mapping between author names and affiliations.
        """
        def _looks_like_affiliation(text: str) -> bool:
            t = (text or '').lower()
            affil_keywords = [
                '@', 'university', 'institute', 'dept', 'faculty', 'ltd',
                'department', 'school', 'lab', 'center', 'organization',
                'city', 'country', 'affiliation', 'vietnam', 'viet nam', 'việt nam',
                'princeton', 'heidelberg', 'germany', 'usa', 'tiergartenstr', 'street', 'road'
            ]
            if any(kw in t for kw in affil_keywords):
                return True
            if re.search(r'\b\d{3,}\b', t):
                return True
            if re.search(r'\b(nj|ca|tx|ny|uk|us)\b', t):
                return True
            return False

        def _looks_like_ieee_membership_suffix(text: str) -> bool:
            t = (text or '').lower()
            membership_keywords = [
                'member',
                'senior member',
                'fellow',
                'life fellow',
                'student member',
                'ieee',
            ]
            return any(kw in t for kw in membership_keywords)

        authors = []
        current = None

        # Tiền xử lý: tách các dòng chứa nhiều tên tác giả (phân tách bằng ',' hoặc ' and ')
        expanded = []
        for info in authors_raw:
            clean = info.strip()
            if not clean:
                continue
            # Kiểm tra: dòng này có chứa nhiều tác giả không?
            # Điều kiện: có dấu phẩy hoặc ' and ' NHƯNG không phải affiliation
            is_affil = _looks_like_affiliation(clean)
            if (',' in clean or ' and ' in clean) and not is_affil:
                # Tách bằng ' and ' trước, rồi ','
                parts = re.split(r'\s+and\s+', clean)
                all_parts = []
                for p in parts:
                    if ',' in p and _looks_like_ieee_membership_suffix(p):
                        all_parts.append(p.strip())
                    else:
                        all_parts.extend([x.strip() for x in p.split(',') if x.strip()])
                # Nếu tất cả các phần đều ngắn (tên người), tách thành nhiều tác giả
                if len(all_parts) > 1 and all(len(p) < 60 for p in all_parts):
                    expanded.extend(all_parts)
                    continue
            expanded.append(clean)

        # Fast-path for common publisher layout: all author names first, then affiliations.
        # Without this branch, sequential parsing can incorrectly attach all affiliations to
        # the last author only.
        if expanded:
            first_affil_idx = None
            has_name_after_affil = False
            for idx, item in enumerate(expanded):
                is_affil = _looks_like_affiliation(item)
                if is_affil and first_affil_idx is None:
                    first_affil_idx = idx
                elif first_affil_idx is not None and (not is_affil) and len(item) < 60:
                    has_name_after_affil = True
                    break

            if first_affil_idx is not None and first_affil_idx >= 2 and not has_name_after_affil:
                names = expanded[:first_affil_idx]
                affils = expanded[first_affil_idx:]
                authors = [{"name": n, "affiliations": []} for n in names]

                if len(affils) >= len(authors):
                    # 1-1 by index; any extra lines (often email/address continuation)
                    # are appended to the last author.
                    for i, af in enumerate(affils):
                        target_idx = i if i < len(authors) else len(authors) - 1
                        authors[target_idx]["affiliations"].append(af)
                elif affils:
                    # Fewer affiliations than names: map by index then reuse the last one.
                    for i, a in enumerate(authors):
                        mapped = affils[i] if i < len(affils) else affils[-1]
                        a["affiliations"].append(mapped)

                current = None

        if not authors:
            for info in expanded:
                clean = info.strip()
                if not clean:
                    continue
                is_affil = _looks_like_affiliation(clean)
                if not current:
                    current = {"name": clean, "affiliations": []}
                elif is_affil or (len(clean) >= 40 and '@' in clean):
                    current["affiliations"].append(clean)
                elif len(clean) < 50 and not is_affil:
                    authors.append(current)
                    current = {"name": clean, "affiliations": []}
                else:
                    current["affiliations"].append(clean)
            if current:
                authors.append(current)

        # Post-process: map superscript numbers/symbols in names to numbered affiliations
        affil_map = {}  # marker -> affiliation text
        unmapped_affils = []
        all_affils = []
        for a in authors:
            for af in a.get("affiliations", []):
                all_affils.append(af)
                m_num = re.match(r'^(\d+)\s+(.*)', af)
                if m_num:
                    affil_map[m_num.group(1)] = m_num.group(2).strip()
                else:
                    m_sym = re.search(r'([*†‡]+)', af)
                    if m_sym:
                        affil_map[m_sym.group(1)] = af.strip()
                    else:
                        unmapped_affils.append(af.strip())

        # FIX 2: Collect unmapped standalone emails
        email_pattern = re.compile(r'[\w\.\-\+]+@[\w\.\-]+')
        extracted_emails = []
        for af in unmapped_affils:
            extracted_emails.extend(email_pattern.findall(af))

        if affil_map:
            # FIX 2: Map by superscript numbers/symbols (*, †, ‡ included)
            for a in authors:
                name = a["name"]
                # Match trailing superscript markers: digits, *, †, ‡, commas
                num_match = re.search(r'([\d\*†‡,\s]+)$', name)
                if num_match:
                    raw_markers = num_match.group(1)
                    markers = re.findall(r'\d+|[*†‡]', raw_markers)
                    a["name"] = name[:num_match.start()].strip()
                    a["affiliations"] = []
                    for mk in markers:
                        if mk in affil_map:
                            a["affiliations"].append(affil_map[mk])
                    # If * present and no mapped affil found, fall back to all
                    if not a["affiliations"] and '*' in raw_markers and '*' not in affil_map:
                        a["affiliations"] = list(affil_map.values())
                else:
                    # No markers => give all numbered affiliations as fallback
                    a["affiliations"] = [v for k, v in affil_map.items() if str(k).isdigit()]
        elif len(authors) > 1:
            # Fallback: if some have no affil, share all
            any_empty = any(not a.get("affiliations") for a in authors)
            any_multi = any(len(a.get("affiliations", [])) > 1 for a in authors)
            if any_empty and any_multi:
                for a in authors:
                    a["affiliations"] = all_affils[:]

        # FIX 2: Attach remaining extracted emails to authors who have no email yet
        if extracted_emails:
            email_idx = 0
            for a in authors:
                has_email = any('@' in aff for aff in a.get('affiliations', []))
                if not has_email and email_idx < len(extracted_emails):
                    a.setdefault('affiliations', []).append(extracted_emails[email_idx])
                    email_idx += 1

        # Normalize affiliations: split fused email text and remove placeholder metadata
        email_pattern = re.compile(r'[\w\.\-\+]+@[\w\.\-]+')
        for a in authors:
            normalized_affils = []
            for aff in a.get('affiliations', []):
                raw = (aff or '').strip()
                if not raw:
                    continue
                emails = email_pattern.findall(raw)
                non_email = email_pattern.sub('', raw)
                non_email = re.sub(r'\s{2,}', ' ', non_email).strip(' ,;')
                if non_email:
                    normalized_affils.append(non_email)
                for em in emails:
                    normalized_affils.append(em)

            # Deduplicate while preserving order
            seen = set()
            deduped = []
            for item in normalized_affils:
                key = item.lower()
                if key in seen:
                    continue
                seen.add(key)
                deduped.append(item)
            a['affiliations'] = deduped

        # Remove common template placeholders accidentally captured from sample templates
        placeholder_name = re.compile(
            r'^(first|second|third|fourth|fifth|sixth)\s+(?:[a-z]\.?\s+)?author(?:,\s*(?:jr\.?|sr\.?))?$',
            re.IGNORECASE,
        )
        membership_only = re.compile(r'^(fellow|member|senior\s+member|student\s+member|life\s+fellow|ieee)$', re.IGNORECASE)
        placeholder_affil = re.compile(r'(springer\s+heidelberg|tiergartenstr|69121\s+heidelberg)', re.IGNORECASE)
        ieee_template_affil = re.compile(
            r'(dept\.?\s*name\s*of\s*organization|\(of\s*affiliation\)|city,\s*country|email\s*address\s*or\s*orcid)',
            re.IGNORECASE,
        )
        cleaned_authors = []
        for a in authors:
            name = (a.get('name') or '').strip()
            affs = a.get('affiliations', [])
            if not name:
                continue
            if placeholder_name.match(name):
                # Skip placeholder rows from default publisher templates.
                continue
            if membership_only.match(name):
                continue

            # Remove placeholder affiliation lines but keep valid emails/other lines.
            filtered_affs = [x for x in affs if (not placeholder_affil.search(x)) and (not ieee_template_affil.search(x))]
            a['affiliations'] = filtered_affs
            cleaned_authors.append(a)

        authors = cleaned_authors

        return authors

    def _parse_paragraph(self, p: Paragraph, in_table: bool = False) -> Dict:
        """Parse paragraph into an IR node. Basic implementation, focuses on text."""
        text = ""
        has_math = False
        

        
        try:
            style_name = self._get_style_name(p)
        except:
            style_name = ""
        from .config import MAP_STYLE
        style_cmd = MAP_STYLE.get(style_name or "", "")
        
        # Level detection for sections
        level = None
        if style_cmd == r"\section": level = 1
        elif style_cmd == r"\subsection": level = 2
        elif style_cmd == r"\subsubsection": level = 3
        elif style_name.lower().startswith("heading"):
            try:
                level = int(re.sub(r'[^\d]', '', style_name))
            except:
                level = 1
        
        # We must process the text at the Run-level, to insert math between text nodes
        if p.runs:
            for run in p.runs:
                # Check for math inside this run's vicinity (actually math is sibling to run usually)
                pass
                
        # Better approach: Iterate the XML elements of the paragraph in exact order via recursive descent
        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_m = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        ns_o = "urn:schemas-microsoft-com:office:office"
        ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        ns_v = "urn:schemas-microsoft-com:vml"
        ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"

        def _ty_le_rong_hinh_tu_drawing(drawing_node):
            """Estimate image width ratio from Word drawing extent (EMU)."""
            try:
                ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                extent = drawing_node.find(f".//{{{ns_wp}}}extent")
                if extent is None:
                    return None
                cx = extent.get("cx")
                if not cx:
                    return None
                cx_emu = float(cx)
                if not self.doc.sections:
                    return None
                sec = self.doc.sections[0]
                usable_emu = float(sec.page_width - sec.left_margin - sec.right_margin)
                if usable_emu <= 0:
                    return None
                ratio = cx_emu / usable_emu
                return max(0.2, min(0.95, ratio))
            except Exception:
                return None
        
        def traverse_node(node):
            nonlocal text, has_math
            if not hasattr(node, "tag") or not isinstance(node.tag, str):
                return

            if node.tag == f"{{{ns_m}}}oMathPara":
                has_math = True
                self.total_formulas += 1
                latex_math = self.bo_toan.omml_element_to_latex(node)
                try:
                    omml_str = etree.tostring(node, encoding='unicode')
                    omml_b64 = base64.b64encode(omml_str.encode('utf-8')).decode('utf-8')
                    text += f"\n\\begin{{equation}}\n«OMML:{omml_b64}»\n\\end{{equation}}\n"
                except Exception:
                    text += f"\n\\begin{{equation}}\n{latex_math}\n\\end{{equation}}\n"
                return
            
            if node.tag == f"{{{ns_m}}}oMath":
                has_math = True
                self.total_formulas += 1
                latex_math = self.bo_toan.omml_element_to_latex(node)
                # Capture raw OMML XML for high-fidelity rendering
                try:
                    omml_str = etree.tostring(node, encoding='unicode')
                    omml_b64 = base64.b64encode(omml_str.encode('utf-8')).decode('utf-8')
                    text += f" «OMML:{omml_b64}» "
                except Exception:
                    text += f" ${latex_math}$ "
                return
            elif node.tag == f"{{{ns_w}}}r":
                # Robust run conversion: handle w:t, w:sym, w:instrText, w:br
                run_obj = Run(node, p)
                run_text_acc = ""
                
                # Check formatting from run properties
                is_bold = False
                is_italic = False
                try:
                    is_bold = bool(run_obj.bold)
                    is_italic = bool(run_obj.italic)
                except: pass

                for r_child in node:
                    c_tag = r_child.tag.split("}")[-1] if hasattr(r_child, "tag") else ""
                    full_tag = r_child.tag if hasattr(r_child, "tag") else ""
                    
                    if c_tag == "t":
                        if r_child.text: run_text_acc += r_child.text
                    elif c_tag == "sym":
                        char_hex = r_child.get(f"{{{ns_w}}}char")
                        if char_hex:
                            try:
                                run_text_acc += chr(int(char_hex, 16))
                            except: pass
                    elif c_tag == "instrText":
                        # Skip field-code instructions (e.g., SEQ Table * ARABIC).
                        continue
                    elif c_tag == "br":
                        run_text_acc += "\n"
                    elif c_tag == "tab":
                        run_text_acc += "\t"
                    elif c_tag == "drawing" or c_tag == "object" or c_tag == "pict":
                        # Process inline graphic elements by recursing into them
                        # But we need to flush current accumulated text first
                        if run_text_acc:
                            # In word2word mode, preserve raw Unicode text (no LaTeX escaping).
                            text += loc_ky_tu(run_text_acc) if self.mode != "word2word" else run_text_acc
                            run_text_acc = ""
                        traverse_node(r_child)
                
                if run_text_acc:
                    # In word2word mode, preserve raw Unicode text instead of LaTeX-escaping.
                    run_text = loc_ky_tu(run_text_acc) if self.mode != "word2word" else run_text_acc
                    if run_text.strip() and level is None and self.mode != "word2word":
                        if is_bold and is_italic:
                            run_text = f"\\textbf{{\\textit{{{run_text}}}}}"
                        elif is_bold:
                            run_text = f"\\textbf{{{run_text}}}"
                        elif is_italic:
                            run_text = f"\\textit{{{run_text}}}"
                    text += run_text
                return 
            elif node.tag == f"{{{ns_w}}}object":
                ole_obj = node.find(f".//{{{ns_o}}}OLEObject")
                if ole_obj is not None:
                    prog_id = ole_obj.get("ProgID", "")
                    if "Equation" in prog_id:
                        r_id = ole_obj.get(f"{{{ns_r}}}id")
                        if r_id:
                            try:
                                rel = p.part.rels[r_id]
                                ole_bin = rel.target_part.blob
                                latex = ole_equation_to_latex(ole_bin)
                                if latex:
                                    has_math = True
                                    self.total_formulas += 1
                                    text += f" ${latex}$ "
                                    return
                            except Exception:
                                pass
                # If OLE equation conversion failed, keep walking children so
                # embedded preview images can be extracted as a fallback.
                for child in node:
                    traverse_node(child)
                return
            elif node.tag == f"{{{ns_w}}}pict":
                imagedata = node.find(f".//{{{ns_v}}}imagedata")
                if imagedata is not None:
                    r_id = imagedata.get(f"{{{ns_r}}}id")
                    if r_id:
                        try:
                            rel = p.part.rels[r_id]
                            latex_path = self._save_image_from_relationship(rel)
                            if not latex_path:
                                return
                            
                            if in_table:
                                img_opts = self._includegraphics_options("\\linewidth")
                                latex_img = f"\n\\begin{{center}}\n\\includegraphics[{img_opts}]{{{latex_path}}}\n\\end{{center}}\n"
                            else:
                                self.dem_anh += 1
                                img_opts = self._includegraphics_options("\\columnwidth")
                                latex_img = f"\n\\begin{{figure}}[H]\n\\centering\n\\includegraphics[{img_opts}]{{{latex_path}}}\n\\caption{{}}\n\\label{{fig:img_{self.dem_anh}}}\n\\end{{figure}}\n"
                            text += latex_img
                        except Exception:
                            pass
                return
            elif node.tag == f"{{{ns_w}}}drawing":
                blip = node.find(f".//{{{ns_a}}}blip")
                if blip is not None:
                    r_id = blip.get(f"{{{ns_r}}}embed")
                    if r_id:
                        try:
                            rel = p.part.rels[r_id]
                            latex_path = self._save_image_from_relationship(rel)
                            if not latex_path:
                                return
                            if in_table:
                                img_opts = self._includegraphics_options("\\linewidth")
                                latex_img = f"\n\\begin{{center}}\n\\includegraphics[{img_opts}]{{{latex_path}}}\n\\end{{center}}\n"
                            else:
                                self.dem_anh += 1
                                img_opts = self._includegraphics_options("\\columnwidth")
                                latex_img = f"\n\\begin{{figure}}[H]\n\\centering\n\\includegraphics[{img_opts}]{{{latex_path}}}\n\\caption{{}}\n\\label{{fig:img_{self.dem_anh}}}\n\\end{{figure}}\n"
                            text += latex_img
                        except Exception:
                            pass
                return
            elif node.tag == f"{{{ns_a}}}blip":
                r_id = node.get(f"{{{ns_r}}}embed")
                if r_id:
                    try:
                        rel = p.part.rels[r_id]
                        latex_path = self._save_image_from_relationship(rel)
                        if not latex_path:
                            return
                        if in_table:
                            img_opts = self._includegraphics_options("\\linewidth")
                            latex_img = f"\n\\begin{{center}}\n\\includegraphics[{img_opts}]{{{latex_path}}}\n\\end{{center}}\n"
                        else:
                            self.dem_anh += 1
                            img_opts = self._includegraphics_options("\\columnwidth")
                            latex_img = f"\n\\begin{{figure}}[H]\n\\centering\n\\includegraphics[{img_opts}]{{{latex_path}}}\n\\caption{{}}\n\\label{{fig:img_{self.dem_anh}}}\n\\end{{figure}}\n"
                        text += latex_img
                    except Exception:
                        pass
                return
                
            for child in node:
                traverse_node(child)
                
        traverse_node(p._p)
        text = text.strip()

        # Fix tab-layout equations parsed as normal text (prevents \textit and \quad corruption)
        raw_p_text = getattr(p, "text", "")
        if raw_p_text and re.match(r'^\t.*\t\([\w\.\-\*]+\)$', raw_p_text):
            m = re.match(r'^\t(.*)\t\(([\w\.\-\*]+)\)$', raw_p_text)
            if m:
                eq_text = m.group(1).strip()
                eq_num = m.group(2).strip()
                formula_text = eq_text.replace('½', '\\frac{1}{2}')
                formula_text = formula_text.replace('×', '\\times')
                formula_text = formula_text.replace('−', '-')
                text = f"\\begin{{equation}}\n{formula_text}\n\\tag{{{eq_num}}}\n\\end{{equation}}"
                has_math = True
                level = None
        
        # Merge multiple figures inside the same paragraph into a single figure block
        # This handles paragraphs with several inline images that each got their own
        # \begin{figure}...\end{figure} — collapse them into one grouped figure.
        if '\\includegraphics' in text and text.count('\\begin{figure}') > 1:
            try:
                # Remove empty-caption intermediate figure boundaries to merge images
                merge_pattern = (
                    r'\\caption\{\}\n\\label\{[a-zA-Z0-9:_]+\}\n\\end\{figure\}\n\n'
                    r'\\begin\{figure\}\[H\]\n\\centering\n'
                )
                prev = None
                while prev != text:
                    prev = text
                    text = re.sub(merge_pattern, '\n', text)
            except Exception:
                pass

        # Heuristics for headings
        if not level:
            try:
                for pattern, latex_cmd in HEADING_PATTERNS:
                    if not pattern or not isinstance(pattern, str):
                        continue
                    if re.match(pattern, text, re.IGNORECASE):
                        if latex_cmd == r"\section": level = 1
                        elif latex_cmd == r"\subsection": level = 2
                        elif latex_cmd == r"\subsubsection": level = 3
                        break
            except Exception as e:
                # Silently skip bad regex or unexpected text content
                pass

        # IEEE-style heading detection fallback (Roman numerals, letter subsections)
        if not level:
            stripped = text.strip()
            # "I. INTRODUCTION", "II. RELATED WORK", "III. METHODOLOGY", etc.
            if re.match(r'^[IVXLCDM]+\.\s+[A-Z]', stripped) and len(stripped) < 120:
                level = 1
            # "A. Dataset Description", "B. Evaluation Metrics", etc.
            elif re.match(r'^[A-Z]\.\s+\S', stripped) and len(stripped) < 120:
                level = 2
                    
        if level:
            # We strip numbering (e.g., "1. Introduction" -> "Introduction")
            clean_text = re.sub(r'^\s*(?:[A-Z0-9IVX]+(?:\.\d+)*\.?|\d+(?:\.\d+)*)\s+', '', text)
            return {"type": "section", "level": level, "text": clean_text}
            
        # Standard paragraph
        # Ideally, we would preserve bold/italics here. For now, just raw text.
        return {"type": "paragraph", "text": text or "", "has_math": has_math}

    def _lay_gridspan(self, tc) -> int:
        try:
            tcPr = tc.tcPr
            if tcPr is None: return 1
            gridSpan = tcPr.gridSpan
            if gridSpan is None: return 1
            val = gridSpan.get(qn('w:val'))
            return max(1, int(val)) if val else 1
        except:
            return 1

    def _lay_vmerge(self, tc):
        try:
            tcPr = tc.tcPr
            if tcPr is None: return None
            vMerge = tcPr.vMerge
            if vMerge is None: return None
            val = vMerge.get(qn('w:val'))
            return str(val) if val else 'continue'
        except:
            return None

    def _lay_ty_le_rong_bang(self, t: Table):
        """Read table preferred width from Word XML and map to page ratio (0..1]."""
        try:
            tbl_pr = getattr(t._tbl, "tblPr", None)
            if tbl_pr is None:
                return None

            tblw = tbl_pr.find(qn('w:tblW'))
            if tblw is None:
                return None

            w_type = (tblw.get(qn('w:type')) or "").lower()
            w_val = tblw.get(qn('w:w'))
            if not w_val:
                return None

            if w_type == 'pct':
                # In WordprocessingML, pct is stored in fiftieths of a percent.
                ratio = float(w_val) / 5000.0
                return max(0.2, min(0.95, ratio))

            if w_type == 'dxa':
                twips = float(w_val)
                if not self.doc.sections:
                    return None
                sec = self.doc.sections[0]
                usable_twips = (
                    sec.page_width.twips
                    - sec.left_margin.twips
                    - sec.right_margin.twips
                )
                if usable_twips <= 0:
                    return None
                ratio = twips / float(usable_twips)
                return max(0.2, min(0.95, ratio))
        except Exception:
            return None

        return None

    def _parse_table(self, t: Table) -> Dict:
        """Parse table including rowspan (vMerge) and colspan (gridSpan)."""
        tbl = t._tbl
        tr_list = list(tbl.tr_lst)
        width_ratio = self._lay_ty_le_rong_bang(t)

        so_cot = 0
        try:
            grid_cols = tbl.tblGrid.gridCol_lst
            so_cot = len(grid_cols)
        except: pass

        if so_cot <= 0:
            for tr in tr_list:
                so_cot = max(so_cot, len(list(tr.tc_lst)))

        luoi = [[None for _ in range(so_cot)] for _ in range(len(tr_list))]
        meta = {}

        for r, tr in enumerate(tr_list):
            c = 0
            for tc in list(tr.tc_lst):
                while c < so_cot and luoi[r][c] is not None:
                    c += 1
                if c >= so_cot:
                    break

                colspan = self._lay_gridspan(tc)
                vmerge = self._lay_vmerge(tc)

                cell_id = id(tc)
                if vmerge in ('continue', 'cont') and r > 0 and luoi[r - 1][c] is not None:
                    cell_id = meta.get((r - 1, c), {}).get('id', cell_id)

                meta[(r, c)] = {
                    'id': cell_id,
                    'tc': tc,
                    'colspan': colspan,
                    'vmerge': vmerge,
                    'start': not (vmerge in ('continue', 'cont')),
                    'col_start': c,
                }

                for k in range(colspan):
                    if c + k < so_cot:
                        luoi[r][c + k] = cell_id
                        if k > 0:
                            # Horizontal merge continuation cells must not be treated
                            # as independent starts, otherwise content is duplicated.
                            meta[(r, c + k)] = {
                                'id': cell_id,
                                'tc': tc,
                                'colspan': 1,
                                'vmerge': vmerge,
                                'start': False,
                                'col_start': c,
                            }

                c += colspan

        # Compute rowspan
        rowspan_map = {}
        for (r, c), info in list(meta.items()):
            if not info.get('start'):
                continue
            if info.get('col_start', c) != c:
                continue

            cell_id = info['id']
            rowspan = 1
            rr = r + 1
            while rr < len(tr_list):
                info_down = meta.get((rr, c))
                if not info_down or info_down.get('id') != cell_id or info_down.get('start'):
                    break
                rowspan += 1
                rr += 1
            rowspan_map[cell_id] = max(rowspan_map.get(cell_id, 1), rowspan)
            
        # Build logical grid data for IR
        parsed_rows = []
        for r in range(len(tr_list)):
            row_data = []
            c = 0
            while c < so_cot:
                info = meta.get((r, c))
                if not info or not info.get('start') or meta.get((r, c)) != info:
                    # It's a merged cell or empty
                    row_data.append({"type": "empty", "colspan": 1, "rowspan": 1, "text": ""})
                    c += 1
                    continue
                    
                colspan = int(info.get('colspan') or 1)
                cell_id = info['id']
                rowspan = int(rowspan_map.get(cell_id, 1))
                
                # Retrieve pure text for now
                cell_obj = None
                try:
                    cell_obj = t.rows[r].cells[0]
                    for candidate in t.rows[r].cells:
                        if id(candidate._tc) == id(info['tc']):
                            cell_obj = candidate
                            break
                except: pass
                
                text_content = ""
                if cell_obj:
                    cell_texts = []
                    # Standard paragraphs
                    for p in cell_obj.paragraphs:
                        p_data = self._parse_paragraph(p, in_table=True)
                        cell_texts.append(p_data.get("text", ""))
                    
                    # Nested tables
                    for nested_tbl in cell_obj.tables:
                        for row in nested_tbl.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    p_data = self._parse_paragraph(p, in_table=True)
                                    cell_texts.append(p_data.get("text", ""))
                    
                    text_content = "\n".join(cell_texts).strip()
                    if not text_content:
                        text_content = loc_ky_tu((cell_obj.text or "").strip())
                    
                row_data.append({
                    "type": "cell",
                    "text": text_content,
                    "colspan": colspan,
                    "rowspan": rowspan,
                    "is_merged": colspan > 1 or rowspan > 1
                })
                
                c += colspan
                
            parsed_rows.append(row_data)

        # Trích xuất Header (Heuristic đơn giản: lấy row 0 làm header nếu có text)
        is_header = False
        if len(parsed_rows) > 0 and len(tr_list) > 1:
            is_header = any(cell["text"] != "" for cell in parsed_rows[0])
            
        is_floating_word_table = t._tbl.find(f".//{{{W_NAMESPACE}}}tblpPr") is not None

        return {
            "type": "table",
            "rows": len(tr_list), 
            "cols": so_cot,
            "has_header": is_header,
            "data": parsed_rows,
            "width_ratio": width_ratio,
            "is_floating_word_table": bool(is_floating_word_table),
        }
