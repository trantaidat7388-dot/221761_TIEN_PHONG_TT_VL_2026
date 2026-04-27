"""IEEE Word Renderer — converts IR data into an IEEE-formatted Word document.

Strategy (Plan A — Refactor Renderer):
  When an uploaded IEEE template is provided the renderer **clears the template body**
  (preserving page layout / section properties) and **rebuilds all content from scratch**
  using the template's inherent styles.  This avoids the fragile fill-in-place approach
  that broke when IR node count differed from template paragraph count.

  Rich-text formatting embedded as LaTeX commands in the IR (e.g. \\textbf{...})
  is parsed and mapped to python-docx Run properties so that bold/italic/hyperlinks
  are preserved in the output Word document.

  IEEE compliance v2 (Apr 2026):
  - Table borders: top/bottom/insideH only (no left/right/insideV), thin (sz=4)
  - References: 'Heading 5' unnumbered heading + 'references' auto-numbered style
  - Author block: Word columns + 'Author' style paragraphs (not table)
  - All headings prefer template styles over manual formatting
  - Bullet list support via 'bullet list' style
"""

import re
import os
import base64
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from lxml import etree

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table

from .word_loader import mo_tai_lieu_word_co_fallback


# ---------------------------------------------------------------------------
# Regex helpers
# ---------------------------------------------------------------------------
_FIG_PATH_RE = re.compile(r"\\includegraphics(?:\[[^\]]*\])?{([^}]+)}")
_CAPTION_RE = re.compile(r"\\caption{([^}]*)}")
_LABEL_RE = re.compile(r"\\label{[^}]*}")
_FIG_ENV_RE = re.compile(r"\\begin{(figure\*?)}", re.IGNORECASE)

# LaTeX rich-text token patterns — used to convert IR text to Word runs
_RICH_TEXT_TOKENS = re.compile(
    r"(\\textbf\{|\\textit\{|\\texttt\{|\\emph\{|\\underline\{"
    r"|\\textsuperscript\{|\\textsubscript\{"
    r"|\\href\{[^}]*\}\{|\\textcolor(?:\[[^\]]*\])?\{[^}]*\}\{"
    r"|\\colorbox\{[^}]*\}\{"
    r"|\{|\})"
)

# IEEE heading roman numerals
_ROMAN_MAP = [
    (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
    (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
    (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
]

# IEEE conference two-column content width approximation on A4 with current margins.
_IEEE_COLUMN_WIDTH_INCH = 3.45
_IEEE_TWO_COL_GUTTER_INCH = 0.24
_IEEE_BODY_LINE_SPACING = 1.0
_IEEE_BODY_SPACE_AFTER_PT = 2
_IEEE_CAPTION_LINE_SPACING = 1.0


def _to_roman(n: int) -> str:
    n = max(1, int(n))
    parts: list[str] = []
    for value, numeral in _ROMAN_MAP:
        while n >= value:
            parts.append(numeral)
            n -= value
    return "".join(parts)


class IEEEWordRenderer:
    """Render IR output into an IEEE-formatted Word document.

    The renderer supports two modes:
      1) **From-scratch**: creates a new blank document with IEEE-like styling.
      2) **Template-based**: opens an uploaded IEEE .docx template, clears the body
         and rebuilds all content using the template's styles/page-layout.
    """

    def render(
        self,
        ir_data: Dict[str, Any],
        output_path: str,
        image_root_dir: str | None = None,
        ieee_template_path: str | None = None,
    ) -> str:
        temp_files: list[str] = []
        try:
            if ieee_template_path:
                try:
                    doc, temp_files = mo_tai_lieu_word_co_fallback(ieee_template_path)
                except Exception as e:
                    raise RuntimeError(f"Không thể mở mẫu IEEE Word: {e}")
                self._using_uploaded_template = True
            else:
                doc = Document()
                self._configure_ieee_document(doc)
                self._using_uploaded_template = False

            self._heading_counters = {1: 1, 2: 1, 3: 1}
            self._heading_counters_figure = 1
            self._heading_counters_table = 1
            self._figure_index = 0
            self._table_index = 0
            self._section_index = 0
            self._subsection_counters: Dict[int, int] = {}  # section_idx -> sub count
            self._image_root_dir = Path(image_root_dir) if image_root_dir else Path(output_path).parent
            self._available_styles = {s.name for s in doc.styles if getattr(s, "name", None)}
            self._render_metrics: Dict[str, int] = {
                "table_span_full_width": 0,
                "table_span_landscape": 0,
                "table_caption_normalized": 0,
                "figure_caption_normalized": 0,
            }

            metadata = ir_data.get("metadata", {}) or {}
            body_nodes = ir_data.get("body", []) or []
            references = ir_data.get("references", []) or []

            if self._using_uploaded_template:
                self._rebuild_on_uploaded_template(doc, metadata, body_nodes, references)
            else:
                self._add_title_section(doc, metadata)
                authors = metadata.get("authors") or []
                if authors:
                    self._add_authors_table(doc, authors)
                self._start_two_column_body(doc)
                self._add_abstract_and_keywords(doc, metadata)
                self._add_body(doc, body_nodes)
                self._add_references(doc, references)

            self.last_render_metrics = dict(self._render_metrics)

            doc.save(output_path)
            return output_path
        finally:
            for temp_path in temp_files:
                try:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                except Exception:
                    pass

    # ========================================================================
    # Plan B: Find and Replace In-Place (Maintains 100% IEEE Template Layout)
    # ========================================================================

    def _rebuild_on_uploaded_template(
        self,
        doc: Document,
        metadata: Dict[str, Any],
        body_nodes: List[Dict[str, Any]],
        references: List[Dict[str, Any]],
    ) -> None:
        """Rebuild content using uploaded template styles while removing template guide text.

        We preserve document/section settings (margins, columns, page size) by keeping
        the trailing sectPr in the body XML and recreating only user content.
        """
        self._clear_document_body_preserve_layout(doc)
        self._add_title_section(doc, metadata)
        authors = metadata.get("authors") or []
        if authors:
            self._add_authors_table(doc, authors)
        self._start_two_column_body(doc)
        self._add_abstract_and_keywords(doc, metadata)
        self._add_body(doc, body_nodes)
        self._add_references(doc, references)

    def _clear_document_body_preserve_layout(self, doc: Document) -> None:
        """Remove all paragraphs/tables but keep section properties for layout fidelity."""
        body = doc._element.body
        children = list(body)
        sect_pr = None
        for child in reversed(children):
            if child.tag == qn("w:sectPr"):
                sect_pr = deepcopy(child)
                break

        for child in children:
            body.remove(child)

        if sect_pr is not None:
            body.append(sect_pr)

    def _fill_into_existing_template(
        self,
        doc: Document,
        metadata: Dict[str, Any],
        body_nodes: List[Dict[str, Any]],
        references: List[Dict[str, Any]],
    ) -> None:
        """Fills data into the IEEE template by finding key anchor paragraphs
        and replacing their text while preserving their formatting, columns, and section breaks."""

        # 1. Fill Metadata (Title, Authors, Abstract, Keywords)
        self._replace_metadata(doc, metadata)

        # 2. Fill Body & References
        self._replace_body_and_refs(doc, body_nodes, references)

    def _replace_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        title = (metadata.get("title") or "").strip()
        abstract = self._latex_to_plain(metadata.get("abstract") or "")
        keywords = [self._latex_to_plain(str(k)) for k in (metadata.get("keywords") or [])]
        authors = metadata.get("authors") or []

        authors_counter = 0

        # Pass 1: Parse all paragraphs and table cells to find anchors
        for p in self._iter_all_paragraphs(doc):
            text_upper = p.text.upper()

            # --- Title ---
            if "PAPER TITLE" in text_upper:
                if title:
                    self._replace_paragraph_text_keep_formatting(p, self._latex_to_plain(title), uppercase=True)
                continue

            # --- Abstract ---
            if p.text.strip().startswith("Abstract—") or p.text.strip().startswith("Abstract-"):
                if abstract:
                    self._replace_paragraph_text_keep_formatting(p, f"Abstract—{abstract}", italic=True)
                continue
            
            if p.text.strip() == "ABSTRACT":
                # some old templates have just "ABSTRACT"
                p.clear()
                continue

            # --- Keywords ---
            if p.text.strip().startswith("Keywords—") or p.text.strip().startswith("Index Terms—"):
                if keywords:
                    kw_str = ", ".join(keywords) + "."
                    self._replace_paragraph_text_keep_formatting(p, f"Index Terms—{kw_str}", italic=True)
                continue

            # --- Authors ---
            if "GIVEN NAME SURNAME" in text_upper or "NAME OF ORGANIZATION" in text_upper:
                # We expect templates to have up to 6 author blocks.
                # Since multiple lines belong to 1 author block, we try to clear or fill them.
                if "GIVEN NAME SURNAME" in text_upper:
                    if authors_counter < len(authors):
                        author_name = self._latex_to_plain(authors[authors_counter].get("name", ""))
                        self._replace_paragraph_text_keep_formatting(p, author_name)
                    else:
                        p.clear() # Clear unused author slots
                    authors_counter += 1
                elif "DEPT. NAME OF ORGANIZATION" in text_upper or "(OF AFFILIATION)" in text_upper or "NAME OF ORGANIZATION" in text_upper:
                    # Very rough heuristic: replace affiliation lines.
                    p.clear() # Cleared for simplicity, or we could meticulously map affiliations.
                elif "CITY, COUNTRY" in text_upper or "EMAIL ADDRESS OR ORCID" in text_upper:
                    p.clear() # Clear placeholder metadata

    def _replace_body_and_refs(
        self,
        doc: Document,
        body_nodes: List[Dict[str, Any]],
        references: List[Dict[str, Any]],
    ) -> None:
        """Finds the INTRODUCTION heading, anchors there, and replaces the rest of the body."""
        intro_idx = -1
        ref_idx = -1
        all_paras = list(doc.paragraphs)

        # Locate INTRODUCTION
        for i, p in enumerate(all_paras):
            text_upper = p.text.strip().upper()
            if "INTRODUCTION" in text_upper and len(text_upper) < 30:
                intro_idx = i
            if "REFERENCES" in text_upper and len(text_upper) < 30:
                ref_idx = i

        if intro_idx == -1:
            return  # Could not find anchor

        # Clear existing text from Intro downwards (but NOT the paragraph objects to keep section breaks)
        end_idx = ref_idx if ref_idx != -1 else len(all_paras)
        for i in range(intro_idx + 1, end_idx):
            all_paras[i].clear()

        # Anchor paragraph for body insertions is the one right after INTRODUCTION
        anchor_p = all_paras[intro_idx + 1] if intro_idx + 1 < end_idx else all_paras[intro_idx]
        
        # Replace the INTRODUCTION heading itself with properly numbered version
        first_section_text = body_nodes[0].get("text", "INTRODUCTION") if body_nodes and body_nodes[0].get("type") == "section" else "INTRODUCTION"
        self._replace_paragraph_text_keep_formatting(all_paras[intro_idx], f"I. {first_section_text.upper()}")

        # Insert body nodes BEFORE the anchor to maintain layout flow properties
        nodes_to_insert = body_nodes[1:] if body_nodes and body_nodes[0].get("type") == "section" else body_nodes
        
        for node in nodes_to_insert:
            node_type = node.get("type", "")

            if node_type == "section":
                level = int(node.get("level", 1) or 1)
                text = self._latex_to_plain(node.get("text") or "")
                if text:
                    self._insert_heading_before(anchor_p, text, level)
                continue

            if node_type == "table":
                self._insert_table_before(doc, anchor_p, node)
                continue

            if node_type == "paragraph":
                text = node.get("text") or ""
                # Handle equations slightly differently
                if r"\begin{equation}" in text or r"$$" in text or r"\[" in text or "«OMML:" in text:
                    self._insert_equation_before(doc, anchor_p, text)
                else:
                    self._insert_rich_paragraph_before(anchor_p, text)
                continue

        # Now handle References
        if ref_idx != -1:
            for i in range(ref_idx + 1, len(all_paras)):
                all_paras[i].clear() # Clear placeholder refs
            
            ref_anchor = all_paras[ref_idx + 1] if ref_idx + 1 < len(all_paras) else all_paras[ref_idx]
            for i, ref in enumerate(references):
                ref_text = self._latex_to_plain(ref.get("text") or "")
                if not ref_text.startswith("["):
                    ref_text = f"[{i+1}] {ref_text}"
                self._insert_paragraph_before(ref_anchor, ref_text, size=8)

    # -----------------------------------------------------------------------
    # Insertion Helpers (Inserting BEFORE an anchor keeps the anchor's sectPr safe)
    # -----------------------------------------------------------------------

    def _insert_paragraph_before(self, anchor_p, text: str, size: int = 10, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        new_p = anchor_p.insert_paragraph_before()
        new_p.alignment = align
        run = new_p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        return new_p

    def _insert_rich_paragraph_before(self, anchor_p, raw_text: str):
        # 1. Strip figures if any
        raw_text, figures = self._extract_figures_from_text(raw_text)
        
        # 2. Add the text paragraph
        if raw_text.strip():
            new_p = anchor_p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._add_rich_runs(new_p, raw_text)
            
        # 3. Add any figures found
        for fig in figures:
            self._insert_figure_before(anchor_p, fig)

    def _insert_equation_before(self, doc: Document, anchor_p, raw_text: str) -> None:
        """Render an equation node as centered OMML or stylized text before an anchor using an invisible table for layout."""
        # Check for OMML marker first
        omml_match = re.search(r"«OMML:([A-Za-z0-9+/=]+)»", raw_text)
        tag_match = re.search(r"\\tag\{([^}]*)\}", raw_text)
        if tag_match:
            # Normalize tag content to avoid hidden newlines that break as "(1" then ")".
            tag_text = re.sub(r"\s+", "", tag_match.group(1) or "")
            eq_num = f"({tag_text})" if tag_text else ""
        else:
            eq_num = ""

        if not eq_num:
            # Fallback to simple paragraph if no equation number
            p = anchor_p.insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            if omml_match:
                try:
                    b64_xml = omml_match.group(1)
                    self._insert_omml_to_paragraph(p, b64_xml)
                except Exception:
                    clean = re.sub(r"«OMML:([A-Za-z0-9+/=]+)»", "", raw_text)
                    run = p.add_run(self._latex_math_to_readable(clean).strip())
                    run.font.name = "Times New Roman"
                    run.font.size = self._equation_font_size_for_text(run.text)
                    run.italic = True
            else:
                clean = re.sub(r"\\begin\{equation\*?\}", "", raw_text)
                clean = re.sub(r"\\end\{equation\*?\}", "", clean)
                clean = re.sub(r"\\\[", "", clean)
                clean = re.sub(r"\\\]", "", clean)
                run = p.add_run(self._latex_math_to_readable(clean).strip())
                run.font.name = "Times New Roman"
                run.font.size = self._equation_font_size_for_text(run.text)
                run.italic = True
            return

        # For plain-text LaTeX equations, use IEEE tab-stop layout to keep equation centered
        # and number right-aligned without table artifacts.
        if not omml_match:
            clean = re.sub(r"\\begin\{equation\*?\}", "", raw_text)
            clean = re.sub(r"\\end\{equation\*?\}", "", clean)
            clean = re.sub(r"\\tag\{([^}]*)\}", "", clean)
            clean = re.sub(r"\\\[", "", clean)
            clean = re.sub(r"\\\]", "", clean)
            eq_text = self._latex_math_to_readable(clean).strip()
            p = anchor_p.insert_paragraph_before()
            self._apply_ieee_equation_tab_layout(p, doc, eq_text, eq_num)
            return

        # OMML equations keep table-based layout for robust math XML embedding.
        parent = anchor_p._p.getparent()
        try:
            temp_table = doc.add_table(rows=1, cols=2)
        except TypeError:
            temp_table = doc.add_table(rows=1, cols=2, width=Inches(3.3))

        col_width = self._get_current_column_width_inch(doc)
        number_width = 0.34
        equation_width = max(2.8, col_width - number_width)
        temp_table.columns[0].width = Inches(equation_width)
        temp_table.columns[1].width = Inches(number_width)

        # Clear borders
        try:
            for row in temp_table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tc_pr = tc.get_or_add_tcPr()
                    tc_borders = OxmlElement("w:tcBorders")
                    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        edge = OxmlElement(f"w:{border_name}")
                        edge.set(qn("w:val"), "nil")
                        tc_borders.append(edge)
                    tc_pr.append(tc_borders)
        except Exception:
            pass

        # First cell: equation
        cell_eq = temp_table.cell(0, 0)
        self._set_cell_no_wrap(cell_eq)
        p_eq = cell_eq.paragraphs[0]
        p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if omml_match:
            try:
                b64_xml = omml_match.group(1)
                self._insert_omml_to_paragraph(p_eq, b64_xml)
            except Exception:
                clean = re.sub(r"«OMML:([A-Za-z0-9+/=]+)»", "", raw_text)
                run = p_eq.add_run(self._latex_math_to_readable(clean).strip())
                run.font.name = "Times New Roman"
                run.font.size = self._equation_font_size_for_text(run.text)
                run.italic = True
        else:
             clean = re.sub(r"\\begin\{equation\*?\}", "", raw_text)
             clean = re.sub(r"\\end\{equation\*?\}", "", clean)
             clean = re.sub(r"\\tag\{([^}]*)\}", "", clean)
             clean = re.sub(r"\\\[", "", clean)
             clean = re.sub(r"\\\]", "", clean)
             run = p_eq.add_run(self._latex_math_to_readable(clean).strip())
             run.font.name = "Times New Roman"
             run.font.size = self._equation_font_size_for_text(run.text)
             run.italic = True

        # Second cell: number
        cell_num = temp_table.cell(0, 1)
        self._set_cell_no_wrap(cell_num)
        p_num = cell_num.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_num.paragraph_format.space_before = Pt(0)
        run_num = p_num.add_run(eq_num)
        run_num.font.name = "Times New Roman"
        run_num.font.size = Pt(10)
        
        # Insert table before anchor
        tbl_xml = temp_table._element
        parent.remove(tbl_xml)
        anchor_p._p.addprevious(tbl_xml)
        
        # Add a small spacing paragraph
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(2)
        parent.remove(p_space._element)
        anchor_p._p.addprevious(p_space._element)

    def _insert_heading_before(self, anchor_p, text: str, level: int):
        text = text.upper() if level == 1 else text
        text = re.sub(r'^[IVX]+\.\s+', '', text)
        text = re.sub(r'^[A-Z]\.\s+', '', text)
        text = re.sub(r'^\d+\.\s+', '', text)

        if level == 1:
            num = _to_roman(self._heading_counters[1])
            self._heading_counters[1] += 1
            self._heading_counters[2] = 1 # reset
            fmt_text = f"{num}. {text}"
            p = anchor_p.insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(fmt_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        elif level == 2:
            num = chr(64 + self._heading_counters[2])
            self._heading_counters[2] += 1
            fmt_text = f"{num}. {text}"
            p = anchor_p.insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(fmt_text)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        else:
            p = anchor_p.insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    def _insert_table_before(self, doc, anchor_p, node: Dict[str, Any]):
        label = f"TABLE {_to_roman(self._heading_counters_table)}"
        self._heading_counters_table += 1
        caption = self._latex_to_plain(node.get("caption") or "")
        
        # Caption above table
        cap_p = anchor_p.insert_paragraph_before()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"{label}. {caption}" if caption else label)
        cap_run.font.name = "Times New Roman"
        cap_run.font.size = Pt(8)

        # We cannot easily insert a table *before* a paragraph in standard python-docx API.
        # But we can access the parent element and insert the tbl element.
        parent = anchor_p._p.getparent()
        table_data = node.get("data", [])
        if not table_data:
            return

        # Use an existing document property to create a detached table
        try:
            temp_table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
        except TypeError:
            temp_table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data), width=Inches(6.0))
        try:
            temp_table.style = "Table Grid"
        except Exception:
            pass
        
        # Fill table
        for r_idx, row_data in enumerate(table_data):
            for c_idx, cell_data in enumerate(row_data):
                if c_idx < len(temp_table.rows[r_idx].cells):
                    cell = temp_table.cell(r_idx, c_idx)
                    cell.text = self._latex_to_plain(cell_data.get("text", ""))
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(8)
                            
        # Move XML element before anchor_p
        tbl_xml = temp_table._element
        parent.remove(tbl_xml) # detach from end
        anchor_p._p.addprevious(tbl_xml) # insert before anchor
        
    def _insert_figure_before(self, anchor_p, fig_data: Dict[str, Any]):
        label = f"Fig. {self._heading_counters_figure}"
        self._heading_counters_figure += 1
        
        # Insert images
        paths = fig_data.get("paths", [])
        if not paths and "path" in fig_data:
            paths = [fig_data["path"]]
            
        for img_path in paths:
            p = anchor_p.insert_paragraph_before()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            resolved_path = self._resolve_image_path(img_path)
            if resolved_path:
                run = p.add_run()
                try:
                    run.add_picture(str(resolved_path), width=Inches(3.0))
                except Exception:
                    pass
                
        # Caption below figure group
        cap_p = anchor_p.insert_paragraph_before()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"{label}. {fig_data.get('caption', '')}")
        cap_run.font.name = "Times New Roman"
        cap_run.font.size = Pt(8)

    # -----------------------------------------------------------------------
    # Utilities
    # -----------------------------------------------------------------------

    def _replace_paragraph_text_keep_formatting(self, p, new_text: str, uppercase=False, italic=False):
        """Replaces text in a paragraph while attempting to preserve the first run's font styling."""
        if uppercase:
            new_text = new_text.upper()
            
        if not p.runs:
            p.clear()
            new_run = p.add_run(new_text)
            new_run.font.name = "Times New Roman"
            return

        # Steal format of first readable run
        base_run = next((r for r in p.runs if r.text.strip()), p.runs[0])
        font_name = base_run.font.name or "Times New Roman"
        font_size = base_run.font.size
        is_bold = base_run.bold

        p.clear()
        
        # Determine rich text components
        # A simple adaptation since we stripped full rich text to avoid over-complicating
        # for standard titles/abstracts which don't usually have deep nested LaTeX.
        new_run = p.add_run(new_text)
        new_run.font.name = font_name
        new_run.font.size = font_size
        if is_bold:
            new_run.bold = True
        if italic:
            new_run.italic = True

    def _iter_all_paragraphs(self, doc: Document):
        """Yields all paragraphs in document, including those inside tables."""
        for p in doc.paragraphs:
            yield p
        for t in doc.tables:
            for r in t.rows:
                for c in t.columns:
                    for cell in r.cells:
                        for p in cell.paragraphs:
                            yield p

    def _extract_figures_from_text(self, text: str) -> Tuple[str, List[Dict[str, Any]]]:
        """Extracts figure references from text. Supports both \begin{figure} and \begin{figure*}."""
        figures = []
        # Updated pattern: allow figure or figure*
        pattern_block = r'\\begin\{figure\*?\}(.*?)\\end\{figure\*?\}'
        
        for block_match in re.finditer(pattern_block, text, re.DOTALL):
            block_content = block_match.group(1)
            
            # Extract all \includegraphics paths from this block
            paths = []
            img_pattern = r'\\includegraphics(?:\[.*?\])?\{(.*?)\}'
            for img_match in re.finditer(img_pattern, block_content):
                paths.append(img_match.group(1))
                
            # Extract caption
            caption = ""
            cap_pattern = r'\\caption\{(.*?)\}'
            cap_match = re.search(cap_pattern, block_content, re.DOTALL)
            if cap_match:
                caption = cap_match.group(1).strip()
                
            if paths:
                figures.append({
                    "paths": paths,
                    "caption": caption
                })
                
        clean_text = re.sub(pattern_block, '', text, flags=re.DOTALL).strip()
        return clean_text, list(figures)

    def _add_title_section(self, doc: Document, metadata: Dict[str, Any]) -> None:
        title = (metadata.get("title") or "").strip()
        if not title: return
        p = doc.add_paragraph()
        title_style = self._pick_style_name(["paper title", "papertitle", "Title"])
        if title_style:
            try:
                p.style = title_style
            except Exception:
                pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self._latex_to_plain(title))
        # With uploaded IEEE templates, preserve style-defined typography to stay close to IEEEtran output.
        if not self._using_uploaded_template:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(24)

    def _add_abstract_and_keywords(self, doc: Document, metadata: Dict[str, Any]) -> None:
        abstract = self._latex_to_plain(metadata.get("abstract") or "")
        # Strip any existing label prefix to avoid double labeling (e.g. "Abstract—Abstract. ...")
        # Requires separator (dot, dash, colon, em-dash) after label word to avoid false positives
        abstract = re.sub(r"^\s*(?:abstract|t[oó]m\s+t[aắ]t)\s*[:.\u2013\u2014\-]+\s*", "", abstract, flags=re.IGNORECASE).strip()
        if abstract:
            abs_style = self._pick_style_name(["Abstract", "abstract"])
            if abs_style and self._using_uploaded_template:
                # Use template Abstract style — it has built-in formatting
                p = doc.add_paragraph()
                try:
                    p.style = abs_style
                except Exception:
                    pass
                # Keep template typography; only inject semantic text.
                p.add_run(f"Abstract—{abstract}")
            else:
                p = doc.add_paragraph(f"Abstract—{abstract}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = _IEEE_BODY_LINE_SPACING
                p.paragraph_format.space_after = Pt(3)

        keywords = metadata.get("keywords") or []
        if keywords:
            kw = ", ".join(self._latex_to_plain(k) for k in keywords if str(k).strip())
            kw_style = self._pick_style_name(["Keywords", "keywords", "KeyWords"])
            if kw_style and self._using_uploaded_template:
                p = doc.add_paragraph()
                try:
                    p.style = kw_style
                except Exception:
                    pass
                p.add_run(f"Keywords—{kw}.")
            else:
                p = doc.add_paragraph(f"Index Terms—{kw}.")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = _IEEE_BODY_LINE_SPACING
                p.paragraph_format.space_after = Pt(4)

    def _add_authors_table(self, doc: Document, authors: List[Dict[str, Any]]) -> None:
        """Render author block using IEEE 'Author' style paragraphs with Word columns.

        The IEEE Word template uses multi-column section layout (not a table)
        with the 'Author' paragraph style.  Each author block contains:
          line 1: Name
          line 2: Organization / Affiliation
          line 3: City, Country
          line 4: email or ORCID
        """
        if not authors:
            return

        author_style = self._pick_style_name(["Author", "author"])

        # Build author paragraph text blocks
        used_emails: set[str] = set()
        author_blocks: list[Dict[str, List[str] | str]] = []
        for author in authors:
            name = self._clean_author_name(str(author.get("name") or ""))
            raw_affs = [
                self._latex_to_plain(str(x)).strip()
                for x in (author.get("affiliations") or [])
                if str(x).strip() and str(x).strip() not in {"*", "†", "‡"}
            ]
            explicit_email = self._latex_to_plain(str(author.get("email") or "")).strip()

            # Split affiliation and email candidates (some sources put all authors' emails into affiliations).
            email_candidates: list[str] = []
            affs: list[str] = []
            for line in raw_affs:
                if self._looks_like_email(line):
                    email_candidates.append(line)
                else:
                    affs.append(line)

            # Keep affiliation compact in IEEE author row (avoid long repeated blocks).
            affs = affs[:2]

            email = explicit_email if self._looks_like_email(explicit_email) else ""
            if not email:
                email = self._pick_best_email_for_author(name, email_candidates, used_emails)
            if email:
                used_emails.add(email.lower())

            org1 = ""
            org2 = ""
            city_country = ""
            if affs:
                org1, city_country = self._split_org_city_from_affiliation(affs[0])
                if len(affs) > 1:
                    org2 = affs[1]

            # IEEE-like fixed profile lines under name to keep columns visually aligned.
            profile_lines = [org1, org2, city_country, email]

            if name:
                author_blocks.append({"name": name, "lines": profile_lines})

        if not author_blocks:
            return

        # Determine IEEE-like author grid:
        # - up to 3 authors: single row, natural columns
        # - 4+ authors: 3 columns; trailing row with 1 author is centered
        if len(author_blocks) <= 3:
            cols = max(1, len(author_blocks))
        else:
            cols = 3

        # Always use invisible-bordered table (it prevents continuous section break layout issues)
        rows_count = (len(author_blocks) + cols - 1) // cols
        table = doc.add_table(rows=rows_count, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Build placement positions (row, col) for each author block.
        positions: list[tuple[int, int]] = []
        total = len(author_blocks)
        if cols < 3:
            for idx in range(total):
                positions.append((idx // cols, idx % cols))
        else:
            first_row_count = min(3, total)
            for i in range(first_row_count):
                positions.append((0, i))

            remaining = total - first_row_count
            offset = 0
            row = 1
            while remaining > 0:
                row_count = min(3, remaining)
                if row_count == 1:
                    row_cols = [1]   # center single trailing author
                elif row_count == 2:
                    row_cols = [0, 2]  # visually centered pair
                else:
                    row_cols = [0, 1, 2]
                for c in row_cols:
                    positions.append((row, c))
                remaining -= row_count
                offset += row_count
                row += 1

        for idx, block in enumerate(author_blocks):
            r, c = positions[idx]
            cell = table.cell(r, c)
            cell.text = ""
            name_text = str(block.get("name") or "")
            lines = [str(x) for x in (block.get("lines") or []) if str(x).strip()]

            # IEEE template expects one Author-style paragraph per author block,
            # with line breaks inside (not multiple paragraphs), otherwise spacing gets too loose.
            p_block = cell.paragraphs[0]
            p_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if author_style:
                try:
                    p_block.style = author_style
                except Exception:
                    pass

            block_lines: List[str] = []
            if name_text:
                block_lines.append(name_text)
            block_lines.extend(lines)

            for i, line in enumerate(block_lines):
                run = p_block.add_run(line)
                if self._using_uploaded_template:
                    # Keep template's base font sizing; only keep affiliation lines italic like sample.
                    run.italic = i in (1, 2)
                else:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10 if i == 0 else 8)
                    run.italic = i in (1, 2)
                if i < len(block_lines) - 1:
                    run.add_break(WD_BREAK.LINE)

        # Hide borders
        try:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tc_pr = tc.get_or_add_tcPr()
                    tc_borders = OxmlElement("w:tcBorders")
                    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        edge = OxmlElement(f"w:{border_name}")
                        edge.set(qn("w:val"), "nil")
                        tc_borders.append(edge)
                    tc_pr.append(tc_borders)
        except Exception:
            pass

        doc.add_paragraph("")

    def _looks_like_email(self, text: str) -> bool:
        return bool(re.fullmatch(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", (text or "").strip()))

    def _pick_best_email_for_author(self, author_name: str, candidates: List[str], used_emails: set[str]) -> str:
        """Pick one likely email for an author from pooled candidates without reusing the same email."""
        clean_candidates = [c.strip() for c in candidates if self._looks_like_email(c)]
        if not clean_candidates:
            return ""

        tokens = [t.lower() for t in re.findall(r"[A-Za-z]+", author_name or "") if len(t) >= 3]

        # Prefer candidate whose local-part matches any author token and is not already used.
        for c in clean_candidates:
            key = c.lower()
            if key in used_emails:
                continue
            local = key.split("@", 1)[0]
            if any(t in local for t in tokens):
                return c

        # Fallback to first non-used candidate.
        for c in clean_candidates:
            if c.lower() not in used_emails:
                return c

        return ""

    def _split_org_city_from_affiliation(self, affiliation: str) -> Tuple[str, str]:
        """Split affiliation into organization and city-country segments for IEEE author lines."""
        text = (affiliation or "").strip()
        if not text:
            return "", ""

        parts = [p.strip() for p in text.split(",") if p.strip()]
        if len(parts) >= 3:
            org = parts[0]
            city_country = ", ".join(parts[1:])
            return org, city_country
        if len(parts) == 2:
            return parts[0], parts[1]
        return text, ""

    def _clean_author_name(self, raw_name: str) -> str:
        text = (raw_name or "").strip()
        if not text:
            return ""
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r",\s*(member|senior\s+member|fellow|student\s+member|life\s+fellow)\s*,?\s*ieee\.?$", "", text, flags=re.IGNORECASE)
        text = re.sub(r",\s*ieee\.?$", "", text, flags=re.IGNORECASE)
        return text.strip(" ,")

    # NOTE: To prevent API breakage, `_add_body` is retained for calls when no uploaded template is used.
    def _add_body(self, doc: Document, body_nodes: List[Dict[str, Any]]) -> None:
        for node in body_nodes:
            node_type = node.get("type", "")

            if node_type == "section":
                level = int(node.get("level", 1) or 1)
                level = max(1, min(3, level))
                text = self._latex_to_plain(node.get("text") or "")
                if text:
                    self._add_ieee_heading(doc, text, level)
                continue

            if node_type == "table":
                table_layout_mode = self._select_table_layout_mode(doc, node)
                if table_layout_mode == "landscape":
                    self._render_metrics["table_span_landscape"] += 1
                    self._start_landscape_single_column_block(doc)
                    self._add_table_node(doc, node, force_full_width=True)
                    self._resume_portrait_two_column_block(doc)
                elif table_layout_mode == "full":
                    self._render_metrics["table_span_full_width"] += 1
                    self._start_single_column_block(doc)
                    self._add_table_node(doc, node, force_full_width=True)
                    self._resume_two_column_block(doc)
                else:
                    self._add_table_node(doc, node, force_full_width=False)
                continue

            if node_type == "list":
                self._add_list_node(doc, node)
                continue

            if node_type == "paragraph":
                raw_text = str(node.get("text") or "")

                # Detect bullet/numbered lists embedded as LaTeX
                if "\\begin{itemize}" in raw_text or "\\begin{enumerate}" in raw_text:
                    self._add_latex_list_node(doc, raw_text)
                    continue

                if "\\begin{figure" in raw_text or _FIG_PATH_RE.search(raw_text):
                    text_without_fig, figures = self._extract_figures_from_text(raw_text)
                    for fig in figures:
                        fig_tex = (
                            "\\begin{figure}[H]\n"
                        )
                        for path in fig.get("paths", []):
                            fig_tex += f"\\includegraphics{{{path}}}\n"
                        fig_tex += f"\\caption{{{fig.get('caption', '')}}}\n\\end{{figure}}"
                        self._add_figure_node(doc, fig_tex)
                    if text_without_fig.strip():
                        p = doc.add_paragraph()
                        body_style = self._pick_style_name(["BodyText", "Body Text", "Normal"])
                        if body_style:
                            try:
                                p.style = body_style
                            except Exception:
                                pass
                        self._apply_body_paragraph_style(p, indent_first_line=True)
                        self._add_rich_runs(p, text_without_fig)
                    continue

                if self._is_equation_like_paragraph(raw_text):
                    self._add_equation_node(doc, raw_text)
                    continue

                # Regular paragraph — render with rich text
                cleaned = self._latex_to_plain(raw_text)
                if cleaned:
                    p = doc.add_paragraph()
                    body_style = self._pick_style_name(["BodyText", "Body Text", "Normal"])
                    if body_style:
                        try:
                            p.style = body_style
                        except Exception:
                            pass
                    self._apply_body_paragraph_style(p, indent_first_line=True)
                    self._add_rich_runs(p, raw_text)

    def _add_ieee_heading(self, doc: Document, text: str, level: int) -> None:
        """Add an IEEE-style heading with proper numbering.

        When an uploaded template is present, prefer the template's built-in
        heading styles (Heading 1, Heading 2, Heading 3) which already have
        correct font, spacing, and numbering definitions.
        """
        clean = self._latex_to_plain(text)
        if not clean:
            return

        # Strip existing numbering from text (e.g., "I. Introduction" -> "Introduction")
        clean = re.sub(r"^[A-Z0-9IVX]+\.\s+", "", clean)
        clean = re.sub(r"^\d+\.\d*\s*", "", clean)

        p = doc.add_paragraph()

        # Try to use template heading styles
        heading_style_name = f"Heading {level}"
        style_applied = False
        if heading_style_name in self._available_styles:
            try:
                p.style = doc.styles[heading_style_name]
                style_applied = True
                self._remove_paragraph_numbering(p)
            except Exception:
                pass

        if level == 1:
            if not style_applied:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(6)
            # IEEE Level 1: "I. INTRODUCTION"
            self._section_index += 1
            self._subsection_counters[self._section_index] = 0
            numbered_text = f"{_to_roman(self._section_index)}. {clean.upper()}"
            run = p.add_run(numbered_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            if not style_applied:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
            # IEEE Level 2: "A. Subsection Title"
            if self._section_index in self._subsection_counters:
                self._subsection_counters[self._section_index] += 1
            else:
                self._subsection_counters[self._section_index] = 1
            sub_idx = self._subsection_counters[self._section_index]
            letter = chr(ord("A") + sub_idx - 1) if sub_idx <= 26 else str(sub_idx)
            numbered_text = f"{letter}. {clean}"
            run = p.add_run(numbered_text)
            run.bold = True
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        else:
            if not style_applied:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
            # IEEE Level 3: italic, not numbered
            run = p.add_run(clean)
            run.bold = False
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    def _add_table_node(self, doc: Document, node: Dict[str, Any], force_full_width: bool = False) -> None:
        """Render a table IR node as a real Word table."""
        self._table_index += 1
        caption = self._normalize_table_caption(self._latex_to_plain(node.get("caption") or ""))
        label = f"TABLE {_to_roman(self._table_index)}"

        # IEEE table caption goes ABOVE the table
        self._add_table_caption(doc, label, caption)

        cols = int(node.get("cols", 1) or 1)
        rows_data = node.get("data", []) or []
        if cols <= 0 or not rows_data:
            return

        table = doc.add_table(rows=len(rows_data), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.autofit = True
        except Exception:
            pass

        # Apply table style — prefer 'Table Grid' for basic structure
        table_style_name = self._resolve_table_style_name(doc)
        if table_style_name:
            try:
                table.style = table_style_name
            except Exception:
                pass
        # Apply IEEE-correct borders (top/bottom/insideH only, thin lines)
        self._force_table_borders(table)
        target_width = self._current_table_target_width_inch(doc, force_full_width)
        self._set_table_width(table, cols, target_width)

        # Try to find the 'table col head' style for header cells
        table_head_style = self._pick_style_name(["table col head", "Table Col Head"])

        for r_idx, row in enumerate(rows_data):
            c_idx = 0
            for cell_data in row:
                if c_idx >= cols:
                    break

                cell_type = cell_data.get("type")
                colspan = max(1, int(cell_data.get("colspan", 1) or 1))
                rowspan = max(1, int(cell_data.get("rowspan", 1) or 1))
                text = self._clean_cell_text(cell_data.get("text") or "")

                if cell_type == "empty":
                    c_idx += 1
                    continue

                target_cell = table.cell(r_idx, c_idx)
                target_cell.text = text

                # Apply formatting to cell
                is_header = (r_idx == 0)
                for paragraph in target_cell.paragraphs:
                    # Apply 'table head' style to header cells if available
                    if is_header and table_head_style:
                        try:
                            paragraph.style = table_head_style
                        except Exception:
                            pass
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.alignment = self._determine_cell_alignment(text, is_header=is_header)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(8)

                # Handle column merge
                end_col = min(cols - 1, c_idx + colspan - 1)
                if end_col > c_idx:
                    try:
                        table.cell(r_idx, c_idx).merge(table.cell(r_idx, end_col))
                    except Exception:
                        pass

                # Handle row merge
                end_row = min(len(rows_data) - 1, r_idx + rowspan - 1)
                if end_row > r_idx:
                    try:
                        table.cell(r_idx, c_idx).merge(table.cell(end_row, c_idx))
                    except Exception:
                        pass

                c_idx += colspan

        # Bold first row (header)
        if rows_data:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        doc.add_paragraph("")  # spacing after table

    def _add_figure_node(self, doc: Document, latex_figure_text: str) -> None:
        """Insert a figure (image + caption) from LaTeX IR text."""
        self._figure_index += 1

        # Extract paths and caption from LaTeX
        paths = []
        for match in _FIG_PATH_RE.finditer(latex_figure_text):
            paths.append(self._latex_to_plain(match.group(1)))
            
        cap_match = _CAPTION_RE.search(latex_figure_text)
        caption = self._normalize_figure_caption(self._latex_to_plain(cap_match.group(1)) if cap_match else "")

        # Insert all actual images
        for image_path in paths:
            resolved = self._resolve_image_path(image_path)
            if resolved and resolved.exists():
                try:
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_para.paragraph_format.space_before = Pt(6)
                    pic_para.paragraph_format.space_after = Pt(4)
                    pic_para.paragraph_format.line_spacing = _IEEE_CAPTION_LINE_SPACING
                    # Calculate column width dynamically if possible
                    col_width = self._get_current_column_width_inch(doc)
                    pic_para.add_run().add_picture(str(resolved), width=Inches(col_width))
                except Exception:
                    pass
            else:
                # Image file missing - show a visible placeholder
                fallback = doc.add_paragraph("[Image: " + image_path + "]")
                fallback.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in fallback.runs:
                    run.font.size = Pt(8)
                    run.italic = True
        if caption:
            cap_text = f"Fig. {self._figure_index}. {caption}"
        else:
            cap_text = f"Fig. {self._figure_index}."

        cap_para = doc.add_paragraph()
        fig_cap_style = self._pick_style_name(["figurecaption", "figure caption", "Caption"])
        if fig_cap_style:
            try:
                cap_para.style = fig_cap_style
                self._remove_paragraph_numbering(cap_para)
            except Exception:
                pass
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(2)
        cap_para.paragraph_format.space_after = Pt(8)
        cap_para.paragraph_format.line_spacing = _IEEE_CAPTION_LINE_SPACING
        run_cap = cap_para.add_run(cap_text)
        run_cap.font.name = "Times New Roman"
        run_cap.font.size = Pt(8)


    def _add_equation_node(self, doc: Document, raw_text: str) -> None:
        """Render an equation node as centered OMML or stylized text."""
        # Check for OMML marker first
        omml_match = re.search(r"«OMML:([A-Za-z0-9+/=]+)»", raw_text)
        
        # Determine the equation number (tag) if present
        tag_match = re.search(r"\\tag\{([^}]*)\}", raw_text)
        if tag_match:
            tag_text = re.sub(r"\s+", "", tag_match.group(1) or "")
            eq_num = f"({tag_text})" if tag_text else ""
        else:
            eq_num = ""

        clean = raw_text
        clean = re.sub(r"«OMML:([A-Za-z0-9+/=]+)»", "", clean)
        clean = re.sub(r"\\begin\{equation\*?\}", "", clean)
        clean = re.sub(r"\\end\{equation\*?\}", "", clean)
        clean = re.sub(r"\\tag\{([^}]*)\}", "", clean)
        clean = re.sub(r"\\\[", "", clean)
        clean = re.sub(r"\\\]", "", clean)
        clean = self._latex_math_to_readable(clean).strip()

        if eq_num:
            if not omml_match:
                p = doc.add_paragraph()
                self._apply_ieee_equation_tab_layout(p, doc, clean, eq_num)
                doc.add_paragraph("")
                return

            temp_table = doc.add_table(rows=1, cols=2)
            col_width = self._get_current_column_width_inch(doc)
            number_width = 0.34
            equation_width = max(2.8, col_width - number_width)
            temp_table.columns[0].width = Inches(equation_width)
            temp_table.columns[1].width = Inches(number_width)
            try:
                for row in temp_table.rows:
                    for cell in row.cells:
                        tc = cell._tc
                        tc_pr = tc.get_or_add_tcPr()
                        tc_borders = OxmlElement("w:tcBorders")
                        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                            edge = OxmlElement(f"w:{border_name}")
                            edge.set(qn("w:val"), "nil")
                            tc_borders.append(edge)
                        tc_pr.append(tc_borders)
            except Exception:
                pass

            p_eq = temp_table.cell(0, 0).paragraphs[0]
            self._set_cell_no_wrap(temp_table.cell(0, 0))
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if omml_match:
                try:
                    self._insert_omml_to_paragraph(p_eq, omml_match.group(1))
                except Exception:
                    if clean:
                        run = p_eq.add_run(clean)
                        run.font.name = "Times New Roman"
                        run.font.size = self._equation_font_size_for_text(clean)
                        run.italic = True
            elif clean:
                run = p_eq.add_run(clean)
                run.font.name = "Times New Roman"
                run.font.size = self._equation_font_size_for_text(clean)
                run.italic = True

            p_num = temp_table.cell(0, 1).paragraphs[0]
            self._set_cell_no_wrap(temp_table.cell(0, 1))
            p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_num.paragraph_format.space_before = Pt(0)
            run_num = p_num.add_run(eq_num)
            run_num.font.name = "Times New Roman"
            run_num.font.size = Pt(10)

            doc.add_paragraph("")
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        if omml_match:
            try:
                self._insert_omml_to_paragraph(p, omml_match.group(1))
            except Exception:
                if clean:
                    run = p.add_run(clean)
                    run.font.name = "Times New Roman"
                    run.font.size = self._equation_font_size_for_text(clean)
                    run.italic = True
        elif clean:
            run = p.add_run(clean)
            run.font.name = "Times New Roman"
            run.font.size = self._equation_font_size_for_text(clean)
            run.italic = True

    # ========================================================================
    # References
    # ========================================================================

    def _add_references(self, doc: Document, references: List[Dict[str, Any]]) -> None:
        if not references:
            return

        # IEEE Word template uses 'Heading 5' for the References heading (unnumbered)
        ref_heading_style = self._pick_style_name(["Heading 5", "heading 5"])
        if ref_heading_style and self._using_uploaded_template:
            p = doc.add_paragraph()
            try:
                p.style = ref_heading_style
            except Exception:
                pass
            run = p.add_run("References")
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        else:
            # Fallback: add unnumbered heading manually (don't use numbered _add_ieee_heading)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("References")
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

        # IEEE Word template uses 'references' style with auto-numbering
        ref_style = self._pick_style_name(["references", "References", "reference"])

        for idx, ref in enumerate(references, start=1):
            if isinstance(ref, dict):
                text = self._latex_to_plain(ref.get("text") or "")
            else:
                text = self._latex_to_plain(str(ref))
            if text:
                # Strip existing [N] prefix if present
                text = re.sub(r"^\[\d+\]\s*", "", text)
                # Clean remaining LaTeX artifacts
                text = self._clean_reference_text(text)

                p = doc.add_paragraph()

                if ref_style and self._using_uploaded_template:
                    # Use template 'references' style (has auto-numbering)
                    try:
                        p.style = ref_style
                    except Exception:
                        pass
                    # The style handles numbering — just add the text
                    run_text = p.add_run(text)
                    run_text.font.name = "Times New Roman"
                    run_text.font.size = Pt(8)
                else:
                    # Fallback: manual numbering
                    p.paragraph_format.first_line_indent = Inches(0.0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.0
                    run_num = p.add_run(f"[{idx}] ")
                    run_num.font.name = "Times New Roman"
                    run_num.font.size = Pt(8)
                    run_text = p.add_run(text)
                    run_text.font.name = "Times New Roman"
                    run_text.font.size = Pt(8)

    # ========================================================================
    # Rich text: parse LaTeX formatting into Word runs
    # ========================================================================

    def _add_rich_runs(self, paragraph, raw_text: str) -> None:
        """Parse LaTeX formatting commands and create Word runs with proper formatting.
        Also handles embedded OMML math markers.
        """
        # Handle OMML markers first (they might be inline)
        if "«OMML:" in raw_text:
            parts = re.split(r"(«OMML:[A-Za-z0-9+/=]+»)", raw_text)
            for part in parts:
                if part.startswith("«OMML:") and part.endswith("»"):
                    try:
                        b64_xml = part[6:-1]
                        self._insert_omml_to_paragraph(paragraph, b64_xml)
                    except Exception:
                        pass
                elif part.strip():
                    self._add_rich_runs(paragraph, part)
            return

        # If no LaTeX commands, just add plain text
        if "\\" not in raw_text and "{" not in raw_text:
            run = paragraph.add_run(raw_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            return

        # Clean out figure/label/caption blocks first (they're handled elsewhere)
        text = re.sub(r"\\begin\{figure\}.*?\\end\{figure\}", "", raw_text, flags=re.DOTALL)
        text = _LABEL_RE.sub("", text)

        # Parse into segments
        segments = self._parse_latex_to_segments(text)

        for seg_text, formatting in segments:
            if not seg_text:
                continue
            run = paragraph.add_run(seg_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

            if "bold" in formatting:
                run.bold = True
            if "italic" in formatting:
                run.italic = True
            if "underline" in formatting:
                run.underline = True
            if "superscript" in formatting:
                run.font.superscript = True
            if "subscript" in formatting:
                run.font.subscript = True
            if "monospace" in formatting:
                run.font.name = "Courier New"

    def _parse_latex_to_segments(self, text: str) -> List[Tuple[str, set]]:
        """Parse LaTeX text into (plain_text, formatting_set) segments."""
        segments: List[Tuple[str, set]] = []

        # Simplified recursive parser
        def _parse(s: str, active_formats: set) -> None:
            if not s:
                return

            # Find the first LaTeX command
            patterns = [
                (r"\\textbf\{", "bold"),
                (r"\\textit\{", "italic"),
                (r"\\emph\{", "italic"),
                (r"\\texttt\{", "monospace"),
                (r"\\underline\{", "underline"),
                (r"\\textsuperscript\{", "superscript"),
                (r"\\textsubscript\{", "subscript"),
            ]

            best_match = None
            best_pos = len(s)
            best_fmt = ""

            for pattern, fmt in patterns:
                m = re.search(pattern, s)
                if m and m.start() < best_pos:
                    best_match = m
                    best_pos = m.start()
                    best_fmt = fmt

            # Check for \href{url}{text}
            href_match = re.search(r"\\href\{([^}]*)\}\{", s)
            if href_match and href_match.start() < best_pos:
                # Text before href
                before = s[:href_match.start()]
                if before:
                    clean_before = self._strip_remaining_latex(before)
                    if clean_before:
                        segments.append((clean_before, set(active_formats)))

                # Find matching closing brace for link text
                inner_start = href_match.end()
                inner_text, after_pos = self._extract_braced_content(s, inner_start - 1)
                url = href_match.group(1)
                url_clean = url.replace("\\%", "%").replace("\\#", "#")
                link_text = self._strip_remaining_latex(inner_text) or url_clean
                segments.append((link_text, active_formats | {"underline"}))

                # Continue after
                if after_pos < len(s):
                    _parse(s[after_pos:], active_formats)
                return

            # Check for \cite{...} pattern
            cite_match = re.search(r"\\cite\{([^}]*)\}", s)
            if cite_match and cite_match.start() < best_pos:
                before = s[:cite_match.start()]
                if before:
                    clean_before = self._strip_remaining_latex(before)
                    if clean_before:
                        segments.append((clean_before, set(active_formats)))

                # Convert \cite{ref1,ref2} to [1, 2]
                refs = cite_match.group(1).split(",")
                nums = []
                for r in refs:
                    r = r.strip()
                    num_match = re.search(r"\d+", r)
                    if num_match:
                        nums.append(num_match.group())
                cite_text = "[" + ", ".join(nums) + "]" if nums else "[?]"
                segments.append((cite_text, set(active_formats)))

                after_pos = cite_match.end()
                if after_pos < len(s):
                    _parse(s[after_pos:], active_formats)
                return

            if best_match is None:
                # No more commands — output remaining as plain text
                clean = self._strip_remaining_latex(s)
                if clean:
                    segments.append((clean, set(active_formats)))
                return

            # Text before the command
            before = s[:best_pos]
            if before:
                clean_before = self._strip_remaining_latex(before)
                if clean_before:
                    segments.append((clean_before, set(active_formats)))

            # Extract content inside braces
            inner_start = best_match.end() - 1  # points to the {
            inner_text, after_pos = self._extract_braced_content(s, inner_start)

            # Recurse into the inner content with added formatting
            _parse(inner_text, active_formats | {best_fmt})

            # Continue after the closing brace
            if after_pos < len(s):
                _parse(s[after_pos:], active_formats)

        _parse(text, set())
        return segments

    def _extract_braced_content(self, s: str, open_pos: int) -> Tuple[str, int]:
        """Extract content between matching braces starting at open_pos.
        Returns (inner_content, position_after_closing_brace)."""
        if open_pos >= len(s) or s[open_pos] != "{":
            return ("", open_pos + 1)

        depth = 0
        i = open_pos
        while i < len(s):
            if s[i] == "{":
                depth += 1
            elif s[i] == "}":
                depth -= 1
                if depth == 0:
                    return (s[open_pos + 1 : i], i + 1)
            i += 1

        # Unmatched brace — return everything after open
        return (s[open_pos + 1 :], len(s))

    def _strip_remaining_latex(self, text: str) -> str:
        """Remove any remaining LaTeX commands from a text fragment."""
        s = text
        # Remove inline math $...$ but keep content
        s = re.sub(r"\$([^$]*)\$", lambda m: self._latex_math_to_readable(m.group(1)), s)
        # Remove remaining unknown commands but keep arguments
        for _ in range(3):
            s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", s)
        # Remove standalone commands
        s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", " ", s)
        # Remove stray braces
        s = s.replace("{", "").replace("}", "")
        s = s.replace("~", " ")
        s = re.sub(r"\s+", " ", s).strip()
        return s

    # ========================================================================
    # LaTeX conversion helpers
    # ========================================================================

    def _latex_to_plain(self, text: str) -> str:
        """Convert LaTeX-bearing text to clean plain text."""
        if not text:
            return ""

        s = str(text)
        s = s.replace("\\\\", "\n")

        # Handle inline math: convert to readable
        s = re.sub(r"\$([^$]+)\$", lambda m: self._latex_math_to_readable(m.group(1)), s)

        def __replace_cite(m):
            refs = m.group(1).split(",")
            nums = [r.strip().replace("ref", "") for r in refs]
            return f"[{', '.join(nums)}]"
        s = re.sub(r"\\cite\{([^{}]+)\}", __replace_cite, s)

        # Keep command arguments, remove command names (iterate for nesting)
        for _ in range(4):
            s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", s)

        s = re.sub(r"\\begin\{[^}]+\}", " ", s)
        s = re.sub(r"\\end\{[^}]+\}", " ", s)

        # Handle \cite{refN,refM} -> [N, M]
        def _cite_replace(m):
            refs = m.group(1).split(",")
            nums = []
            for r in refs:
                nm = re.search(r"\d+", r.strip())
                if nm:
                    nums.append(nm.group())
            return "[" + ", ".join(nums) + "]" if nums else "[?]"

        s = re.sub(r"\\cite\{([^}]*)\}", _cite_replace, s)

        s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", " ", s)

        s = s.replace("{", " ").replace("}", " ")
        s = s.replace("~", " ")
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _latex_math_to_readable(self, math_text: str) -> str:
        """Convert LaTeX math notation to readable Unicode-ish text."""
        s = math_text
        # Remove common LaTeX line-break/alignment markers to avoid unintended Word wrapping.
        s = s.replace("\\\\", " ")
        s = s.replace("&", " ")
        # Common replacements
        replacements = [
            (r"\\frac\{1\}\{2\}", "½"),
            (r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1)/(\2)"),
            (r"\\sqrt\{([^}]*)\}", r"√(\1)"),
            (r"\\sum", "Σ"),
            (r"\\prod", "∏"),
            (r"\\int", "∫"),
            (r"\\infty", "∞"),
            (r"\\alpha", "α"), (r"\\beta", "β"), (r"\\gamma", "γ"),
            (r"\\delta", "δ"), (r"\\epsilon", "ε"), (r"\\theta", "θ"),
            (r"\\lambda", "λ"), (r"\\mu", "μ"), (r"\\sigma", "σ"),
            (r"\\pi", "π"), (r"\\omega", "ω"), (r"\\phi", "φ"),
            (r"\\times", "×"), (r"\\cdot", "·"), (r"\\pm", "±"),
            (r"\\leq", "≤"), (r"\\geq", "≥"), (r"\\neq", "≠"),
            (r"\\approx", "≈"), (r"\\rightarrow", "→"), (r"\\leftarrow", "←"),
            (r"\\in", "∈"), (r"\\notin", "∉"),
            (r"\\subset", "⊂"), (r"\\supset", "⊃"),
            (r"\\cup", "∪"), (r"\\cap", "∩"),
            (r"\\forall", "∀"), (r"\\exists", "∃"),
            (r"\^(\{[^}]*\}|[0-9a-zA-Z])", lambda m: self._to_superscript(m.group(1))),
            (r"_(\{[^}]*\}|[0-9a-zA-Z])", lambda m: self._to_subscript(m.group(1))),
        ]
        for pattern, repl in replacements:
            if callable(repl):
                s = re.sub(pattern, repl, s)
            else:
                s = re.sub(pattern, repl, s)

        # Strip remaining LaTeX commands
        s = re.sub(r"\\[a-zA-Z]+\*?", " ", s)
        s = s.replace("{", "").replace("}", "")
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _equation_font_size_for_text(self, equation_text: str):
        """Adaptive equation font size to reduce wrapping for long formulas."""
        length = len((equation_text or "").strip())
        if length >= 110:
            return Pt(8)
        if length >= 80:
            return Pt(9)
        return Pt(10)

    def _to_superscript(self, text: str) -> str:
        """Convert text to Unicode superscript where possible."""
        t = text.strip("{}")
        sup_map = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
                    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
                    "+": "⁺", "-": "⁻", "=": "⁼", "n": "ⁿ", "i": "ⁱ"}
        return "".join(sup_map.get(c, c) for c in t)

    def _insert_omml_to_paragraph(self, paragraph, base64_xml: str) -> None:
        """Insert a native OMML math element into a docx paragraph."""
        try:
            xml_str = base64.b64decode(base64_xml).decode('utf-8')
            # The XML might have namespace declarations that need fixing for lxml insertion
            # We use etree to parse and then append the element
            omml_elem = etree.fromstring(xml_str.encode('utf-8'))
            tag_name = etree.QName(omml_elem.tag).localname
            if tag_name == "oMathPara":
                math_nodes = list(omml_elem.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"))
                if not math_nodes:
                    math_nodes = [omml_elem]
                for node in math_nodes:
                    paragraph._element.append(deepcopy(node))
            else:
                paragraph._element.append(omml_elem)
        except Exception as e:
            print(f"[Cảnh báo] Lỗi chèn OMML: {e}")

    def _is_equation_like_paragraph(self, raw_text: str) -> bool:
        text = str(raw_text or "")
        if "\\begin{equation" in text or "\\[" in text or "$$" in text:
            return True
        if "\\tag{" in text:
            return True
        if "«OMML:" in text:
            without_markers = re.sub(r"«OMML:[A-Za-z0-9+/=]+»", "", text)
            without_markers = re.sub(r"\\tag\{[^}]*\}", "", without_markers)
            without_markers = re.sub(r"[\\\s{}\[\]()]+", "", without_markers)
            if not without_markers:
                return True
        return False

    def _to_subscript(self, text: str) -> str:
        """Convert text to Unicode subscript where possible."""
        t = text.strip("{}")
        sub_map = {"0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
                    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
                    "+": "₊", "-": "₋", "=": "₌"}
        return "".join(sub_map.get(c, c) for c in t)

    # ========================================================================
    # Image resolution
    # ========================================================================

    def _resolve_image_path(self, image_path: str) -> Optional[Path]:
        if not image_path:
            return None

        candidate = Path(image_path)
        if candidate.is_absolute() and candidate.exists():
            return candidate

        root = self._image_root_dir
        joined = (root / image_path).resolve()
        if joined.exists():
            return joined

        # Common case: LaTeX path starts with images/...
        if image_path.startswith("images/"):
            joined_alt = (root / image_path[len("images/"):]).resolve()
            if joined_alt.exists():
                return joined_alt

        # Try just the basename
        basename = Path(image_path).name
        for search_dir in [root, root / "images"]:
            if search_dir.is_dir():
                candidate_path = search_dir / basename
                if candidate_path.exists():
                    return candidate_path

        return None

    # ========================================================================
    # Document configuration (from-scratch mode)
    # ========================================================================

    def _configure_ieee_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.78)
        section.bottom_margin = Cm(1.78)
        section.left_margin = Cm(1.65)
        section.right_margin = Cm(1.65)
        self._set_section_columns(section, 1)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.line_spacing = 1.0

    # ========================================================================
    # Utility: resolve table style
    # ========================================================================

    def _resolve_table_style_name(self, doc: Document) -> str | None:
        preferred = ["Table Grid", "Table Normal", "Normal Table"]
        available = {s.name for s in doc.styles if getattr(s, "name", None)}
        for name in preferred:
            if name in available:
                return name
        return None

    def _start_two_column_body(self, doc: Document) -> None:
        """Switch to IEEE two-column layout for main content and references."""
        if not doc.sections:
            return

        for sec in doc.sections:
            self._set_section_columns(sec, 1)

        body_section = doc.add_section(WD_SECTION.CONTINUOUS)
        body_section.left_margin = doc.sections[0].left_margin
        body_section.right_margin = doc.sections[0].right_margin
        body_section.top_margin = doc.sections[0].top_margin
        body_section.bottom_margin = doc.sections[0].bottom_margin
        self._set_section_columns(body_section, 2)

    def _set_section_columns(self, section, num: int, space_twips: int = 360) -> None:
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), str(max(1, int(num))))
        cols.set(qn("w:space"), str(max(0, int(space_twips))))

    def _force_table_borders(self, table: Table) -> None:
        """Apply IEEE-standard table borders.

        Readability-first profile for conversion output:
        - Top/bottom/left/right: single, thin
        - InsideH/insideV: single, thin
        """
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)

        # Borders to SHOW (thin lines)
        for name in ["top", "bottom", "left", "right", "insideH", "insideV"]:
            edge = borders.find(qn(f"w:{name}"))
            if edge is None:
                edge = OxmlElement(f"w:{name}")
                borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "4")  # thin (0.5pt) — IEEE uses very thin lines
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "000000")

    def _add_table_caption(self, doc: Document, label: str, caption: str) -> None:
        """IEEE conference table caption: two lines (label + uppercase title)."""
        label_para = doc.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_para.paragraph_format.space_before = Pt(6)
        label_para.paragraph_format.space_after = Pt(0)
        label_para.paragraph_format.line_spacing = _IEEE_CAPTION_LINE_SPACING
        run_label = label_para.add_run(label)
        run_label.font.name = "Times New Roman"
        run_label.font.size = Pt(8)

        if caption:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(2)
            title_para.paragraph_format.line_spacing = _IEEE_CAPTION_LINE_SPACING
            run_title = title_para.add_run(caption.upper())
            run_title.font.name = "Times New Roman"
            run_title.font.size = Pt(8)

    def _determine_cell_alignment(self, text: str, is_header: bool = False):
        value = (text or "").strip()
        if is_header:
            return WD_ALIGN_PARAGRAPH.CENTER
        if not value:
            return WD_ALIGN_PARAGRAPH.LEFT

        compact = value.replace(",", "").replace("%", "").replace("$", "").strip()
        if re.fullmatch(r"[+-]?\d+(?:\.\d+)?", compact):
            return WD_ALIGN_PARAGRAPH.RIGHT
        if re.fullmatch(r"\d+[\-\/]\d+", compact):
            return WD_ALIGN_PARAGRAPH.CENTER
        return WD_ALIGN_PARAGRAPH.LEFT

    def _set_table_width(self, table: Table, cols: int, width_inch: float) -> None:
        if cols <= 0:
            return

        col_width = Inches(max(1.0, width_inch) / cols)
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < cols:
                    cell.width = col_width

        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(max(1.0, width_inch) * 1440)))

    def _current_table_target_width_inch(self, doc: Document, force_full_width: bool) -> float:
        if not doc.sections:
            return _IEEE_COLUMN_WIDTH_INCH

        sec = doc.sections[-1]
        try:
            content_width = (sec.page_width - sec.left_margin - sec.right_margin).inches
        except Exception:
            content_width = _IEEE_COLUMN_WIDTH_INCH * 2 + _IEEE_TWO_COL_GUTTER_INCH

        if force_full_width:
            return max(_IEEE_COLUMN_WIDTH_INCH, content_width)

        col_width = (content_width - _IEEE_TWO_COL_GUTTER_INCH) / 2.0
        return max(2.8, min(_IEEE_COLUMN_WIDTH_INCH, col_width))

    def _select_table_layout_mode(self, doc: Document, node: Dict[str, Any]) -> str:
        required_width = self._estimate_table_required_width_inch(node)
        single_col_width = self._current_table_target_width_inch(doc, force_full_width=False)
        full_width = self._current_table_target_width_inch(doc, force_full_width=True)

        cols = int(node.get("cols", 1) or 1)
        # Keep 3-column tables in single-column flow by default.
        # Only 4+ columns are considered for spanning full width.
        if cols <= 3:
            return "column"

        # 4-column tables: try hard to keep in one column first, avoid oversized look.
        if cols == 4:
            if required_width <= single_col_width * 1.35:
                return "column"
            if required_width <= full_width * 1.08:
                return "full"
            return "landscape"

        # 5+ columns: prefer full-width for genuinely wider tables.
        if cols >= 5:
            if required_width <= full_width * 1.20:
                return "full"
            return "landscape"

        # Keep very small tables in-column.
        if required_width <= single_col_width * 1.18:
            return "column"

        # Use landscape only for truly extreme wide tables.
        if cols >= 8 and required_width > full_width * 1.02:
            return "landscape"
        if required_width <= full_width * 1.12:
            return "full"
        if cols <= 5 and required_width <= full_width * 1.35:
            return "full"
        return "landscape"

    def _apply_ieee_equation_tab_layout(self, paragraph, doc: Document, eq_text: str, eq_num: str) -> None:
        """Apply IEEE-style equation layout: centered equation + right equation number."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

        col_width = self._get_current_column_width_inch(doc)
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(max(1.4, col_width * 0.50)), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Inches(max(2.6, col_width)), WD_TAB_ALIGNMENT.RIGHT)

        run = paragraph.add_run("\t" + (eq_text or ""))
        run.font.name = "Times New Roman"
        run.font.size = self._equation_font_size_for_text(eq_text)
        run.italic = True

        if eq_num:
            run_num = paragraph.add_run("\t" + eq_num)
            run_num.font.name = "Times New Roman"
            run_num.font.size = Pt(10)

    def _set_cell_no_wrap(self, cell) -> None:
        """Prevent Word from wrapping equation text inside a narrow table cell."""
        try:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            no_wrap = tc_pr.find(qn("w:noWrap"))
            if no_wrap is None:
                no_wrap = OxmlElement("w:noWrap")
                tc_pr.append(no_wrap)
            no_wrap.set(qn("w:val"), "1")
        except Exception:
            pass

    def _estimate_table_required_width_inch(self, node: Dict[str, Any]) -> float:
        cols = int(node.get("cols", 1) or 1)
        rows_data = node.get("data", []) or []
        if cols <= 0:
            return _IEEE_COLUMN_WIDTH_INCH

        max_chars_per_col = [0] * cols
        for r_idx, row in enumerate(rows_data):
            c_idx = 0
            for cell in row:
                if c_idx >= cols:
                    break
                colspan = max(1, int(cell.get("colspan", 1) or 1))
                text = self._latex_to_plain(cell.get("text") or "")
                lines = text.splitlines() or [text]
                effective_chars = max((len(line.strip()) for line in lines), default=0)
                if r_idx == 0:
                    effective_chars = int(effective_chars * 1.15)

                share = max(1, colspan)
                per_col = int(max(1, effective_chars) / share)
                for offset in range(share):
                    col_idx = c_idx + offset
                    if col_idx < cols:
                        max_chars_per_col[col_idx] = max(max_chars_per_col[col_idx], per_col)
                c_idx += colspan

        estimated = 0.0
        for chars in max_chars_per_col:
            # 8pt Times New Roman rough fit model with side padding
            estimated += max(0.72, min(2.8, chars * 0.045 + 0.26))
        return max(1.5, estimated)

    def _should_span_table_two_columns(self, node: Dict[str, Any]) -> bool:
        cols = int(node.get("cols", 1) or 1)
        rows_data = node.get("data", []) or []
        if cols >= 5:
            return True

        if not rows_data:
            return False

        longest_cell = 0
        long_cells = 0
        for row in rows_data:
            for cell in row:
                text = self._latex_to_plain(cell.get("text") or "")
                length = len(text.strip())
                if length > longest_cell:
                    longest_cell = length
                if length >= 28:
                    long_cells += 1

        return cols >= 4 or longest_cell >= 40 or long_cells >= 4

    def _start_single_column_block(self, doc: Document) -> None:
        sec = doc.add_section(WD_SECTION.CONTINUOUS)
        self._copy_section_margins(doc.sections[0], sec)
        self._copy_section_page_size(doc.sections[0], sec)
        sec.orientation = WD_ORIENT.PORTRAIT
        self._set_section_columns(sec, 1)

    def _resume_two_column_block(self, doc: Document) -> None:
        sec = doc.add_section(WD_SECTION.CONTINUOUS)
        self._copy_section_margins(doc.sections[0], sec)
        self._copy_section_page_size(doc.sections[0], sec)
        sec.orientation = WD_ORIENT.PORTRAIT
        self._set_section_columns(sec, 2)

    def _start_landscape_single_column_block(self, doc: Document) -> None:
        sec = doc.add_section(WD_SECTION.CONTINUOUS)
        self._copy_section_margins(doc.sections[0], sec)
        base = doc.sections[0]
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = base.page_height
        sec.page_height = base.page_width
        self._set_section_columns(sec, 1)

    def _resume_portrait_two_column_block(self, doc: Document) -> None:
        sec = doc.add_section(WD_SECTION.CONTINUOUS)
        self._copy_section_margins(doc.sections[0], sec)
        self._copy_section_page_size(doc.sections[0], sec)
        sec.orientation = WD_ORIENT.PORTRAIT
        self._set_section_columns(sec, 2)

    def _copy_section_margins(self, source, target) -> None:
        target.left_margin = source.left_margin
        target.right_margin = source.right_margin
        target.top_margin = source.top_margin
        target.bottom_margin = source.bottom_margin

    def _copy_section_page_size(self, source, target) -> None:
        target.page_width = source.page_width
        target.page_height = source.page_height

    def _normalize_table_caption(self, caption: str) -> str:
        text = (caption or "").strip()
        original = text
        text = re.sub(r"^TABLE\s+[IVXLCDM]+(?:\.|:)\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"^Table\s+\d+(?:\.|:)\s*", "", text, flags=re.IGNORECASE)
        normalized = text.strip(" .:-")
        if original and normalized != original:
            self._render_metrics["table_caption_normalized"] += 1
        return normalized

    def _normalize_figure_caption(self, caption: str) -> str:
        text = (caption or "").strip()
        original = text
        text = re.sub(r"^Fig(?:ure)?\.?\s*\d+\.?\s*", "", text, flags=re.IGNORECASE)
        normalized = text.strip(" .:-")
        if original and normalized != original:
            self._render_metrics["figure_caption_normalized"] += 1
        return normalized

    def _apply_body_paragraph_style(self, paragraph, indent_first_line: bool = False) -> None:
        """Apply IEEE body text formatting.

        When using an uploaded template with 'Body Text' style, the style
        already defines line spacing, font, etc.  We only override indent
        when the style is NOT applied (from-scratch mode).
        """
        if not self._using_uploaded_template:
            paragraph.paragraph_format.line_spacing = _IEEE_BODY_LINE_SPACING
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(_IEEE_BODY_SPACE_AFTER_PT)
            if indent_first_line:
                paragraph.paragraph_format.first_line_indent = Cm(0.5)

    def _pick_style_name(self, candidates: List[str]) -> str | None:
        if not candidates:
            return None
        lowered = {name.lower(): name for name in self._available_styles}
        for cand in candidates:
            if cand in self._available_styles:
                return cand
            key = cand.lower()
            if key in lowered:
                return lowered[key]
        return None

    # ========================================================================
    # New helpers for IEEE v2 upgrade
    # ========================================================================

    def _clean_cell_text(self, raw: str) -> str:
        """Clean a table cell text value: convert LaTeX to plain and remove escapes."""
        text = self._latex_to_plain(raw)
        # Remove LaTeX escape artifacts that _latex_to_plain may miss
        text = text.replace("\\%", "%")
        text = text.replace("\\&", "&")
        text = text.replace("\\#", "#")
        text = text.replace("\\$", "$")
        text = text.replace("\\\\%", "%")
        text = text.replace("\\\\&", "&")
        return text

    def _clean_reference_text(self, text: str) -> str:
        """Additional cleaning for reference entry text."""
        text = text.replace("\\%", "%")
        text = text.replace("\\&", "&")
        text = text.replace("\\#", "#")
        text = text.replace("\\$", "$")
        text = text.replace("\\\\%", "%")
        text = text.replace("\\\\&", "&")
        # Remove stray backslashes before common chars
        text = re.sub(r"\\(?=[%&#$])", "", text)
        return text.strip()

    def _get_current_column_width_inch(self, doc: Document) -> float:
        """Calculate the actual column width from section properties."""
        if not doc.sections:
            return _IEEE_COLUMN_WIDTH_INCH
        sec = doc.sections[-1]
        try:
            content_width = (sec.page_width - sec.left_margin - sec.right_margin).inches
        except Exception:
            content_width = _IEEE_COLUMN_WIDTH_INCH * 2 + _IEEE_TWO_COL_GUTTER_INCH

        # Check if we're in a 2-column section
        cols_el = sec._sectPr.find(qn("w:cols"))
        num_cols = 1
        if cols_el is not None:
            try:
                num_cols = int(cols_el.get(qn("w:num"), "1"))
            except (ValueError, TypeError):
                pass

        if num_cols >= 2:
            return (content_width - _IEEE_TWO_COL_GUTTER_INCH) / 2.0
        return content_width

    def _add_list_node(self, doc: Document, node: Dict[str, Any]) -> None:
        """Render a list node (bullet or numbered) using IEEE styles."""
        items = node.get("items", [])
        bullet_style = self._pick_style_name(["bullet list", "List Bullet", "List Number"])

        for item in items:
            text = self._latex_to_plain(str(item.get("text", "") if isinstance(item, dict) else item))
            if not text.strip():
                continue
            p = doc.add_paragraph()
            if bullet_style:
                try:
                    p.style = bullet_style
                except Exception:
                    pass
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    def _add_latex_list_node(self, doc: Document, raw_text: str) -> None:
        """Parse LaTeX \\begin{itemize}/\\begin{enumerate} and render as Word list."""
        bullet_style = self._pick_style_name(["bullet list", "List Bullet"])

        # Extract items from \\item ...
        # Remove environment wrappers
        clean = re.sub(r"\\begin\{(itemize|enumerate)\}", "", raw_text)
        clean = re.sub(r"\\end\{(itemize|enumerate)\}", "", clean)

        items = re.split(r"\\item\b\s*", clean)
        for item_text in items:
            item_text = self._latex_to_plain(item_text.strip())
            if not item_text:
                continue
            p = doc.add_paragraph()
            if bullet_style:
                try:
                    p.style = bullet_style
                except Exception:
                    pass
            else:
                # Fallback: add bullet character manually
                item_text = f"• {item_text}"
            run = p.add_run(item_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    def _remove_paragraph_numbering(self, paragraph) -> None:
        """Force paragraph to not use auto-numbering (overriding style default)."""
        pPr = paragraph._p.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numId = numPr.find(qn('w:numId'))
        if numId is None:
            numId = OxmlElement('w:numId')
            numPr.append(numId)
        numId.set(qn('w:val'), '0')
