"""Springer Word Renderer built on top of IEEE renderer infrastructure.

This renderer keeps the robust IR handling/equation/table logic from IEEEWordRenderer,
but switches layout and styles toward Springer-like Word output.
"""

import re
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .word_ieee_renderer import IEEEWordRenderer


class SpringerWordRenderer(IEEEWordRenderer):
    """Render IR output into a Springer-like Word document."""

    def render(
        self,
        ir_data: Dict[str, Any],
        output_path: str,
        image_root_dir: str | None = None,
        springer_template_path: str | None = None,
    ) -> str:
        rendered_path = super().render(
            ir_data=ir_data,
            output_path=output_path,
            image_root_dir=image_root_dir,
            ieee_template_path=springer_template_path,
        )

        # Template headers often keep placeholder short title unless replaced explicitly.
        try:
            doc = Document(rendered_path)
            self._sync_template_header_title(doc, ir_data.get("metadata") or {})
            doc.save(rendered_path)
        except Exception as e:
            print(f"[WARN] Failed to sync Springer template header: {e}")

        return rendered_path

    def _rebuild_on_uploaded_template(
        self,
        doc: Document,
        metadata: Dict[str, Any],
        body_nodes: List[Dict[str, Any]],
        references: List[Dict[str, Any]],
    ) -> None:
        '''Override native clear-body approach to use fill-in-place for Springer templates.
        This preserves Springer macros, bookmarks, fields, and headers/footers exactly.
        '''
        self._fill_springer_template(doc, metadata, body_nodes, references)

    def _fill_springer_template(
        self,
        doc: Document,
        metadata: Dict[str, Any],
        body_nodes: List[Dict[str, Any]],
        references: List[Dict[str, Any]],
    ) -> None:
        title = (metadata.get("title") or "").strip()
        abstract = self._latex_to_plain(metadata.get("abstract") or "")
        abstract = re.sub(r"^\s*(?:abstract|t[oó]m\s+t[aắ]t)\s*[:.\u2013\u2014\-]+\s*", "", abstract, flags=re.IGNORECASE)
        abstract = re.sub(r"^\s*[\-\u2013\u2014]{2,}\s*", "", abstract).strip()
        
        keywords_raw = metadata.get("keywords") or []
        keywords = self._sanitize_keywords([self._latex_to_plain(str(k)) for k in keywords_raw])
        authors = metadata.get("authors") or []

        all_paras = list(self._iter_all_paragraphs(doc))
        title_idx, abs_idx, kw_idx, first_sec_idx, ref_idx = -1, -1, -1, -1, -1

        for i, p in enumerate(all_paras):
            text_upper = p.text.strip().upper()
            if p.style.name == "papertitle" and title_idx == -1:
                title_idx = i
            if p.style.name == "abstract" and abs_idx == -1:
                abs_idx = i
            if p.style.name == "keywords" and kw_idx == -1:
                kw_idx = i
            if p.style.name == "heading1" and "FIRST SECTION" in text_upper and first_sec_idx == -1:
                first_sec_idx = i
            if p.style.name == "heading1" and "REFERENCES" in text_upper and ref_idx == -1:
                ref_idx = i

        # Replace Title
        if title_idx != -1 and title:
            self._replace_paragraph_text_keep_formatting(all_paras[title_idx], self._latex_to_plain(title))

        # Handle Authors: Clear everything between Title and Abstract, then insert
        if title_idx != -1 and abs_idx != -1:
            for i in range(title_idx + 1, abs_idx):
                p_el = all_paras[i]._element; p_parent = p_el.getparent();
                if p_parent is not None:
                    p_parent.remove(p_el)
            if authors:
                anchor_p = all_paras[abs_idx]
                self._insert_springer_authors_before(anchor_p, authors)

        # Handle Abstract
        if abs_idx != -1 and abstract:
            p = all_paras[abs_idx]
            p.clear()
            r1 = p.add_run("Abstract. ")
            r1.bold = True
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(9)
            r2 = p.add_run(abstract)
            r2.bold = False
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(9)

        # Handle Keywords
        if kw_idx != -1 and keywords:
            kw_str = ", ".join(keywords) + "."
            p = all_paras[kw_idx]
            p.clear()
            r1 = p.add_run("Keywords: ")
            r1.bold = True
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(9)
            r2 = p.add_run(kw_str)
            r2.bold = False
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(9)

        # Handle Body
        if first_sec_idx != -1:
            end_idx = ref_idx if ref_idx != -1 else len(all_paras)

            # Delete any existing tables in the template body to prevent leftover lines/shapes
            for t in doc.tables:
                t_el = t._element
                if t_el.getparent() is not None:
                    t_el.getparent().remove(t_el)

            # Remove all paragraphs EXCEPT first_sec_idx to use it as anchor
            for i in range(first_sec_idx + 1, end_idx):
                if all_paras[i]._element.getparent() is not None:
                    p_el = all_paras[i]._element; p_parent = p_el.getparent();
                    if p_parent is not None:
                        p_parent.remove(p_el)
            
            anchor_p = all_paras[first_sec_idx]
            self._insert_springer_body_before(doc, anchor_p, body_nodes)
            # Now remove the old anchor paragraph entirely from the XML
            anchor_el = anchor_p._element; anchor_parent = anchor_el.getparent();
            if anchor_parent is not None:
                anchor_parent.remove(anchor_el)

        # Handle References
        if ref_idx != -1:
            # We want to use ref_anchor. If ref_idx+1 exists, use it as anchor. However ref_idx itself
            # might just be the "References" title which we keep!
            # Wait, no. "References" needs to be kept by us. But we clear from ref_idx+1.
            for i in range(ref_idx + 1, len(all_paras)):
                if all_paras[i]._element.getparent() is not None:
                    p_el = all_paras[i]._element; p_parent = p_el.getparent();
                    if p_parent is not None:
                        p_parent.remove(p_el)
            
            ref_anchor = all_paras[ref_idx] # We shouldn't use ref_idx+1 as anchor since we just deleted it!
            # So insert references after ref_anchor? No, insert BEFORE ref_anchor doesn't make sense since ref_anchor is the title.
            # But doc.add_paragraph() is totally fine here! Because References is at the END.
            ref_style = self._pick_style_name(["referenceitem", "referencelist", "ReferenceLine", "Normal"])

            for idx, ref in enumerate(references, start=1):
                text = self._latex_to_plain(ref.get("text") or "" if isinstance(ref, dict) else str(ref))
                if not text: continue
                text = self._clean_reference_text(text)
                while True:
                    updated = re.sub(r"^\s*\[?\d+\]?\s*[\.)]?\s*", "", text)
                    if updated == text: break
                    text = updated
                text = re.sub(r"(?<!\s)(https?://)", r" \1", text)

                p = doc.add_paragraph()
                if ref_style:
                    try: p.style = ref_style
                    except: pass
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.left_indent = Inches(0.2)
                
                run_text = p.add_run(f"{text}")
                run_text.font.name = "Times New Roman"
                run_text.font.size = Pt(9)

    def _insert_springer_authors_before(self, anchor_p, authors: List[Dict[str, Any]]) -> None:
        author_style = self._pick_style_name(["author", "Author"])
        address_style = self._pick_style_name(["address", "Address", "institute"])
        email_style = self._pick_style_name(["email", "Email", "address", "Normal"])

        valid_authors = []
        extra_affs = []
        
        for a in authors:
            raw_name = str(a.get("name") or "").strip()
            if not raw_name: continue
            cleaned = self._clean_author_name(raw_name)
            if cleaned and self._looks_like_person_name(cleaned):
                valid_authors.append(a)
            elif cleaned:
                extra_affs.append(cleaned)

        unique_affs: List[str] = []
        author_aff_map: List[List[int]] = []
        all_emails = []

        for a in authors:
            email = self._latex_to_plain(str(a.get("email") or "")).strip()
            if email and self._looks_like_email(email):
                all_emails.append(email)

        for a in valid_authors:
            idx_list = []
            affs = []
            for x in (a.get("affiliations") or []):
                val = self._latex_to_plain(str(x)).strip()
                if not val or self._looks_like_pure_author_name(val):
                    continue
                if self._looks_like_email(val):
                    all_emails.append(val)
                    continue
                affs.append(val)
            aff_line = "; ".join(affs)
            if aff_line:
                if aff_line not in unique_affs:
                    unique_affs.append(aff_line)
                idx_list.append(unique_affs.index(aff_line) + 1)
            author_aff_map.append(idx_list)

        for aff in extra_affs:
            if aff not in unique_affs and not self._looks_like_email(aff):
                unique_affs.append(aff)

        if valid_authors:
            p_name = anchor_p.insert_paragraph_before()
            if author_style:
                try: p_name.style = author_style
                except: pass
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for i, a in enumerate(valid_authors):
                name = self._clean_author_name(str(a.get("name") or ""))
                if i > 0:
                    if i == len(valid_authors) - 1:
                        p_name.add_run(" and " if len(valid_authors) == 2 else ", and ")
                    else:
                        p_name.add_run(", ")
                run = p_name.add_run(name)

                indices = author_aff_map[i]
                if len(unique_affs) > 1 and indices:
                    idx_str = ",".join(str(idx) for idx in indices)
                    run_sup = p_name.add_run(idx_str)
                    run_sup.font.superscript = True

        if unique_affs:
            for i, aff in enumerate(unique_affs, start=1):
                p_aff = anchor_p.insert_paragraph_before()
                if address_style:
                    try: p_aff.style = address_style
                    except: pass
                p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
                prefix = f"{i} " if len(unique_affs) > 1 else ""
                p_aff.add_run(prefix + aff)

        all_emails = list(dict.fromkeys(all_emails))
        if all_emails:
            p_email = anchor_p.insert_paragraph_before()
            if email_style:
                try: p_email.style = email_style
                except:
                    try: p_email.style = address_style
                    except: pass
            p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_email.add_run(", ".join(all_emails))

    def _insert_springer_body_before(self, doc: Document, anchor_p, body_nodes: List[Dict[str, Any]]) -> None:
        self._section_index = 0
        self._subsection_counters = {}
        for idx, node in enumerate(body_nodes):
            node_type = node.get("type", "")

            if node_type == "section":
                level = int(node.get("level", 1) or 1)
                text = self._latex_to_plain(node.get("text") or "")
                if text:
                    self._insert_springer_heading_before(anchor_p, text, level)
                continue

            if node_type == "table":
                self._insert_springer_table_before(doc, anchor_p, node)
                continue

            if node_type == "list":
                list_text = self._latex_to_plain(node.get("text") or "")
                self._insert_rich_paragraph_before(anchor_p, list_text)
                continue

            if node_type == "paragraph":
                raw_text = str(node.get("text") or "")
                plain = self._latex_to_plain(raw_text).strip()
                if not plain: continue
                if self._is_duplicate_table_title_paragraph(plain, body_nodes, idx): continue
                if re.match(r"^\s*(?:TABLE|BANG|BẢNG)\s+[IVXLCDM\d]+\s*[:.\-]?\s*$", plain, re.IGNORECASE): continue

                if self._is_equation_like_paragraph(raw_text):
                    self._insert_springer_equation_before(doc, anchor_p, raw_text)
                    continue

                if "\\begin{figure" in raw_text or "\\includegraphics" in raw_text:
                    self._insert_springer_figure_before(doc, anchor_p, raw_text)
                    continue

                if re.match(r"^Fig\.?\s*\d+\.?", plain, re.IGNORECASE):
                    cap_clean = self._normalize_springer_caption(plain, "figure")
                    fig_idx_match = re.match(r"^Fig\.?\s*(\d+)\.?", plain, re.IGNORECASE)
                    self._figure_index = int(fig_idx_match.group(1)) if fig_idx_match else self._figure_index + 1
                    
                    p_cap = anchor_p.insert_paragraph_before()
                    fig_style = self._pick_style_name(["figurecaption", "Figure Caption", "Caption"])
                    if fig_style:
                        try: p_cap.style = fig_style
                        except: pass
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.add_run(f"Fig. {self._figure_index}. {cap_clean}" if cap_clean else f"Fig. {self._figure_index}.")
                    continue

                p = anchor_p.insert_paragraph_before()
                is_first_para = (idx == 0) or (body_nodes[idx - 1].get("type") in ("section", "table", "figure"))
                style_candidates = ["p1a", "Normal"] if is_first_para else ["Normal", "p1a"]
                
                body_style = self._pick_style_name(style_candidates)
                if body_style:
                    try: p.style = body_style
                    except: pass
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(plain)

    def _insert_springer_heading_before(self, anchor_p, text: str, level: int) -> None:
        clean = self._latex_to_plain(text)
        if not clean:
            return
        clean = self._strip_existing_heading_number(clean)
        if clean.isupper() and len(clean) > 3:
            clean = clean.title()

        p = anchor_p.insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_candidates = []
        if level == 1: style_candidates = ["heading1", "Heading 1"]
        elif level == 2: style_candidates = ["heading2", "Heading 2"]
        else: style_candidates = ["heading3", "Heading 3", f"Heading {min(level, 4)}"]
        
        heading_style = self._pick_style_name(style_candidates)
        if heading_style:
            try: p.style = heading_style
            except: pass

        if level == 1:
            self._section_index += 1
            self._subsection_counters[self._section_index] = 0
            run = p.add_run(clean)
            if not heading_style:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            return

        if self._section_index not in self._subsection_counters:
            self._subsection_counters[self._section_index] = 0
        self._subsection_counters[self._section_index] += 1

        if level == 2:
            run = p.add_run(clean)
            if not heading_style:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
        else:
            run = p.add_run(clean)
            if not heading_style:
                run.bold = True
                run.italic = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    def _insert_springer_table_before(self, doc: Document, anchor_p, node: Dict[str, Any]) -> None:
        self._table_index = getattr(self, "_table_index", 0) + 1
        caption = self._normalize_springer_caption(str(node.get("caption") or ""), "table")
        text = f"Table {self._table_index}. {caption}" if caption else f"Table {self._table_index}."
        
        cap_p = anchor_p.insert_paragraph_before()
        cap_style = self._pick_style_name(["tablecaption", "Table Caption", "Caption"])
        if cap_style:
            try: cap_p.style = cap_style
            except: pass
        cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_cap = cap_p.add_run(text)
        run_cap.font.name = "Times New Roman"
        run_cap.font.size = Pt(9)

        parent = anchor_p._p.getparent()
        table_data = node.get("data", [])
        if not table_data: return

        try: temp_table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data))
        except TypeError: temp_table = doc.add_table(rows=len(table_data), cols=max(len(r) for r in table_data), width=Inches(6.0))
        
        # Springer tables use "Normal Table" (invisible borders) with explicit horizontal rules
        # Let's apply top border to first row, bottom border to first row, and bottom border to last row
        for r_idx, row_data in enumerate(table_data):
            for c_idx, cell_data in enumerate(row_data):
                if c_idx < len(temp_table.rows[r_idx].cells):
                     cell = temp_table.cell(r_idx, c_idx)
                     # Clear default paragraph space
                     for p in cell.paragraphs:
                         p.paragraph_format.space_before = Pt(1)
                         p.paragraph_format.space_after = Pt(1)
                     cell.text = self._latex_to_plain(cell_data.get("text", ""))
                     # Format font
                     for p in cell.paragraphs:
                         for r in p.runs:
                             r.font.name = "Times New Roman"
                             r.font.size = Pt(9)
                             if r_idx == 0:
                                 r.bold = True
                     
                     # Apply horizontal borders
                     tcPr = cell._tc.get_or_add_tcPr()
                     tcBorders = OxmlElement('w:tcBorders')
                     
                     def add_border(edge):
                         border = OxmlElement(f'w:{edge}')
                         border.set(qn('w:val'), 'single')
                         border.set(qn('w:sz'), '4')
                         border.set(qn('w:space'), '0')
                         border.set(qn('w:color'), '000000')
                         tcBorders.append(border)

                     if r_idx == 0:
                         add_border('top')
                         add_border('bottom')
                     elif r_idx == len(table_data) - 1:
                         add_border('bottom')
                         
                     tcPr.append(tcBorders)
                            
        tbl_xml = temp_table._element
        parent.remove(tbl_xml)
        anchor_p._p.addprevious(tbl_xml)

    def _insert_springer_figure_before(self, doc: Document, anchor_p, latex_figure_text: str) -> None:
        self._figure_index += 1
        path_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^\}]+)\}", latex_figure_text)
        cap_match = re.search(r"\\caption\{([^\}]*)\}", latex_figure_text)
        image_path = self._latex_to_plain(path_match.group(1)) if path_match else ""
        caption = self._normalize_springer_caption(
            self._latex_to_plain(cap_match.group(1)) if cap_match else "",
            "figure",
        )
        resolved = self._resolve_image_path(image_path)
        if resolved and resolved.exists():
            try:
                from PIL import Image as PILImage
                img_width_inches = 4.8 # Default max width
                try:
                    with PILImage.open(str(resolved)) as img:
                        ppi = img.info.get("dpi", (72.0, 72.0))
                        if not hasattr(ppi, "__iter__"): ppi = (72.0, 72.0)
                        img_width_inches = img.width / ppi[0]
                except Exception:
                    pass

                pic_para = anchor_p.insert_paragraph_before()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Springer LNCS text width is ~12.2 cm (4.8 inches)
                if img_width_inches > 4.8:
                    pic_para.add_run().add_picture(str(resolved), width=Inches(4.8))
                else:
                    pic_para.add_run().add_picture(str(resolved))
            except: pass

        cap_text = f"Fig. {self._figure_index}. {caption}" if caption else f"Fig. {self._figure_index}."
        p = anchor_p.insert_paragraph_before()
        fig_style = self._pick_style_name(["figurecaption", "Figure Caption", "Caption"])
        if fig_style:
            try: p.style = fig_style
            except: pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(cap_text)

    def _insert_springer_equation_before(self, doc: Document, anchor_p, raw_text: str) -> None:
        omml_match = re.search(r"«OMML:([A-Za-z0-9+/=]+)»", raw_text)
        tag_match = re.search(r"\\tag\{([^\}]*)\}", raw_text)
        if tag_match:
            tag_text = re.sub(r"\s+", "", tag_match.group(1) or "")
            eq_num = f"({tag_text})" if tag_text else ""
        else:
            trailing_num = re.search(r"\((\d{1,3}[A-Za-z]?)\)\s*$", str(raw_text or ""))
            if trailing_num and "=" in str(raw_text or ""):
                eq_num = f"({trailing_num.group(1)})"
                raw_text = str(raw_text or "")[: trailing_num.start()].rstrip()
            else: eq_num = ""

        clean = raw_text
        clean = re.sub(r"«OMML:([A-Za-z0-9+/=]+)»", "", clean)
        clean = re.sub(r"\\begin\{equation\*?\}", "", clean)
        clean = re.sub(r"\\end\{equation\*?\}", "", clean)
        clean = re.sub(r"\\tag\{([^\}]*)\}", "", clean)
        clean = re.sub(r"\\\[", "", clean)
        clean = re.sub(r"\\\]", "", clean)
        clean = self._latex_math_to_readable(clean).strip()

        p = anchor_p.insert_paragraph_before()
        eq_style = self._pick_style_name(["equation", "Equation"])
        if eq_style:
            try: p.style = eq_style
            except: pass

        r = p.add_run("\t")
        if omml_match:
            try: self._insert_omml_to_paragraph(p, omml_match.group(1))
            except:
                if clean: p.add_run(clean)
        elif clean:
            p.add_run(clean)

        if eq_num:
            p.add_run(f"\t{eq_num}")


    def _sync_template_header_title(self, doc: Document, metadata: Dict[str, Any]) -> None:
        import re
        title = (metadata.get("title") or "").strip()
        title_plain = self._latex_to_plain(title).strip()
        
        authors = metadata.get("authors") or []
        short_authors = ""
        
        if authors:
            valid_authors = [self._clean_author_name(str(a.get("name") or "")).strip() for a in authors if self._clean_author_name(str(a.get("name") or "")).strip()]
            if len(valid_authors) == 1:
                short_authors = valid_authors[0]
            elif len(valid_authors) == 2:
                short_authors = f"{valid_authors[0]} and {valid_authors[1]}"
            elif len(valid_authors) > 2:
                short_authors = f"{valid_authors[0]} et al."
        
        for section in doc.sections:
            header = section.header
            if not header.is_linked_to_previous:
                for p in header.paragraphs:
                    if "contribution title (shortened if too long)" in p.text.lower():
                        new_text = re.sub(
                            r"contribution title \(shortened if too long\)", 
                            title_plain if title_plain else "Paper Title", 
                            p.text, 
                            flags=re.IGNORECASE
                        )
                        self._replace_paragraph_text_keep_formatting(p, new_text)
                        
            even_header = section.even_page_header
            if even_header and not even_header.is_linked_to_previous:
                for p in even_header.paragraphs:
                    if "f. author and s. author" in p.text.lower() or "f. author" in p.text.lower():
                        new_text = re.sub(
                            r"(F\.|f\.)\s*author\s*and\s*(S\.|s\.)\s*author", 
                            short_authors if short_authors else "Authors", 
                            p.text, 
                            flags=re.IGNORECASE
                        )
                        if new_text == p.text:
                            new_text = short_authors
                        if new_text:
                            self._replace_paragraph_text_keep_formatting(p, new_text)

    def _configure_ieee_document(self, doc: Document) -> None:
        """Configure a Springer-like one-column document."""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        self._set_section_columns(section, 1)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(4)
        normal_style.paragraph_format.line_spacing = 1.15

    def _start_two_column_body(self, doc: Document) -> None:
        """Springer body stays one-column (no IEEE two-column switch)."""
        if not doc.sections:
            return
        for sec in doc.sections:
            self._set_section_columns(sec, 1)

    def _add_title_section(self, doc: Document, metadata: Dict[str, Any]) -> None:
        title = (metadata.get("title") or "").strip()
        if not title:
            return
        p = doc.add_paragraph()
        title_style = self._pick_style_name(["papertitle", "paper title", "Title"])
        if title_style:
            try:
                p.style = title_style
            except Exception:
                pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self._latex_to_plain(title))
        if not self._using_uploaded_template:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(8)

    def _add_authors_table(self, doc: Document, authors: List[Dict[str, Any]]) -> None:
        if not authors:
            return
        
        author_style = self._pick_style_name(["author", "Author"])
        address_style = self._pick_style_name(["address", "Address", "institute"])
        email_style = self._pick_style_name(["email", "Email", "address", "Normal"])

        # 1. Extract valid authors and collect all unique affiliation lines (excluding emails)
        valid_authors = []
        extra_affs = []
        
        for a in authors:
            raw_name = str(a.get("name") or "").strip()
            if not raw_name:
                continue
            cleaned = self._clean_author_name(raw_name)
            if cleaned and self._looks_like_person_name(cleaned):
                valid_authors.append(a)
            elif cleaned:
                extra_affs.append(cleaned)

        unique_affs: List[str] = []
        author_aff_map: List[List[int]] = [] # Maps each valid author to their 1-based affiliation indices

        for a in valid_authors:
            idx_list = []
            affs = [
                self._latex_to_plain(str(x)).strip()
                for x in (a.get("affiliations") or [])
                if str(x).strip() and not self._looks_like_pure_author_name(str(x))
            ]
            aff_line = "; ".join(affs)
            if aff_line:
                if aff_line not in unique_affs:
                    unique_affs.append(aff_line)
                idx_list.append(unique_affs.index(aff_line) + 1)
            author_aff_map.append(idx_list)

        for aff in extra_affs:
            if aff not in unique_affs and not self._looks_like_email(aff):
                unique_affs.append(aff)

        # 2. Render Author Line with Superscripts
        if valid_authors:
            p_name = doc.add_paragraph()
            if author_style:
                try:
                    p_name.style = author_style
                except Exception:
                    pass
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for i, a in enumerate(valid_authors):
                name = self._clean_author_name(str(a.get("name") or ""))
                
                # Add joining text before name
                if i > 0:
                    if i == len(valid_authors) - 1:
                        p_name.add_run(" and " if len(valid_authors) == 2 else ", and ")
                    else:
                        p_name.add_run(", ")

                run = p_name.add_run(name)
                if not self._using_uploaded_template:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

                # Add superscripts
                indices = author_aff_map[i]
                if len(unique_affs) > 1 and indices:
                    idx_str = ",".join(str(idx) for idx in indices)
                    run_sup = p_name.add_run(idx_str)
                    run_sup.font.superscript = True

        # 3. Render Address Lines
        if unique_affs:
            for i, aff in enumerate(unique_affs, start=1):
                p_aff = doc.add_paragraph()
                if address_style:
                    try:
                        p_aff.style = address_style
                    except Exception:
                        pass
                p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Only use numbers if there are multiple affiliations
                prefix = f"{i} " if len(unique_affs) > 1 else ""
                run_aff = p_aff.add_run(prefix + aff)
                
                # Usually Springer italicizes the address if using normal fallback fonts
                if not self._using_uploaded_template:
                    run_aff.font.name = "Times New Roman"
                    run_aff.font.size = Pt(9)
                    run_aff.italic = True

        # 4. Render Email Line
        all_emails = []
        for a in authors:
            email = self._latex_to_plain(str(a.get("email") or "")).strip()
            if email and self._looks_like_email(email):
                all_emails.append(email)
        
        all_emails = list(dict.fromkeys(all_emails))
        if all_emails:
            p_email = doc.add_paragraph(", ".join(all_emails))
            if email_style:
                try:
                    p_email.style = email_style
                except Exception:
                    try:
                        p_email.style = address_style
                    except Exception:
                        pass
            p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not self._using_uploaded_template:
                for run in p_email.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
        
        doc.add_paragraph("")

    def _add_abstract_and_keywords(self, doc: Document, metadata: Dict[str, Any]) -> None:
        abstract = self._latex_to_plain(metadata.get("abstract") or "")
        abstract = re.sub(r"^\s*(?:abstract|t[oó]m\s+t[aắ]t)\s*[:.\u2013\u2014\-]+\s*", "", abstract, flags=re.IGNORECASE)
        abstract = re.sub(r"^\s*[-\u2013\u2014]{2,}\s*", "", abstract).strip()
        if abstract:
            p = doc.add_paragraph()
            abs_style = self._pick_style_name(["abstract", "Abstract"])
            if abs_style:
                try:
                    p.style = abs_style
                except Exception:
                    pass
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_label = p.add_run("Abstract. ")
            if not self._using_uploaded_template:
                run_label.bold = True
                run_label.font.name = "Times New Roman"
                run_label.font.size = Pt(10)
            run_body = p.add_run(abstract)
            if not self._using_uploaded_template:
                run_body.font.name = "Times New Roman"
                run_body.font.size = Pt(10)

        keywords = self._sanitize_keywords(metadata.get("keywords") or [])
        if keywords:
            kw = ", ".join(self._latex_to_plain(str(k)) for k in keywords if str(k).strip())
            kw = re.sub(r"^\s*keywords?\s*[:\-–—]*\s*", "", kw, flags=re.IGNORECASE)
            p = doc.add_paragraph()
            kw_style = self._pick_style_name(["keywords", "Keywords"])
            if kw_style:
                try:
                    p.style = kw_style
                except Exception:
                    pass
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_label = p.add_run("Keywords: ")
            if not self._using_uploaded_template:
                run_label.bold = True
                run_label.font.name = "Times New Roman"
                run_label.font.size = Pt(10)
            run_body = p.add_run(kw)
            if not self._using_uploaded_template:
                run_body.font.name = "Times New Roman"
                run_body.font.size = Pt(10)

    def _add_body(self, doc: Document, body_nodes: List[Dict[str, Any]]) -> None:
        for idx, node in enumerate(body_nodes):
            node_type = node.get("type", "")

            if node_type == "section":
                level = int(node.get("level", 1) or 1)
                level = max(1, min(3, level))
                text = self._latex_to_plain(node.get("text") or "")
                if text:
                    self._add_ieee_heading(doc, text, level)
                continue

            if node_type == "table":
                self._add_table_node(doc, node, force_full_width=False)
                continue

            if node_type == "list":
                self._add_list_node(doc, node)
                continue

            if node_type != "paragraph":
                continue

            raw_text = str(node.get("text") or "")
            plain = self._latex_to_plain(raw_text).strip()
            if not plain:
                continue

            if self._is_duplicate_table_title_paragraph(plain, body_nodes, idx):
                continue

            # Drop orphan IEEE caption labels that should be merged into table captions.
            if re.match(r"^\s*(?:TABLE|BANG|BẢNG)\s+[IVXLCDM\d]+\s*[:.\-]?\s*$", plain, re.IGNORECASE):
                continue

            if self._is_equation_like_paragraph(raw_text):
                self._add_equation_node(doc, raw_text)
                continue

            if "\\begin{figure" in raw_text or "\\includegraphics" in raw_text:
                self._add_figure_node(doc, raw_text)
                continue

            # Some IEEE extractions produce plain caption paragraph instead of figure node.
            if re.match(r"^Fig\.?\s*\d+\.?", plain, re.IGNORECASE):
                cap_clean = self._normalize_springer_caption(plain, "figure")
                fig_idx_match = re.match(r"^Fig\.?\s*(\d+)\.?", plain, re.IGNORECASE)
                fig_idx = int(fig_idx_match.group(1)) if fig_idx_match else max(1, self._figure_index + 1)
                self._figure_index = max(self._figure_index, fig_idx)
                p_cap = doc.add_paragraph(f"Fig. {fig_idx}. {cap_clean}" if cap_clean else f"Fig. {fig_idx}.")
                fig_style = self._pick_style_name(["figurecaption", "Figure Caption", "Caption"])
                if fig_style:
                    try:
                        p_cap.style = fig_style
                    except Exception:
                        pass
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            p = doc.add_paragraph()
            
            # Springer formatting: first paragraph is not indented ('p1a'), subsequent are indented ('Normal')
            is_first_para = (idx == 0) or (body_nodes[idx - 1].get("type") in ("section", "table", "figure"))
            style_candidates = ["p1a", "Normal"] if is_first_para else ["Normal", "p1a"]
            
            body_style = self._pick_style_name(style_candidates)
            if body_style:
                try:
                    p.style = body_style
                except Exception:
                    pass
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(plain)

    def _is_duplicate_table_title_paragraph(self, plain: str, body_nodes: List[Dict[str, Any]], idx: int) -> bool:
        next_node = body_nodes[idx + 1] if idx + 1 < len(body_nodes) else None
        if not next_node or next_node.get("type") != "table":
            return False

        caption = self._normalize_springer_caption(str(next_node.get("caption") or ""), "table")
        if not caption:
            return False

        p_norm = re.sub(r"[^a-z0-9]+", " ", plain.lower()).strip()
        c_norm = re.sub(r"[^a-z0-9]+", " ", caption.lower()).strip()
        if len(p_norm) < 12:
            return False

        return p_norm == c_norm or p_norm in c_norm or c_norm in p_norm

    def _add_ieee_heading(self, doc: Document, text: str, level: int) -> None:
        """Springer-like heading numbering (arabic)."""
        clean = self._latex_to_plain(text)
        if not clean:
            return
        clean = self._strip_existing_heading_number(clean)

        # Convert IEEE ALL CAPS headings (e.g. "INTRODUCTION") to title case for Springer
        if clean.isupper() and len(clean) > 3:
            clean = clean.title()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        heading_style = self._pick_style_name(
            [
                "heading1" if level == 1 else "heading2" if level == 2 else "Heading 3",
                f"Heading {min(level, 4)}",
            ]
        )
        if heading_style:
            try:
                p.style = heading_style
            except Exception:
                pass

        if level == 1:
            self._section_index += 1
            self._subsection_counters[self._section_index] = 0
            numbered = f"{self._section_index} {clean}"
            run = p.add_run(numbered)
            if not self._using_uploaded_template:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            return

        if self._section_index not in self._subsection_counters:
            self._subsection_counters[self._section_index] = 0
        self._subsection_counters[self._section_index] += 1
        sub_idx = self._subsection_counters[self._section_index]

        if level == 2:
            numbered = f"{self._section_index}.{sub_idx} {clean}"
            run = p.add_run(numbered)
            if not self._using_uploaded_template:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        else:
            run = p.add_run(clean)
            if not self._using_uploaded_template:
                run.bold = True
                run.italic = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)

    def _add_table_caption(self, doc: Document, label: str, caption: str) -> None:
        """Springer-style table caption as a single line."""
        index = getattr(self, "_table_index", 1)
        caption = self._normalize_springer_caption(caption, "table")
        text = f"Table {index}. {caption}" if caption else f"Table {index}."
        p = doc.add_paragraph()
        cap_style = self._pick_style_name(["tablecaption", "Table Caption", "Caption"])
        if cap_style:
            try:
                p.style = cap_style
            except Exception:
                pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        if not self._using_uploaded_template:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

    def _add_references(self, doc: Document, references: List[Dict[str, Any]]) -> None:
        if not references:
            return
        p = doc.add_paragraph("References")
        heading_style = self._pick_style_name(["heading1", "Heading 1", "headings"])
        if heading_style:
            try:
                p.style = heading_style
            except Exception:
                pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if not self._using_uploaded_template:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        ref_style = self._pick_style_name(["referenceitem", "referencelist", "ReferenceLine", "Normal"])

        for idx, ref in enumerate(references, start=1):
            if isinstance(ref, dict):
                text = self._latex_to_plain(ref.get("text") or "")
            else:
                text = self._latex_to_plain(str(ref))
            if not text:
                continue
            text = self._clean_reference_text(text)
            while True:
                updated = re.sub(r"^\s*\[?\d+\]?\s*[\.)]?\s*", "", text)
                if updated == text:
                    break
                text = updated
            text = re.sub(r"(?<!\s)(https?://)", r" \1", text)

            p = doc.add_paragraph()
            if ref_style:
                try:
                    p.style = ref_style
                except Exception:
                    pass

            # Some Springer templates auto-number referenceitem, some don't.
            # We'll just prepend '1. ' manually if it's not present, as LNCS expects numeric references
            # unless the style handles it automatically (we'll assume manual for safety).
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.left_indent = Inches(0.2)
            
            run_num = p.add_run(f"{idx}. ")
            run_text = p.add_run(text)
            
            if not self._using_uploaded_template:
                run_num.font.name = "Times New Roman"
                run_num.font.size = Pt(9)
                run_text.font.name = "Times New Roman"
                run_text.font.size = Pt(9)

    def _select_table_layout_mode(self, doc: Document, node: Dict[str, Any]) -> str:
        """Springer/LNCS is single-column, keep tables in-column to match template."""
        return "column"

    def _current_table_target_width_inch(self, doc: Document, force_full_width: bool) -> float:
        if not doc.sections:
            return 6.0
        sec = doc.sections[-1]
        try:
            content_width = (sec.page_width - sec.left_margin - sec.right_margin).inches
            return max(4.8, content_width)
        except Exception:
            return 6.0

    def _add_equation_node(self, doc: Document, raw_text: str) -> None:
        """Render equation using Springer equation paragraph style when available."""
        omml_match = re.search(r"«OMML:([A-Za-z0-9+/=]+)»", raw_text)
        tag_match = re.search(r"\\tag\{([^}]*)\}", raw_text)
        if tag_match:
            tag_text = re.sub(r"\s+", "", tag_match.group(1) or "")
            eq_num = f"({tag_text})" if tag_text else ""
        else:
            trailing_num = re.search(r"\((\d{1,3}[A-Za-z]?)\)\s*$", str(raw_text or ""))
            if trailing_num and "=" in str(raw_text or ""):
                eq_num = f"({trailing_num.group(1)})"
                raw_text = str(raw_text or "")[: trailing_num.start()].rstrip()
            else:
                eq_num = ""

        clean = raw_text
        clean = re.sub(r"«OMML:([A-Za-z0-9+/=]+)»", "", clean)
        clean = re.sub(r"\\begin\{equation\*?\}", "", clean)
        clean = re.sub(r"\\end\{equation\*?\}", "", clean)
        clean = re.sub(r"\\tag\{([^}]*)\}", "", clean)
        clean = re.sub(r"\\\[", "", clean)
        clean = re.sub(r"\\\]", "", clean)
        clean = self._latex_math_to_readable(clean).strip()

        p = doc.add_paragraph()
        eq_style = self._pick_style_name(["equation", "Equation"])
        if eq_style:
            try:
                p.style = eq_style
            except Exception:
                pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if omml_match:
            try:
                self._insert_omml_to_paragraph(p, omml_match.group(1))
            except Exception:
                if clean:
                    p.add_run(clean)
        elif clean:
            p.add_run(clean)

        if eq_num:
            p.add_run(f"  {eq_num}")

    def _is_equation_like_paragraph(self, raw_text: str) -> bool:
        if super()._is_equation_like_paragraph(raw_text):
            return True

        plain = self._latex_to_plain(str(raw_text or "")).strip()
        if not plain or len(plain) > 180:
            return False
        if re.search(r"https?://", plain, re.IGNORECASE):
            return False

        has_trailing_number = bool(re.search(r"\(\d{1,3}[A-Za-z]?\)\s*$", plain))
        has_math_relation = "=" in plain or any(sym in plain for sym in ["≤", "≥", "∑", "∫", "∞"])
        has_math_token = bool(re.search(r"\b(BA|F1|AUC|ROC|RMSE|MSE|Precision|Recall|Specificity|Sensitivity)\b", plain, re.IGNORECASE))

        if has_math_relation and (has_trailing_number or has_math_token):
            return True

        return False

    def _add_figure_node(self, doc: Document, latex_figure_text: str) -> None:
        """Render Springer figure with centered image and `figurecaption` style below."""
        self._figure_index += 1

        path_match = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", latex_figure_text)
        cap_match = re.search(r"\\caption\{([^}]*)\}", latex_figure_text)
        image_path = self._latex_to_plain(path_match.group(1)) if path_match else ""
        caption = self._normalize_springer_caption(
            self._latex_to_plain(cap_match.group(1)) if cap_match else "",
            "figure",
        )

        resolved = self._resolve_image_path(image_path)
        if resolved and resolved.exists():
            try:
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                width = max(3.2, min(5.6, self._current_table_target_width_inch(doc, False)))
                pic_para.add_run().add_picture(str(resolved), width=Inches(width))
            except Exception:
                pass

        cap_text = f"Fig. {self._figure_index}. {caption}" if caption else f"Fig. {self._figure_index}."
        p = doc.add_paragraph(cap_text)
        fig_style = self._pick_style_name(["figurecaption", "Figure Caption", "Caption"])
        if fig_style:
            try:
                p.style = fig_style
            except Exception:
                pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _sanitize_keywords(self, keywords: List[Any]) -> List[str]:
        cleaned: list[str] = []
        for item in keywords:
            raw = self._latex_to_plain(str(item or "")).strip()
            if not raw:
                continue
            # Guard against parser over-capturing whole body into keywords.
            candidates = [x.strip(" .;:-") for x in raw.split(",") if x.strip()]
            for c in candidates:
                c = re.sub(r"^keywords?\s*[-–—:]*\s*", "", c, flags=re.IGNORECASE).strip()
                if not c:
                    continue
                if len(c) > 80:
                    continue
                if len(c.split()) > 8:
                    continue
                if "." in c:
                    continue
                if c.lower() in {"however", "therefore", "thus", "and", "or", "are", "is"}:
                    continue
                if re.search(r"\b(section|introduction|references|table|figure)\b", c, re.IGNORECASE):
                    continue
                cleaned.append(c)
        return list(dict.fromkeys(cleaned))

    def _normalize_springer_caption(self, caption: str, kind: str) -> str:
        text = (caption or "").replace("\n", " ").strip()
        if not text:
            return ""

        if kind == "table":
            text = re.sub(r"^\s*table\s*(?:[ivxlcdm]{1,8}|\d+)\b\s*[:.\-]?\s*", "", text, flags=re.IGNORECASE)
            text = re.sub(r"^\s*TABLE\s*(?:[IVXLCDM]{1,8}|\d+)\b\s*[:.\-]?\s*", "", text)
        else:
            text = re.sub(r"^\s*fig(?:ure)?\.?\s*\d+\s*[:.\-]?\s*", "", text, flags=re.IGNORECASE)

        text = re.sub(r"\s+", " ", text).strip(" .")
        return text

    def _strip_existing_heading_number(self, text: str) -> str:
        clean = (text or "").strip()
        # Remove duplicated numbering patterns like "3 3 TITLE" or "3.2 3.2 TITLE".
        clean = re.sub(r"^\s*(\d+(?:\.\d+)?)\s+\1\s+", "", clean)
        # Remove existing (Heading X) artifacts from template texts
        clean = re.sub(r"\s*\(Heading\s+\d+\)\s*", "", clean, flags=re.IGNORECASE)
        # Remove one leading numbering token that likely comes from source style.
        clean = re.sub(r"^\s*(?:\d+(?:\.\d+)?|[IVXLCDM]+|[A-Z])\s*[.)]?\s+", "", clean)
        return clean.strip()

    def _looks_like_affiliation_text(self, text: str) -> bool:
        value = (text or "").lower()
        if self._looks_like_email(value):
            return True
        hints = [
            "university", "institute", "department", "faculty", "school",
            "city", "country", "street", "road", "heidelberg", "princeton",
            "vietnam", "việt", "viet nam", "college", "laboratory", "lab",
            "ltd", "inc", "corp", "company"
        ]
        return any(h in value for h in hints) or bool(re.search(r"\d", value))

    def _looks_like_person_name(self, text: str) -> bool:
        value = (text or "").strip()
        if not value:
            return False
        if self._looks_like_email(value):
            return False

        normalized = re.sub(r"\[[^\]]+\]", "", value)
        normalized = re.sub(r"\d+", "", normalized).strip()
        words = [w for w in re.split(r"\s+", normalized) if w]
        if len(words) < 2 or len(words) > 6:
            return False

        lower = normalized.lower()
        bad_tokens = {
            "university", "institute", "department", "faculty", "school",
            "springer", "street", "road", "city", "country",
            "usa", "germany", "vietnam", "heidelberg", "princeton", "tiergartenstr", "nj",
        }
        if any(tok in lower for tok in bad_tokens):
            return False
        if re.search(r"\b\d{4,}\b", value):
            return False
        return True

    def _looks_like_pure_author_name(self, text: str) -> bool:
        value = self._clean_author_name(text)
        if not value:
            return False
        if self._looks_like_affiliation_text(value):
            return False
        words = [w for w in re.split(r"\s+", value) if w]
        return 1 <= len(words) <= 6 and all(not re.search(r"\d", w) for w in words)

    def _clean_author_name(self, raw_name: str) -> str:
        text = super()._clean_author_name(raw_name)
        text = re.sub(r"\[[^\]]+\]", "", text)
        text = re.sub(r"\s+", " ", text).strip(" ,;")
        return text
