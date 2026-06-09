# xu_ly_bang.py - Bộ xử lý bảng: phân loại bảng + chuyển đổi sang LaTeX

import os
import re

from docx.table import Table
from docx.oxml.ns import qn

from .config import OMML_NAMESPACE, W_NAMESPACE, OLE_NAMESPACE, VML_NAMESPACE, R_NAMESPACE, A_NAMESPACE, REL_NAMESPACE
from .utils import loc_ky_tu, phat_hien_loai_tai_lieu
from .xu_ly_ole_equation import ole_equation_to_latex


class BoXuLyBang:
    """Bộ xử lý bảng, tách khỏi controller để đảm bảo SRP."""

    def __init__(self, bo_chuyen, doc_class: str = ""):
        """Khởi tạo với tham chiếu ChuyenDoiWordSangLatex để tái sử dụng logic.

        Args:
            bo_chuyen: Tham chiếu đến ChuyenDoiWordSangLatex.
            doc_class: Loại template ('ieee', 'springer', 'acm', ...). Nếu rỗng,
                       sẽ tự suy ra từ đường dẫn template của bo_chuyen.
        """
        self.bo_chuyen = bo_chuyen
        # Suy ra doc_class từ template nếu chưa được cung cấp
        if doc_class:
            self.doc_class = doc_class
        else:
            try:
                template_path = getattr(bo_chuyen, 'duong_dan_template', '')
                template_src = open(template_path, 'r', encoding='utf-8', errors='ignore').read()
                self.doc_class = phat_hien_loai_tai_lieu(template_src)
            except Exception:
                self.doc_class = "generic"

    def la_table_of_contents(self, bang: Table) -> bool:
        """Phát hiện bảng Mục lục (TOC) dựa trên từ khóa và cấu trúc."""
        try:
            if len(bang.rows) < 5:
                return False

            if self.bo_chuyen.tong_so_phan_tu > 0:
                vi_tri_phan_tram = (self.bo_chuyen.vi_tri_hien_tai / self.bo_chuyen.tong_so_phan_tu) * 100
                if vi_tri_phan_tram > 30:
                    return False

            toan_bo_text = ''
            for hang in bang.rows[:min(5, len(bang.rows))]:
                for cell in hang.cells:
                    toan_bo_text += (cell.text or "").strip().upper() + ' '

            co_tu_khoa_toc = False
            tu_khoa_muc_luc = ['MỤC LỤC', 'TABLE OF CONTENTS']
            for tu in tu_khoa_muc_luc:
                if tu in toan_bo_text:
                    co_tu_khoa_toc = True
                    break

            dem_dau_cham = 0
            dem_so_trang_cuoi = 0
            dem_cau_truc_muc = 0

            for hang in bang.rows[:min(20, len(bang.rows))]:
                if len(hang.cells) == 0:
                    continue

                text_hang = ''.join([c.text for c in hang.cells])

                if '.....' in text_hang or '…' in text_hang:
                    dem_dau_cham += 1

                if len(hang.cells) >= 2:
                    cell_cuoi = (hang.cells[-1].text or "").strip()
                    if cell_cuoi.isdigit() and 1 <= len(cell_cuoi) <= 4:
                        dem_so_trang_cuoi += 1

                    cell_dau = (hang.cells[0].text or "").strip().upper()
                    if re.search(r'(CH[UƯ][ƠƯ]NG|CHAPTER|PH[ẦẦ]N|PART|M[ỤỦ]C)\s*\d', cell_dau):
                        dem_cau_truc_muc += 1
                    if re.search(r'^\d+\.?\d*\.?\s+[A-ZÀ-Ỹ]', cell_dau):
                        dem_cau_truc_muc += 1

            so_hang_kiem_tra = min(20, len(bang.rows))

            if co_tu_khoa_toc:
                if dem_dau_cham >= 3 or dem_so_trang_cuoi >= 5:
                    return True

            if (dem_dau_cham > so_hang_kiem_tra * 0.5
                    and dem_so_trang_cuoi > so_hang_kiem_tra * 0.5
                    and dem_cau_truc_muc >= 3):
                return True
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 78: {e}')
            pass
        return False

    def la_bang_chua_anh(self, bang: Table) -> bool:
        """Phát hiện bảng chứa chủ yếu ảnh (figure layout)."""
        try:
            so_cell_co_anh = 0
            so_cell_co_text_dai = 0
            tong_cell = 0
            cells_da_kiem = set()

            for hang in bang.rows:
                for cell in hang.cells:
                    cell_id = id(getattr(cell, '_tc', cell))
                    if cell_id in cells_da_kiem:
                        continue
                    cells_da_kiem.add(cell_id)

                    tong_cell += 1
                    cell_text = (cell.text or "").strip()
                    co_anh = False

                    for para in cell.paragraphs:
                        for run in para.runs:
                            blips = run._element.findall(f'.//{{{A_NAMESPACE}}}blip')
                            if blips:
                                co_anh = True
                                break
                        drawings = para._element.findall(f'.//{{{A_NAMESPACE}}}blip')
                        if drawings:
                            co_anh = True

                    if co_anh:
                        so_cell_co_anh += 1

                    if re.match(r'^[\(\[]*[a-zA-Z0-9][\)\]]*\.?$', cell_text):
                        pass
                    elif re.match(r'^(Hình|Figure|Fig|Bảng|Table)\s*\d+', cell_text, re.IGNORECASE):
                        pass
                    elif len(cell_text) > 20:
                        so_cell_co_text_dai += 1

            if tong_cell == 0:
                return False

            if so_cell_co_anh >= 1:
                if so_cell_co_text_dai <= 1:
                    return True
                if so_cell_co_anh / tong_cell >= 0.3:
                    return True
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 129: {e}')
            pass
        return False

    def la_bang_layout(self, bang: Table) -> bool:
        """Phát hiện bảng layout metadata (đầu bài báo)."""
        try:
            if self.bo_chuyen.tong_so_phan_tu > 0:
                vi_tri_phan_tram = (self.bo_chuyen.vi_tri_hien_tai / self.bo_chuyen.tong_so_phan_tu) * 100
                if vi_tri_phan_tram > 25:
                    return False

            # Bảng dữ liệu thật: nhiều hàng hoặc nhiều cột → không phải layout
            so_hang = len(bang.rows)
            so_cot = len(bang.columns) if bang.rows else 0
            if so_hang >= 10:
                return False
            if so_cot >= 4:
                return False

            toan_bo_text = ''
            for hang in bang.rows[:min(10, len(bang.rows))]:
                for cell in hang.cells:
                    toan_bo_text += (cell.text or "").strip().upper() + ' '

            tu_khoa_layout = [
                'ARTICLE INFORMATION', 'ARTICLE TITLE', 'JOURNAL:',
                'ISSN:', 'ABSTRACT', 'KEYWORDS:', 'TỪ KHÓA:',
                'AUTHOR', 'AFFILIATION', 'CORRESPONDENCE', 'CITATION',
                'RECEIVED:', 'ACCEPTED:', 'PUBLISHED:', 'DOI:',
                'OPEN ACCESS', 'TÓM TẮT', 'VOLUME:', 'ISSUE:',
            ]

            dem_tu_khoa = sum(1 for tu in tu_khoa_layout if tu in toan_bo_text)
            if dem_tu_khoa >= 3:
                return True
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 165: {e}')
            pass
        return False

    def la_bang_metadata(self, bang: Table) -> bool:
        """Nhận diện bảng Metadata dựa trên từ khóa đầu bài."""
        try:
            if self.bo_chuyen.dem_bang > 5:
                return False

            text_dau = ""
            for hang in bang.rows[:2]:
                for cell in hang.cells:
                    text_dau += (cell.text or "").upper() + " "

            tu_khoa = [
                "ARTICLE INFO", "ARTICLE INFORMATION", "ABSTRACT",
                "TÓM TẮT", "THÔNG TIN BÀI BÁO",
            ]
            for k in tu_khoa:
                if k in text_dau:
                    return True
            return False
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 188: {e}')
            return False

    def tao_latex_minipage_metadata(self, bang: Table) -> str:
        """Chuyển bảng Metadata thành hai cột minipage (Info | Abstract)."""
        try:
            latex = []
            latex.append(r"\vspace{0.5cm}")
            latex.append(r"\noindent")

            noi_dung_cot_1 = []
            noi_dung_cot_2 = []

            for hang in bang.rows:
                cells = hang.cells
                if len(cells) >= 1:
                    noi_dung_cot_1.append(self.xu_ly_doan_van_trong_cell(cells[0]))
                if len(cells) >= 2:
                    noi_dung_cot_2.append(self.xu_ly_doan_van_trong_cell(cells[1]))

            text_trai = "\n".join(noi_dung_cot_1)
            text_phai = "\n".join(noi_dung_cot_2)

            latex.append(r"\begin{minipage}[t]{0.30\textwidth}")
            latex.append(text_trai)
            latex.append(r"\end{minipage}")
            latex.append(r"\hfill")
            latex.append(r"\begin{minipage}[t]{0.65\textwidth}")
            latex.append(text_phai)
            latex.append(r"\end{minipage}")
            latex.append(r"\vspace{0.5cm}")

            return "\n".join(latex)
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 221: {e}')
            return ""

    def la_bang_tieu_su(self, bang: Table) -> bool:
        """Nhận diện bảng tiểu sử tác giả (ảnh + đoạn text dài)."""
        try:
            if len(bang.columns) != 2:
                return False

            cells = bang.rows[0].cells
            co_anh = False
            for cell in cells:
                for para in cell.paragraphs:
                    if para._element.findall(f'.//{{{A_NAMESPACE}}}blip'):
                        co_anh = True
                        break
                if co_anh:
                    break

            cells_hang_0 = bang.rows[0].cells
            if len(cells_hang_0) < 2:
                return False
            text_len = len(cells_hang_0[0].text) + len(cells_hang_0[1].text)
            return co_anh and text_len > 50
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 242: {e}')
            return False

    def trich_xuat_anh_trong_cell(self, cell) -> list:
        """Trích xuất ảnh từ một ô bảng bằng cách duyệt các paragraph."""
        danh_sach_anh = []
        for para in cell.paragraphs:
            anh_list, _ = self.bo_chuyen.trich_xuat_anh(para)
            danh_sach_anh.extend(anh_list)
        return danh_sach_anh

    def tao_latex_tieu_su_tac_gia(self, bang: Table) -> str:
        """Tạo layout tiểu sử tác giả bằng minipage ảnh + text."""
        try:
            latex = []
            for hang in bang.rows:
                cells = hang.cells
                if len(cells) < 2:
                    continue

                anh_list_0 = self.trich_xuat_anh_trong_cell(cells[0])
                anh_list_1 = self.trich_xuat_anh_trong_cell(cells[1])

                noi_dung_text = ""
                file_anh = ""

                if anh_list_0:
                    file_anh = anh_list_0[0]
                    noi_dung_text = self.xu_ly_doan_van_trong_cell(cells[1], che_do_inline=True)
                elif anh_list_1:
                    file_anh = anh_list_1[0]
                    noi_dung_text = self.xu_ly_doan_van_trong_cell(cells[0], che_do_inline=True)
                else:
                    continue

                ten_thu_muc = os.path.basename(self.bo_chuyen.thu_muc_anh)
                latex.append(r"\vspace{0.3cm}")
                latex.append(r"\noindent")
                latex.append(r"\begin{minipage}[t]{0.2\textwidth}")
                latex.append(r"\vspace{0pt}")
                latex.append(rf"\includegraphics[width=\linewidth, height=3.5cm, keepaspectratio]{{{ten_thu_muc}/{file_anh}}}")
                latex.append(r"\end{minipage}")
                latex.append(r"\hfill")
                latex.append(r"\begin{minipage}[t]{0.75\textwidth}")
                latex.append(r"\vspace{0pt}")
                latex.append(noi_dung_text)
                latex.append(r"\end{minipage}")
                latex.append(r"\vspace{0.3cm}")

            return "\n".join(latex)
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 292: {e}')
            return ""

    def xu_ly_doan_van_trong_cell(self, cell, che_do_inline: bool = True) -> str:
        """Gộp và xử lý nội dung các paragraph bên trong một ô bảng."""
        noi_dung = []
        for p in cell.paragraphs:
            text = self.bo_chuyen.xu_ly_doan_van(p, che_do_inline=che_do_inline)
            if text:
                noi_dung.append(text)
        return "\n".join(noi_dung)

    def la_bang_cong_thuc(self, bang: Table) -> bool:
        """Phát hiện bảng công thức: 2 cột, cột cuối là số thứ tự (1), (2)..."""
        try:
            if len(bang.columns) != 2:
                return False

            dem_so_thu_tu = 0
            for hang in bang.rows:
                if len(hang.cells) >= 2:
                    cell_cuoi = (hang.cells[-1].text or "").strip()
                    if re.match(r'^\(\d+\)$', cell_cuoi):
                        dem_so_thu_tu += 1

            if len(bang.rows) > 0 and dem_so_thu_tu / len(bang.rows) >= 0.5:
                return True
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 319: {e}')
            pass
        return False

    def trich_xuat_noi_dung_bang_layout(self, bang: Table) -> str:
        """Trích xuất nội dung text từ bảng layout (bỏ cấu trúc bảng)."""
        ket_qua = []
        da_xuat_para = set()
        try:
            for hang in bang.rows:
                for cell in hang.cells:
                    text = (cell.text or "").strip()
                    if text and len(text) > 2:
                        for para in cell.paragraphs:
                            noi_dung = ""
                            for run in para.runs:
                                noi_dung += self.bo_chuyen.xu_ly_run(run)
                            noi_dung_clean = noi_dung.strip()
                            if noi_dung_clean and noi_dung_clean not in da_xuat_para:
                                da_xuat_para.add(noi_dung_clean)
                                ket_qua.append(noi_dung + "\n\n")
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 340: {e}')
            pass
        return ''.join(ket_qua)

    def trich_xuat_omml_tu_cell(self, cell) -> str:
        """Trích xuất công thức (OMML hoặc OLE) từ một cell của bảng."""
        cong_thuc_parts = []
        try:
            for para in cell.paragraphs:
                omath_list = para._element.findall(f'.//{{{OMML_NAMESPACE}}}oMath')
                for omath in omath_list:
                    latex = self.bo_chuyen.bo_toan.omml_element_to_latex(omath)
                    if latex.strip():
                        cong_thuc_parts.append(latex)

                if not omath_list:
                    objects = para._element.findall(f'.//{{{W_NAMESPACE}}}object')
                    for obj in objects:
                        ole = obj.find(f'.//{{{OLE_NAMESPACE}}}OLEObject')
                        if ole is not None:
                            ole_rid = ole.get(f'{{{R_NAMESPACE}}}id')
                            if ole_rid:
                                ole_part = self.bo_chuyen.tai_lieu.part.related_parts.get(ole_rid)
                                if ole_part:
                                    try:
                                        latex_from_mtef = ole_equation_to_latex(ole_part.blob)
                                        if latex_from_mtef.strip():
                                            cong_thuc_parts.append(latex_from_mtef)
                                            continue
                                    except Exception as e:
                                        print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 369: {e}')
                                        pass

                            imagedata = obj.find(f'.//{{{VML_NAMESPACE}}}imagedata')
                            if imagedata is not None:
                                rid = imagedata.get(f'{{{R_NAMESPACE}}}id')
                                if rid:
                                    part = self.bo_chuyen.tai_lieu.part.related_parts.get(rid)
                                    if part:
                                        self.bo_chuyen.dem_anh += 1
                                        content_type = getattr(part, 'content_type', '')
                                        ext = 'png'
                                        if 'wmf' in content_type:
                                            ext = 'wmf'
                                        elif 'emf' in content_type:
                                            ext = 'emf'
                                        elif 'jpeg' in content_type:
                                            ext = 'jpg'

                                        ten_anh_goc = f'formula_{self.bo_chuyen.dem_anh}.{ext}'
                                        if not os.path.exists(self.bo_chuyen.thu_muc_anh):
                                            os.makedirs(self.bo_chuyen.thu_muc_anh, exist_ok=True)

                                        duong_dan_anh = os.path.join(self.bo_chuyen.thu_muc_anh, ten_anh_goc)
                                        with open(duong_dan_anh, 'wb') as f:
                                            f.write(part.blob)

                                        ten_anh = ten_anh_goc
                                        if ext in ('wmf', 'emf'):
                                            try:
                                                from PIL import Image
                                                img = Image.open(duong_dan_anh)
                                                ten_anh = f'formula_{self.bo_chuyen.dem_anh}.png'
                                                duong_dan_png = os.path.join(self.bo_chuyen.thu_muc_anh, ten_anh)
                                                new_size = (img.size[0] * 3, img.size[1] * 3)
                                                # Pillow >= 10: LANCZOS chuyển vào Image.Resampling
                                                _resample = getattr(Image.Resampling, 'LANCZOS', None) or getattr(Image, 'LANCZOS', None) or getattr(Image, 'BICUBIC', 3)
                                                img_resized = img.resize(new_size, _resample)
                                                img_resized.save(duong_dan_png)
                                                os.remove(duong_dan_anh)
                                            except Exception as e:
                                                print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 407: {e}')
                                                pass

                                        ten_thu_muc = os.path.basename(self.bo_chuyen.thu_muc_anh)
                                        cong_thuc_parts.append(
                                            rf'\\includegraphics[height=1.5em]{{{ten_thu_muc}/{ten_anh}}}'
                                        )

                if not omath_list and not cong_thuc_parts:
                    text = (para.text or "").strip()
                    if text:
                        cong_thuc_parts.append(loc_ky_tu(text))

            if not cong_thuc_parts:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    cong_thuc_parts.append(loc_ky_tu(cell_text))
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 424: {e}')
            pass
        return ' '.join(cong_thuc_parts)

    def xu_ly_bang_cong_thuc(self, bang: Table) -> str:
        # Chuyển bảng công thức thành equation environment (có \tag)
        latex = ""
        try:
            for hang in bang.rows:
                if len(hang.cells) >= 2:
                    cong_thuc = self.trich_xuat_omml_tu_cell(hang.cells[0])
                    cells = hang.cells
                    cell_so = (cells[1].text or "").strip()

                    so_match = re.match(r'^\((\d+)\)$', cell_so)
                    if so_match:
                        so = so_match.group(1)
                        latex += r"\begin{equation}" + "\n"
                        if cong_thuc.strip():
                            latex += f"  {cong_thuc}\n"
                        else:
                            latex += f"  \\text{{[Công thức {so}]}}\n"
                        latex += rf"  \tag{{{so}}}\label{{eq:eq_{so}}}" + "\n"
                        latex += r"\end{equation}" + "\n\n"
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 447: {e}')
            pass
        return latex

    def trich_xuat_anh_tu_bang(self, bang: Table) -> list:
        # Trích xuất ảnh từ bảng (figure layout), luôn giữ ảnh
        danh_sach_anh = []
        if not self.bo_chuyen.tai_lieu:
            return danh_sach_anh

        for hang in bang.rows:
            for cell in hang.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        blips = run._element.findall(f'.//{{{A_NAMESPACE}}}blip')
                        for blip in blips:
                            embed = blip.get(f'{{{REL_NAMESPACE}}}embed')
                            if not embed:
                                continue
                            part = self.bo_chuyen.tai_lieu.part.related_parts.get(embed)
                            if not part:
                                continue

                            self.bo_chuyen.dem_anh += 1
                            content_type = getattr(part, 'content_type', '')
                            ext = 'png'
                            if 'jpeg' in content_type:
                                ext = 'jpg'

                            ten_anh = f'hinh_{self.bo_chuyen.dem_anh}.{ext}'
                            if not os.path.exists(self.bo_chuyen.thu_muc_anh):
                                os.makedirs(self.bo_chuyen.thu_muc_anh, exist_ok=True)

                            duong_dan_anh = os.path.join(self.bo_chuyen.thu_muc_anh, ten_anh)
                            with open(duong_dan_anh, 'wb') as f:
                                f.write(part.blob)

                            danh_sach_anh.append(ten_anh)
        return danh_sach_anh

    def _tim_caption_con_trong_bang(self, bang: Table) -> list:
        # Tim text caption con (a), (b)... trong cac cell cua bang
        danh_sach = []
        try:
            for hang in bang.rows:
                for cell in hang.cells:
                    text = (cell.text or "").strip()
                    match = re.match(r'^\(([a-z])\)(.*)$', text)
                    if match:
                        nhan = match.group(1)
                        mo_ta = match.group(2).strip()
                        caption = f"({nhan})"
                        if mo_ta:
                            caption += f" {mo_ta}"
                        danh_sach.append(loc_ky_tu(caption))
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 502: {e}')
            pass
        return danh_sach

    def _lay_gridspan(self, tc) -> int:
        # Lấy colspan từ w:gridSpan
        try:
            tcPr = tc.tcPr
            if tcPr is None:
                return 1
            gridSpan = tcPr.gridSpan
            if gridSpan is None:
                return 1
            val = gridSpan.get(qn('w:val'))
            if val is None:
                return 1
            return max(1, int(val))
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 519: {e}')
            return 1

    def _lay_vmerge(self, tc):
        # Lấy trạng thái vMerge: None / 'restart' / 'continue'
        try:
            tcPr = tc.tcPr
            if tcPr is None:
                return None
            vMerge = tcPr.vMerge
            if vMerge is None:
                return None
            val = vMerge.get(qn('w:val'))
            if val is None:
                return 'continue'
            return str(val)
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 535: {e}')
            return None

    def _xay_dung_luoi_o(self, bang: Table):
        # Xây dựng ma trận cell theo vị trí thực để xử lý rowspan/colspan
        # Dùng getattr để bypass type stubs: _tbl là thuộc tính internal không được khai báo trong stubs
        tbl = getattr(bang, '_tbl', None)
        if tbl is None:
            return [], {}, {}, 0, 0
        # Dùng getattr để bypass type stubs của python-docx (tr_lst không có trong stubs)
        tr_list = list(getattr(tbl, 'tr_lst', []))

        # Ước lượng số cột theo tblGrid, fallback theo số tc lớn nhất
        # Dùng findall trên XML thay vì gridCol_lst (không có trong type stubs của python-docx)
        so_cot = 0
        try:
            from docx.oxml.ns import qn as _qn
            grid_els = tbl.tblGrid.findall(_qn('w:gridCol'))
            so_cot = len(grid_els)
        except Exception as e:
            print(f'[Cảnh báo] Lỗi im lặng ở xu_ly_bang.py dòng 548: {e}')
            so_cot = 0

        if so_cot <= 0:
            for tr in tr_list:
                so_cot = max(so_cot, len(list(getattr(tr, 'tc_lst', []))))

        # Khai báo kiểu tường minh để tránh lỗi type-checker khi gán cell_id (int)
        from typing import Optional
        luoi: list[list[Optional[int]]] = [[None for _ in range(so_cot)] for _ in range(len(tr_list))]
        meta = {}

        for r, tr in enumerate(tr_list):
            c = 0
            for tc in list(getattr(tr, 'tc_lst', [])):
                while c < so_cot and luoi[r][c] is not None:
                    c += 1
                if c >= so_cot:
                    break

                colspan = self._lay_gridspan(tc)
                vmerge = self._lay_vmerge(tc)

                cell_id = id(tc)
                if vmerge in ('continue', 'cont') and r > 0 and luoi[r - 1][c] is not None:
                    cell_id = meta.get((r - 1, c), {}).get('id', cell_id)

                meta[(r, c)] = {
                    'id': cell_id,
                    'tc': tc,
                    'colspan': colspan,
                    'vmerge': vmerge,
                    'start': vmerge not in ('continue', 'cont'),
                }

                for k in range(colspan):
                    if c + k < so_cot:
                        luoi[r][c + k] = cell_id
                        if (r, c + k) not in meta:
                            meta[(r, c + k)] = meta[(r, c)]

                c += colspan

        # Tính rowspan cho các cell start
        rowspan_map = {}
        for (r, c), info in list(meta.items()):
            if not info.get('start'):
                continue
            if meta.get((r, c)) != info:
                continue

            cell_id = info['id']
            rowspan = 1
            rr = r + 1
            while rr < len(tr_list):
                info_down = meta.get((rr, c))
                if not info_down:
                    break
                if info_down.get('id') != cell_id:
                    break
                if info_down.get('start'):
                    break
                rowspan += 1
                rr += 1
            rowspan_map[cell_id] = max(rowspan_map.get(cell_id, 1), rowspan)

        return luoi, meta, rowspan_map, so_cot, len(tr_list)

    def _render_tabular_merge(self, bang: Table) -> str:
        # Render bảng dữ liệu sang tabular có hỗ trợ \multirow/\multicolumn
        luoi, meta, rowspan_map, so_cot, so_hang = self._xay_dung_luoi_o(bang)

        # Phân bố độ rộng cột theo tỉ lệ \linewidth để đảm bảo không bị tràn trang (hỗ trợ two-column)
        if so_cot > 0:
            width_frac = 0.98 / so_cot
        else:
            width_frac = 0.15
        
        # Bỏ kẻ sọc dọc '|' theo chuẩn booktabs
        # Sửa lại: Dùng đường kẻ dọc và ngang tiêu chuẩn (giống Word)
        cot = '|' + '|'.join([f"p{{{width_frac:.3f}\\linewidth}}" for _ in range(so_cot)]) + '|'

        # Nếu số cột > 4, khả năng cao cần chiếm 2 cột trong IEEE/ACM
        is_wide = so_cot > 4
        env_name = "table*" if is_wide else "table"
        scale_width = "\\textwidth" if is_wide else "\\columnwidth"

        # Theo chuẩn IEEEtran:
        #   - table* (bảng 2 cột) BẮT BUỘC dùng [t] hoặc [b], không được [H] hay [h]
        #   - table thường dùng [htbp] để LaTeX tự lấp đầy khoảng trắng
        vi_tri = "[t]" if is_wide else "[htbp]"

        latex = rf"\begin{{{env_name}}}{vi_tri}" + "\n"
        latex += r"  \centering" + "\n"
        
        # Caption above for tables (IEEE standard)
        caption = self.bo_chuyen.bat_caption_bang()
        if caption:
            caption_clean = re.sub(r'^(Bảng|Table)\s*\d+\s*[:\.\-–—]?\s*', '', caption, flags=re.IGNORECASE).strip()
            # Bảo vệ \url{} trong caption (Moving Argument)
            caption_clean = caption_clean.replace(r"\url{", r"\protect\url{")
            latex += rf"  \caption{{{caption_clean}}}" + "\n"
        
        latex += rf"  \resizebox{{{scale_width}}}{{!}}{{%." + "\n"
        latex += r"  \setlength{\arrayrulewidth}{0.4pt}" + "\n"
        latex += rf"  \begin{{tabular}}{{{cot}}}" + "\n"
        
        # Dùng \hline thay vì \toprule
        latex += r"  \hline" + "\n"

        # Khởi tạo ma trận đánh dấu các ô bị chiếm bởi rowspan
        occupied_cells = {} # (row, col) -> True
        
        for r in range(so_hang):
            tex_cells = []
            c_logical = 0   # Chỉ số cột thực tế
            
            while c_logical < so_cot:
                # Nếu ô này đã bị chiếm bởi rowspan từ phía trên
                if occupied_cells.get((r, c_logical)):
                    tex_cells.append("")
                    c_logical += 1
                    continue
                
                info = meta.get((r, c_logical))
                if not info or not info.get('start') or meta.get((r, c_logical)) != info:
                    tex_cells.append("")
                    c_logical += 1
                    continue
                
                colspan = int(info.get('colspan') or 1)
                cell_id = info['id']
                rowspan = int(rowspan_map.get(cell_id, 1))
                
                # Đánh dấu các ô bị chiếm trong tương lai
                for dr in range(rowspan):
                    for dc in range(colspan):
                        if dr > 0 or dc > 0:
                            occupied_cells[(r + dr, c_logical + dc)] = True
                
                try:
                    cell_obj = bang.rows[r].cells[0]
                    for candidate in bang.rows[r].cells:
                        if id(getattr(candidate, '_tc', None)) == id(info['tc']):
                            cell_obj = candidate
                            break
                except Exception:
                    cell_obj = None

                noi_dung = self.xu_ly_doan_van_trong_cell(cell_obj) if cell_obj else ""
                token = noi_dung.strip()
                token = token.replace("\r\n", "\n").replace("\r", "\n").replace("\n", r"\\ ")

                if rowspan > 1:
                    token = rf"\multirow{{{rowspan}}}{{*}}{{{token}}}"
                if colspan > 1:
                    mc_width = colspan * width_frac
                    if c_logical == 0:
                        mc_format = rf"|p{{{mc_width:.3f}\linewidth}}|"
                    else:
                        mc_format = rf"p{{{mc_width:.3f}\linewidth}}|"
                    token = rf"\multicolumn{{{colspan}}}{{{mc_format}}}{{{token}}}"

                tex_cells.append(token)
                c_logical += colspan

            # Loại bỏ trailing empty cells và lọc qua multicolumn skip logic
            dong_filtered = []
            skip_mc = 0
            for cell_str in tex_cells:
                if skip_mc > 0:
                    skip_mc -= 1
                    continue
                dong_filtered.append(cell_str)
                mc_match = re.match(r'\\multicolumn\{(\d+)\}', cell_str)
                if mc_match:
                    skip_mc = int(mc_match.group(1)) - 1
            
            latex += "    " + " & ".join(dong_filtered) + r" \\" + "\n"
            
            # Determine horizontal rule: use \cline if any multirow is spanning
            spanning_cols = set()
            for ci in range(so_cot):
                info_ci = meta.get((r, ci))
                if info_ci and info_ci.get('start'):
                    cell_id_ci = info_ci['id']
                    rspan = int(rowspan_map.get(cell_id_ci, 1))
                    if rspan > 1:
                        for dc in range(int(info_ci.get('colspan', 1) or 1)):
                            spanning_cols.add(ci + dc)
            
            if spanning_cols and r < so_hang - 1:
                cline_parts = []
                range_start = None
                for ci in range(so_cot):
                    if ci not in spanning_cols:
                        if range_start is None:
                            range_start = ci
                    else:
                        if range_start is not None:
                            cline_parts.append(rf"\cline{{{range_start + 1}-{ci}}}")
                            range_start = None
                if range_start is not None:
                    cline_parts.append(rf"\cline{{{range_start + 1}-{so_cot}}}")
                latex += "  " + "".join(cline_parts) + "\n" if cline_parts else "  " + r"\hline" + "\n"
            else:
                latex += r"  \hline" + "\n"
        latex += r"  \end{tabular}%" + "\n"
        latex += r"  }" + "\n"
        latex += rf"  \label{{tab:bang{self.bo_chuyen.dem_bang}}}" + "\n"
        latex += rf"\end{{{env_name}}}" + "\n\n"

        return latex

    # -------------------------------------------------------------------------
    # Phân loại và xử lý bảng siêu dài
    # -------------------------------------------------------------------------

    NGUONG_BANG_SIEU_DAI = 12  # Số hàng tối đa trước khi coi là "siêu dài"

    def _la_bang_sieu_dai(self, bang: Table) -> bool:
        """Phát hiện bảng có quá nhiều hàng, vượt quá chiều dọc 1 trang."""
        return len(bang.rows) > self.NGUONG_BANG_SIEU_DAI

    def _render_bang_sieu_dai(self, bang: Table) -> str:
        """
        Render bảng dài vượt trang bằng môi trường longtable.

        Phân nhánh theo doc_class:

        **Springer (LLNCS / svjour3 - 1 cột)**:
          - longtable hoạt động trực tiếp vì định dạng 1 cột.
          - Dùng \\linewidth (= \\textwidth trong 1 cột).
          - KHÔNG cần \\onecolumn / \\twocolumn.

        **IEEE / ACM (twocolumn)**:
          - longtable KHÔNG hoạt động trong chế độ twocolumn.
          - Phải dùng \\onecolumn -> longtable -> \\twocolumn.
          - Dùng \\textwidth (toàn trang sau khi thoát 2 cột).
        """
        try:
            luoi, meta, rowspan_map, so_cot, so_hang = self._xay_dung_luoi_o(bang)

            if so_cot <= 0:
                return self._render_tabular_merge(bang)

            # Phân nhánh: Springer 1-cột vs IEEE/ACM 2-cột
            is_springer = (self.doc_class == 'springer')

            if is_springer:
                # Springer LLNCS: 1 cột - longtable đơn giản với \linewidth
                width_unit = '\\linewidth'
                prefix_lines: list = []
                suffix_lines: list = []
            else:
                # IEEE / ACM twocolumn: phải thoát ra 1 cột trước
                width_unit = '\\textwidth'
                prefix_lines = [r'\onecolumn', '']
                suffix_lines = ['', r'\twocolumn', '']

            width_frac = 0.97 / so_cot
            cot_spec = (
                '|'
                + '|'.join([f'p{{{width_frac:.3f}{width_unit}}}' for _ in range(so_cot)])
                + '|'
            )

            caption = self.bo_chuyen.bat_caption_bang()
            if caption:
                caption_clean = re.sub(
                    r'^(Bảng|Table)\s*\d+\s*[:\.\-–—]?\s*', '', caption, flags=re.IGNORECASE
                ).strip()
            else:
                caption_clean = f'Bảng {self.bo_chuyen.dem_bang}'
            # Bảo vệ \url{} trong caption (Moving Argument)
            caption_clean = caption_clean.replace(r"\url{", r"\protect\url{")

            label = f'tab:bang{self.bo_chuyen.dem_bang}'

            latex = list(prefix_lines)
            latex.append(r'  \setlength{\arrayrulewidth}{0.4pt}')
            latex.append(rf'  \begin{{longtable}}{{{cot_spec}}}')
            # Caption trước nội dung - chuẩn IEEE và Springer
            latex.append(rf'  \caption{{{caption_clean}}}\label{{{label}}}\\\\')
            latex.append(r'  \hline')

            # ---- Render từng hàng ----
            occupied_cells: dict = {}

            for r in range(so_hang):
                tex_cells = []
                c_logical = 0

                while c_logical < so_cot:
                    if occupied_cells.get((r, c_logical)):
                        tex_cells.append('')
                        c_logical += 1
                        continue

                    info = meta.get((r, c_logical))
                    if not info or not info.get('start') or meta.get((r, c_logical)) != info:
                        tex_cells.append('')
                        c_logical += 1
                        continue

                    colspan = int(info.get('colspan') or 1)
                    cell_id = info['id']
                    rowspan = int(rowspan_map.get(cell_id, 1))

                    for dr in range(rowspan):
                        for dc in range(colspan):
                            if dr > 0 or dc > 0:
                                occupied_cells[(r + dr, c_logical + dc)] = True

                    try:
                        cell_obj = bang.rows[r].cells[0]
                        for candidate in bang.rows[r].cells:
                            # _tc là thuộc tính internal, dùng getattr để bypass stubs
                            if id(getattr(candidate, '_tc', None)) == id(info['tc']):
                                cell_obj = candidate
                                break
                    except Exception:
                        cell_obj = None

                    noi_dung = self.xu_ly_doan_van_trong_cell(cell_obj) if cell_obj else ''
                    token = (noi_dung.strip()
                             .replace('\r\n', '\n').replace('\r', '\n')
                             .replace('\n', r'\\ '))

                    if rowspan > 1:
                        token = rf'\multirow{{{rowspan}}}{{*}}{{{token}}}'
                    if colspan > 1:
                        mc_width = colspan * width_frac
                        mc_fmt = (
                            rf'|p{{{mc_width:.3f}{width_unit}}}|'
                            if c_logical == 0 else
                            rf'p{{{mc_width:.3f}{width_unit}}}|'
                        )
                        token = rf'\multicolumn{{{colspan}}}{{{mc_fmt}}}{{{token}}}'

                    tex_cells.append(token)
                    c_logical += colspan

                # Lọc multicolumn skip
                dong_filtered = []
                skip_mc = 0
                for cell_str in tex_cells:
                    if skip_mc > 0:
                        skip_mc -= 1
                        continue
                    dong_filtered.append(cell_str)
                    mc_match = re.match(r'\\multicolumn\{(\d+)\}', cell_str)
                    if mc_match:
                        skip_mc = int(mc_match.group(1)) - 1

                latex.append('    ' + ' & '.join(dong_filtered) + r' \\')

                # Đường kẻ ngang (\hline hoặc \cline)
                spanning_cols: set = set()
                for ci in range(so_cot):
                    info_ci = meta.get((r, ci))
                    if info_ci and info_ci.get('start'):
                        cell_id_ci = info_ci['id']
                        rspan = int(rowspan_map.get(cell_id_ci, 1))
                        if rspan > 1:
                            for dc in range(int(info_ci.get('colspan', 1) or 1)):
                                spanning_cols.add(ci + dc)

                if spanning_cols and r < so_hang - 1:
                    cline_parts = []
                    range_start = None
                    for ci in range(so_cot):
                        if ci not in spanning_cols:
                            if range_start is None:
                                range_start = ci
                        else:
                            if range_start is not None:
                                cline_parts.append(rf'\cline{{{range_start + 1}-{ci}}}')
                                range_start = None
                    if range_start is not None:
                        cline_parts.append(rf'\cline{{{range_start + 1}-{so_cot}}}')
                    latex.append('  ' + ''.join(cline_parts) if cline_parts else r'  \hline')
                else:
                    latex.append(r'  \hline')

            latex.append(r'  \end{longtable}')
            latex.extend(suffix_lines)

            return '\n'.join(latex) + '\n'

        except Exception as e:
            print(f'[Cảnh báo] Lỗi _render_bang_sieu_dai, fallback sang tabular: {e}')
            return self._render_tabular_merge(bang)


    def xu_ly_bang(self, bang: Table) -> str:
        """Phân loại và xử lý bảng theo 4 tầng:

        1. Bảng đặc biệt (Metadata, Tiểu sử, Layout, Công thức, TOC, Ảnh)
        2. Bảng dữ liệu siêu dài (>12 hàng) → onecolumn + longtable + twocolumn
        3. Bảng rộng (>4 cột, ≤12 hàng) → table*[t] + resizebox{textwidth}
        4. Bảng thường (≤4 cột, ≤12 hàng) → table[htbp] + resizebox{columnwidth}
        """
        try:
            if self.la_bang_metadata(bang):
                return self.tao_latex_minipage_metadata(bang)

            if self.la_bang_tieu_su(bang):
                return self.tao_latex_tieu_su_tac_gia(bang)

            if self.la_bang_layout(bang):
                return self.trich_xuat_noi_dung_bang_layout(bang)

            if self.la_bang_cong_thuc(bang):
                return self.xu_ly_bang_cong_thuc(bang)

            if self.la_table_of_contents(bang):
                if not self.bo_chuyen.toc_da_sinh:
                    self.bo_chuyen.toc_da_sinh = True
                    return r"\tableofcontents" + "\n\\newpage\n\n"
                return ""

            if self.la_bang_chua_anh(bang):
                danh_sach_anh = self.trich_xuat_anh_tu_bang(bang)
                if danh_sach_anh:
                    # Ảnh trong bảng: KHÔNG dùng \begin{figure} để tránh lỗi "Not in outer par mode"
                    if len(danh_sach_anh) > 1:
                        caption_con = self._tim_caption_con_trong_bang(bang)
                        caption_chinh = self.bo_chuyen.bat_caption_hinh()
                        return self.bo_chuyen.tao_latex_nhom_hinh(danh_sach_anh, caption_con, caption_chinh, trong_bang=True)
                    caption_chinh = self.bo_chuyen.bat_caption_hinh()
                    return self.bo_chuyen.tao_latex_hinh(danh_sach_anh[0], caption_chinh, trong_bang=True)

            self.bo_chuyen.so_bang_noi_dung += 1
            self.bo_chuyen.dem_bang += 1

            # Bảng siêu dài (>12 hàng): dùng onecolumn + longtable + twocolumn
            # vì longtable không hoạt động trong chế độ twocolumn của IEEE
            if self._la_bang_sieu_dai(bang):
                return self._render_bang_sieu_dai(bang)

            # Bảng bình thường: table[htbp] hoặc table*[t] với resizebox
            return self._render_tabular_merge(bang)
        except Exception as e:
            print(f"Lỗi xử lý bảng: {e}")
            return ""
