"""Round-trip test: Springer → IEEE → Springer.

Verifies that converting a Springer document to IEEE and then back to Springer
preserves key content: title, abstract, keywords, headings, body paragraphs,
and references.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from backend.core_engine.ast_parser import WordASTParser
from backend.core_engine.word_ieee_renderer import IEEEWordRenderer
from backend.core_engine.word_springer_renderer import SpringerWordRenderer

# Templates
IEEE_TEMPLATE = ROOT / "input_data" / "Template_word" / "conference-template-a4 (ieee).docx"
SPRINGER_TEMPLATE = ROOT / "input_data" / "Template_word" / "splnproc2510.docm"

# Sample Springer input (pick a real file if available)
SPRINGER_INPUTS = list((ROOT / "input_data" / "Template_word").glob("*.docm"))
SPRINGER_INPUTS += list((ROOT / "input_data" / "Template_word").glob("*.docx"))
# Filter out templates themselves
SPRINGER_INPUTS = [
    p for p in SPRINGER_INPUTS
    if "conference-template" not in p.name.lower()
    and "ieee" not in p.name.lower()
    and "acm" not in p.name.lower()
    and "Transactions-template" not in p.name
]


def _count_sections(body_nodes):
    return sum(1 for n in body_nodes if n.get("type") == "section")


def _count_paragraphs(body_nodes):
    return sum(1 for n in body_nodes if n.get("type") == "paragraph")


def _dump_ir_summary(ir, label=""):
    meta = ir.get("metadata", {})
    body = ir.get("body", [])
    refs = ir.get("references", [])
    print(f"\n{'='*60}")
    print(f"IR Summary: {label}")
    print(f"{'='*60}")
    print(f"  Title:    {(meta.get('title') or '')[:80]}")
    print(f"  Abstract: {(meta.get('abstract') or '')[:100]}...")
    print(f"  Keywords: {meta.get('keywords', [])}")
    print(f"  Authors:  {len(meta.get('authors', []))} authors")
    print(f"  Body:     {len(body)} nodes ({_count_sections(body)} sections, {_count_paragraphs(body)} paragraphs)")
    print(f"  Refs:     {len(refs)} references")
    print(f"{'='*60}")


@pytest.fixture
def tmp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


class TestRoundTripSpringerIEEE:
    """Test Springer → IEEE → Springer round-trip conversion."""

    def _parse_word(self, path, mode="word2word"):
        images_dir = path.parent / "images"
        images_dir.mkdir(exist_ok=True)
        parser = WordASTParser(str(path), thu_muc_anh=str(images_dir), mode=mode)
        return parser.parse()

    def _get_first_springer_input(self):
        # Try to find a real Springer input file
        candidates = [
            ROOT / "input_data" / "Template_word" / "splnproc2510.docm",
        ]
        # Also check any .docm files that look like Springer papers
        for p in (ROOT / "input_data" / "Template_word").glob("*.docm"):
            if p.name != "splnproc2510.docm":
                candidates.append(p)
        for p in candidates:
            if p.exists():
                return p
        pytest.skip("No Springer input file found")

    def test_parse_ieee_output_metadata(self, tmp_dir):
        """Test that parsing IEEE output correctly extracts metadata."""
        springer_input = self._get_first_springer_input()

        # Step 1: Parse Springer input
        ir_original = self._parse_word(springer_input)
        _dump_ir_summary(ir_original, "Original Springer")

        # Step 2: Render to IEEE
        ieee_output = tmp_dir / "ieee_output.docx"
        ieee_images = tmp_dir / "ieee_images"
        ieee_images.mkdir()
        renderer = IEEEWordRenderer()
        renderer.render(
            ir_data=ir_original,
            output_path=str(ieee_output),
            image_root_dir=str(tmp_dir),
            ieee_template_path=str(IEEE_TEMPLATE) if IEEE_TEMPLATE.exists() else None,
        )
        assert ieee_output.exists(), "IEEE output file not created"

        # Step 3: Parse IEEE output (this is where bugs lived)
        ir_from_ieee = self._parse_word(ieee_output)
        _dump_ir_summary(ir_from_ieee, "Parsed from IEEE output")

        # Verify: title should be preserved
        orig_title = (ir_original["metadata"].get("title") or "").strip()
        ieee_title = (ir_from_ieee["metadata"].get("title") or "").strip()
        if orig_title:
            assert ieee_title, f"Title lost after IEEE parse! Original: {orig_title[:60]}"

        # Verify: abstract should NOT start with em-dash or "Abstract"
        abstract = (ir_from_ieee["metadata"].get("abstract") or "").strip()
        assert not abstract.startswith("—"), \
            f"Abstract starts with em-dash (parsing bug): {abstract[:80]}"
        assert not abstract.lower().startswith("abstract"), \
            f"Abstract still has 'Abstract' prefix: {abstract[:80]}"

        # Verify: keywords should not contain label prefix
        keywords = ir_from_ieee["metadata"].get("keywords", [])
        for kw in keywords:
            assert "Keywords" not in kw, f"Keyword contains label prefix: {kw}"
            assert "Index Terms" not in kw, f"Keyword contains label prefix: {kw}"

        # Verify: body should have sections  
        sections = _count_sections(ir_from_ieee["body"])
        assert sections > 0, "No sections detected in IEEE output — heading detection broken"

    def test_full_round_trip(self, tmp_dir):
        """Test full Springer → IEEE → Springer round-trip."""
        springer_input = self._get_first_springer_input()

        # Step 1: Parse Springer
        ir_original = self._parse_word(springer_input)
        _dump_ir_summary(ir_original, "Step 1: Original Springer")

        # Step 2: Render IEEE
        ieee_output = tmp_dir / "ieee_output.docx"
        IEEEWordRenderer().render(
            ir_data=ir_original,
            output_path=str(ieee_output),
            image_root_dir=str(tmp_dir),
            ieee_template_path=str(IEEE_TEMPLATE) if IEEE_TEMPLATE.exists() else None,
        )

        # Step 3: Parse IEEE → IR
        ir_from_ieee = self._parse_word(ieee_output)
        _dump_ir_summary(ir_from_ieee, "Step 3: Parsed from IEEE")

        # Step 4: Render Springer
        springer_output = tmp_dir / "springer_output.docx"
        SpringerWordRenderer().render(
            ir_data=ir_from_ieee,
            output_path=str(springer_output),
            image_root_dir=str(tmp_dir),
            springer_template_path=str(SPRINGER_TEMPLATE) if SPRINGER_TEMPLATE.exists() else None,
        )
        assert springer_output.exists(), "Springer output not created"

        # Step 5: Parse final Springer output for validation
        ir_final = self._parse_word(springer_output)
        _dump_ir_summary(ir_final, "Step 5: Final Springer output")

        # ---- Validation ----
        orig_meta = ir_original["metadata"]
        final_meta = ir_final["metadata"]

        # Title preserved
        if orig_meta.get("title"):
            assert final_meta.get("title"), "Title lost in round-trip"

        # Abstract preserved (at least partially)
        orig_abs = (orig_meta.get("abstract") or "")[:200]
        final_abs = (final_meta.get("abstract") or "")[:200]
        if orig_abs:
            assert final_abs, f"Abstract lost in round-trip. Original: {orig_abs[:60]}"
            assert not final_abs.startswith("—"), "Abstract starts with em-dash"

        # Sections preserved (count should be roughly similar)
        orig_sections = _count_sections(ir_original["body"])
        final_sections = _count_sections(ir_final["body"])
        if orig_sections > 0:
            ratio = final_sections / max(1, orig_sections)
            assert ratio >= 0.5, \
                f"Too many sections lost: {orig_sections} -> {final_sections}"

        # Body not empty
        assert len(ir_final["body"]) > 0, "Body is empty after round-trip"

        print("\n[OK] Round-trip test PASSED")
        print(f"   Title: '{(final_meta.get('title') or '')[:60]}'")
        print(f"   Sections: {orig_sections} -> {final_sections}")
        print(f"   Body nodes: {len(ir_original['body'])} -> {len(ir_final['body'])}")

    def test_ieee_heading_detection(self, tmp_dir):
        """Test that IEEE Roman numeral headings are correctly detected as sections."""
        from docx import Document

        # Create a minimal IEEE-like document
        doc = Document()
        doc.add_paragraph("Test Paper Title").style = doc.styles["Title"]
        doc.add_paragraph("Abstract—This is the abstract content.")
        doc.add_paragraph("Keywords—keyword1, keyword2.")

        # Add IEEE-style headings
        doc.add_paragraph("I. INTRODUCTION")
        doc.add_paragraph("This is introduction text.")
        doc.add_paragraph("II. RELATED WORK")
        doc.add_paragraph("This is related work text.")
        doc.add_paragraph("A. First Subsection")
        doc.add_paragraph("Subsection text.")
        doc.add_paragraph("III. CONCLUSION")
        doc.add_paragraph("Conclusion text.")

        test_docx = tmp_dir / "test_ieee_headings.docx"
        doc.save(str(test_docx))

        # Parse
        ir = self._parse_word(test_docx)
        _dump_ir_summary(ir, "IEEE Heading Test")

        # Check sections detected
        sections = [n for n in ir["body"] if n.get("type") == "section"]
        section_texts = [n.get("text", "") for n in sections]
        print(f"  Detected sections: {section_texts}")

        assert len(sections) >= 3, \
            f"Expected at least 3 sections, got {len(sections)}: {section_texts}"

        # Check abstract doesn't have em-dash prefix
        abstract = ir["metadata"].get("abstract", "")
        assert not abstract.startswith("—"), f"Abstract has em-dash prefix: {abstract[:50]}"
        assert "abstract content" in abstract.lower(), f"Abstract content lost: {abstract[:50]}"

    def test_keywords_emdash_cleanup(self, tmp_dir):
        """Test that IEEE Keywords—/Index Terms— prefix is properly stripped."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test Title").style = doc.styles["Title"]
        doc.add_paragraph("Abstract—This is the abstract.")
        doc.add_paragraph("Keywords—machine learning, deep learning, classification.")
        doc.add_paragraph("I. INTRODUCTION")
        doc.add_paragraph("Body text.")

        test_docx = tmp_dir / "test_keywords.docx"
        doc.save(str(test_docx))

        ir = self._parse_word(test_docx)
        keywords = ir["metadata"].get("keywords", [])
        print(f"  Keywords: {keywords}")

        # Keywords should not contain "Keywords" label
        for kw in keywords:
            assert "keywords" not in kw.lower(), f"Keyword still has label: {kw}"
            assert "—" not in kw, f"Keyword has em-dash: {kw}"

        # Should have extracted actual keywords
        assert len(keywords) >= 2, f"Expected at least 2 keywords, got {len(keywords)}"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
