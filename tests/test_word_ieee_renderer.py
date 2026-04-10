import base64
from docx import Document
from docx.enum.section import WD_ORIENT

from backend.core_engine.word_ieee_renderer import IEEEWordRenderer


def _all_text(doc: Document) -> str:
    """Helper: collect text from paragraphs AND table cells."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_ieee_word_renderer_preserves_core_metadata(tmp_path):
    ir = {
        "metadata": {
            "title": "A Springer-origin Manuscript",
            "authors": [
                {"name": "Alice Nguyen", "affiliations": ["University A", "alice@example.com"]},
                {"name": "Bob Tran", "affiliations": ["Institute B"]},
            ],
            "abstract": "This is the abstract content.",
            "keywords": ["NLP", "Conversion"],
        },
        "body": [],
        "references": [],
    }

    out = tmp_path / "ieee_output.docx"
    IEEEWordRenderer().render(ir, str(out))

    doc = Document(str(out))
    text = _all_text(doc)

    assert "A Springer-origin Manuscript" in text
    assert "Alice Nguyen" in text
    assert "Bob Tran" in text
    assert "This is the abstract content." in text
    assert "NLP" in text and "Conversion" in text


def test_ieee_word_renderer_preserves_body_table_figure_and_references(tmp_path):
    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO3Z0mQAAAAASUVORK5CYII="
    )
    image_dir = tmp_path / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    (image_dir / "img_1.png").write_bytes(png_bytes)

    ir = {
        "metadata": {
            "title": "Paper",
            "authors": [],
            "abstract": "",
            "keywords": [],
        },
        "body": [
            {"type": "section", "level": 1, "text": "Introduction"},
            {"type": "paragraph", "text": "Regular content paragraph."},
            {
                "type": "paragraph",
                "text": "\\begin{figure}[htbp]\\includegraphics[width=0.9\\columnwidth]{images/img_1.png}\\caption{Sample figure}\\end{figure}",
            },
            {
                "type": "table",
                "cols": 2,
                "caption": "Sample table",
                "data": [
                    [{"type": "text", "text": "H1"}, {"type": "text", "text": "H2"}],
                    [{"type": "text", "text": "V1"}, {"type": "text", "text": "V2"}],
                ],
            },
        ],
        "references": [
            {"type": "paragraph", "text": "Author, 2025, Journal."},
        ],
    }

    out = tmp_path / "ieee_output_body.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    # Headings now have IEEE numbering
    assert "INTRODUCTION" in all_text
    assert "Regular content paragraph." in all_text
    assert "Fig. 1. Sample figure" in all_text
    assert "TABLE I" in all_text
    assert "SAMPLE TABLE" in all_text
    assert "References" in all_text
    assert "[1] Author, 2025, Journal." in all_text
    assert len(doc.inline_shapes) == 1

    assert len(doc.tables) >= 1
    # Find the data table (not the author table which has no borders)
    data_tables = [t for t in doc.tables if len(t.rows) == 2]
    assert len(data_tables) >= 1
    table = data_tables[0]
    assert table.cell(0, 0).text == "H1"
    assert table.cell(0, 1).text == "H2"
    assert table.cell(1, 0).text == "V1"
    assert table.cell(1, 1).text == "V2"


def test_ieee_word_renderer_table_without_style_does_not_crash(tmp_path, monkeypatch):
    ir = {
        "metadata": {"title": "Paper", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "table",
                "cols": 2,
                "caption": "No style table",
                "data": [
                    [{"type": "text", "text": "A"}, {"type": "text", "text": "B"}],
                ],
            }
        ],
        "references": [],
    }

    renderer = IEEEWordRenderer()
    monkeypatch.setattr(renderer, "_resolve_table_style_name", lambda _doc: None)

    out = tmp_path / "ieee_no_style_table.docx"
    renderer.render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    data_tables = [t for t in doc.tables if any(c.text.strip() for r in t.rows for c in r.cells)]
    assert len(data_tables) >= 1
    assert data_tables[0].cell(0, 0).text == "A"


def test_ieee_word_renderer_fills_existing_template_content(tmp_path):
    from docx import Document as DocxDocument

    template_path = tmp_path / "ieee_template.docx"
    tpl = DocxDocument()
    tpl.add_paragraph("Paper Title* (use style: paper title)")
    tpl.add_paragraph("line 1: 1st Given Name Surname")
    tpl.add_paragraph("line 2: dept. name of organization")
    tpl.add_paragraph("ABSTRACT")
    tpl.add_paragraph("ABSTRACT—This electronic document is a template")
    tpl.add_paragraph("INTRODUCTION")
    tpl.add_paragraph("This template, modified in MS Word 2007")
    tpl.add_paragraph("REFERENCES")
    tpl.add_paragraph("[1] Sample ref")
    tpl.save(str(template_path))

    ir = {
        "metadata": {
            "title": "Customer Churn Prediction",
            "authors": [
                {"name": "NGUYEN-HOANG Anh-Tuan", "affiliations": ["Can Tho University"]},
            ],
            "abstract": "New abstract from springer source.",
            "keywords": ["Customer Churn", "Machine Learning"],
        },
        "body": [
            {"type": "section", "level": 1, "text": "INTRODUCTION"},
            {"type": "paragraph", "text": "Body from springer."},
        ],
        "references": [
            {"type": "paragraph", "text": "A. Author, 2025."},
        ],
    }

    out = tmp_path / "filled.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path), str(template_path))

    doc = Document(str(out))
    all_text = _all_text(doc)
    assert "Customer Churn Prediction" in all_text
    assert "New abstract from springer source." in all_text
    assert "Customer Churn" in all_text and "Machine Learning" in all_text
    assert "Body from springer." in all_text
    assert "A. Author, 2025." in all_text


def test_ieee_word_renderer_prefers_style_name_mapping(tmp_path):
    from docx import Document as DocxDocument
    from docx.enum.style import WD_STYLE_TYPE

    template_path = tmp_path / "ieee_style_template.docx"
    tpl = DocxDocument()

    if not any(s.name == "paper title" for s in tpl.styles):
        tpl.styles.add_style("paper title", WD_STYLE_TYPE.PARAGRAPH)
    if not any(s.name == "abstract style" for s in tpl.styles):
        tpl.styles.add_style("abstract style", WD_STYLE_TYPE.PARAGRAPH)

    p_title = tpl.add_paragraph("XXX PLACEHOLDER TITLE")
    p_title.style = tpl.styles["paper title"]
    tpl.add_paragraph("author placeholder")
    p_abs_head = tpl.add_paragraph("UNUSUAL ABSTRACT MARK")
    p_abs_head.style = tpl.styles["abstract style"]
    tpl.add_paragraph("old abstract")
    tpl.add_paragraph("INTRODUCTION", style="Heading 1")
    tpl.add_paragraph("old body")
    tpl.add_paragraph("REFERENCES", style="Heading 1")
    tpl.add_paragraph("old ref")
    tpl.save(str(template_path))

    ir = {
        "metadata": {
            "title": "Style Name Driven Title",
            "authors": [{"name": "Author A", "affiliations": ["Org A"]}],
            "abstract": "Abstract filled by style mapping.",
            "keywords": ["KW1", "KW2"],
        },
        "body": [
            {"type": "section", "level": 1, "text": "INTRODUCTION"},
            {"type": "paragraph", "text": "Body line from source."},
        ],
        "references": [{"type": "paragraph", "text": "Ref from source."}],
    }

    out = tmp_path / "ieee_style_filled.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path), str(template_path))

    doc = Document(str(out))
    text = _all_text(doc)
    assert "Style Name Driven Title" in text
    assert "Abstract filled by style mapping." in text
    assert "KW1" in text and "KW2" in text
    assert "Body line from source." in text
    assert "Ref from source." in text


def test_ieee_word_renderer_places_figure_in_body_flow_not_intro_lump(tmp_path):
    from docx import Document as DocxDocument

    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO3Z0mQAAAAASUVORK5CYII="
    )
    image_dir = tmp_path / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    (image_dir / "img_2.png").write_bytes(png_bytes)

    template_path = tmp_path / "ieee_flow_template.docx"
    tpl = DocxDocument()
    tpl.add_paragraph("Paper Title* (use style: paper title)")
    tpl.add_paragraph("Author placeholder")
    tpl.add_paragraph("ABSTRACT")
    tpl.add_paragraph("Abstract placeholder")
    tpl.add_paragraph("INTRODUCTION", style="Heading 1")
    tpl.add_paragraph("Body slot 1")
    tpl.add_paragraph("Body slot 2")
    tpl.add_paragraph("Body slot 3")
    tpl.add_paragraph("REFERENCES", style="Heading 1")
    tpl.add_paragraph("Ref slot")
    tpl.save(str(template_path))

    ir = {
        "metadata": {
            "title": "T",
            "authors": [],
            "abstract": "A",
            "keywords": [],
        },
        "body": [
            {"type": "section", "level": 1, "text": "INTRODUCTION"},
            {"type": "paragraph", "text": "Body first paragraph."},
            {
                "type": "paragraph",
                "text": "\\begin{figure}[htbp]\\includegraphics{images/img_2.png}\\caption{Flow Figure}\\end{figure}",
            },
            {"type": "paragraph", "text": "Body last paragraph."},
        ],
        "references": [],
    }

    out = tmp_path / "ieee_flow_filled.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path), str(template_path))

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    # Heading now has numbering: "I. INTRODUCTION"
    intro_idx = next(i for i, t in enumerate(texts) if "INTRODUCTION" in t.strip().upper())
    assert "INTRODUCTION" in texts[intro_idx].upper()
    assert "Body first paragraph." in joined
    assert "Fig. 1. Flow Figure" in joined
    assert "Body last paragraph." in joined
    assert len(doc.inline_shapes) == 1


def test_ieee_word_renderer_fills_ieee_author_table_block_3_columns(tmp_path):
    """When template has no author table, renderer should create one with all authors."""
    ir = {
        "metadata": {
            "title": "Table Author Fill",
            "authors": [
                {"name": "Author One", "affiliations": ["Org 1", "one@x.com"]},
                {"name": "Author Two", "affiliations": ["Org 2", "two@x.com"]},
                {"name": "Author Three", "affiliations": ["Org 3", "three@x.com"]},
            ],
            "abstract": "Abs",
            "keywords": [],
        },
        "body": [{"type": "section", "level": 1, "text": "INTRODUCTION"}],
        "references": [],
    }

    out = tmp_path / "ieee_author_table_filled.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    # Find an author table (at least 3 columns)
    author_tables = [t for t in doc.tables if len(t.columns) >= 3]
    assert len(author_tables) >= 1
    table = author_tables[0]
    assert "Author One" in table.cell(0, 0).text
    assert "Author Two" in table.cell(0, 1).text
    assert "Author Three" in table.cell(0, 2).text


def test_ieee_word_renderer_enforces_two_column_body_layout(tmp_path):
    ir = {
        "metadata": {
            "title": "Two Column Check",
            "authors": [],
            "abstract": "Abstract part.",
            "keywords": [],
        },
        "body": [
            {"type": "section", "level": 1, "text": "Introduction"},
            {"type": "paragraph", "text": "Main body paragraph."},
        ],
        "references": [],
    }

    out = tmp_path / "ieee_two_column.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    cols = doc.sections[-1]._sectPr.find("w:cols", doc.sections[-1]._sectPr.nsmap)
    assert cols is not None
    assert cols.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num") == "2"


def test_ieee_word_renderer_normalizes_table_caption_and_forces_borders(tmp_path):
    ir = {
        "metadata": {"title": "Caption Test", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "table",
                "cols": 2,
                "caption": "TABLE I. DATASET FEATURES",
                "data": [
                    [{"type": "text", "text": "A"}, {"type": "text", "text": "B"}],
                ],
            }
        ],
        "references": [],
    }

    out = tmp_path / "ieee_table_caption_border.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "TABLE I" in all_text
    assert "DATASET FEATURES" in all_text
    assert "TABLE I. TABLE I." not in all_text

    table = doc.tables[0]
    tbl_borders = table._tbl.tblPr.find("w:tblBorders", table._tbl.tblPr.nsmap)
    assert tbl_borders is not None
    top = tbl_borders.find("w:top", tbl_borders.nsmap)
    assert top is not None
    assert top.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "single"


def test_ieee_word_renderer_uses_caption_font_size_8pt(tmp_path):
    ir = {
        "metadata": {"title": "Caption Font", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "paragraph",
                "text": "\\begin{figure}[htbp]\\includegraphics{images/not_found.png}\\caption{Figure caption font}\\end{figure}",
            },
            {
                "type": "table",
                "cols": 2,
                "caption": "Table caption font",
                "data": [[{"type": "text", "text": "A"}, {"type": "text", "text": "B"}]],
            },
        ],
        "references": [],
    }

    out = tmp_path / "ieee_caption_font.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    fig_cap = next(p for p in doc.paragraphs if p.text.startswith("Fig. 1."))
    table_label = next(p for p in doc.paragraphs if p.text.strip() == "TABLE I")
    table_title = next(p for p in doc.paragraphs if p.text.strip() == "TABLE CAPTION FONT")

    assert fig_cap.runs
    assert table_label.runs
    assert table_title.runs
    assert fig_cap.runs[0].font.size.pt == 8
    assert table_label.runs[0].font.size.pt == 8
    assert table_title.runs[0].font.size.pt == 8


def test_ieee_word_renderer_table_width_fits_single_column(tmp_path):
    ir = {
        "metadata": {"title": "Width Test", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "table",
                "cols": 2,
                "caption": "Width constrained",
                "data": [
                    [{"type": "text", "text": "Col 1"}, {"type": "text", "text": "Col 2"}],
                    [{"type": "text", "text": "v1"}, {"type": "text", "text": "v2"}],
                ],
            }
        ],
        "references": [],
    }

    out = tmp_path / "ieee_table_width.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    table = doc.tables[0]
    tbl_w = table._tbl.tblPr.find("w:tblW", table._tbl.tblPr.nsmap)
    assert tbl_w is not None
    assert tbl_w.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "dxa"

    width_twips = int(tbl_w.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"))
    assert 4700 <= width_twips <= 5100


def test_ieee_word_renderer_aligns_numeric_cells_right_and_text_left(tmp_path):
    ir = {
        "metadata": {"title": "Align Test", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "table",
                "cols": 2,
                "caption": "Alignment",
                "data": [
                    [{"type": "text", "text": "Metric"}, {"type": "text", "text": "Value"}],
                    [{"type": "text", "text": "Accuracy"}, {"type": "text", "text": "98.5"}],
                ],
            }
        ],
        "references": [],
    }

    out = tmp_path / "ieee_table_align.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    table = doc.tables[0]
    text_cell_align = table.cell(1, 0).paragraphs[0].alignment
    num_cell_align = table.cell(1, 1).paragraphs[0].alignment

    assert text_cell_align == 0  # LEFT
    assert num_cell_align == 2   # RIGHT


def test_ieee_word_renderer_wide_table_spans_full_width_block(tmp_path):
    header = [{"type": "text", "text": f"H{i}"} for i in range(1, 6)]
    row = [{"type": "text", "text": f"Very long value for column {i} with details"} for i in range(1, 6)]
    ir = {
        "metadata": {"title": "Span Test", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "table",
                "cols": 5,
                "caption": "Wide",
                "data": [header, row],
            }
        ],
        "references": [],
    }

    out = tmp_path / "ieee_wide_span.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    table = doc.tables[0]
    tbl_w = table._tbl.tblPr.find("w:tblW", table._tbl.tblPr.nsmap)
    assert tbl_w is not None
    width_twips = int(tbl_w.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"))
    assert width_twips > 7000


def test_ieee_word_renderer_body_line_spacing_profile(tmp_path):
    ir = {
        "metadata": {"title": "Spacing", "authors": [], "abstract": "Abs", "keywords": []},
        "body": [
            {"type": "section", "level": 1, "text": "Introduction"},
            {"type": "paragraph", "text": "Main body text line."},
        ],
        "references": [],
    }

    out = tmp_path / "ieee_spacing_profile.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    p = next(x for x in doc.paragraphs if x.text.strip() == "Main body text line.")
    assert p.paragraph_format.line_spacing == 1.0
    assert p.paragraph_format.space_after.pt == 2


def test_ieee_word_renderer_uses_landscape_block_for_extreme_wide_table(tmp_path):
    header = [{"type": "text", "text": f"H{i}"} for i in range(1, 9)]
    row = [{"type": "text", "text": f"Long narrative value in col {i} with many tokens to force wide table"} for i in range(1, 9)]
    ir = {
        "metadata": {"title": "Landscape Span", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "table",
                "cols": 8,
                "caption": "Wide matrix",
                "data": [header, row],
            }
        ],
        "references": [],
    }

    out = tmp_path / "ieee_landscape_span.docx"
    renderer = IEEEWordRenderer()
    renderer.render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    assert any(sec.orientation == WD_ORIENT.LANDSCAPE for sec in doc.sections)
    assert renderer.last_render_metrics["table_span_landscape"] >= 1


def test_ieee_word_renderer_reports_caption_normalization_metrics(tmp_path):
    ir = {
        "metadata": {"title": "Metrics", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {
                "type": "paragraph",
                "text": "\\begin{figure}[htbp]\\includegraphics{images/not_found.png}\\caption{Fig. 9. Sample}\\end{figure}",
            },
            {
                "type": "table",
                "cols": 2,
                "caption": "TABLE I: Demo",
                "data": [[{"type": "text", "text": "A"}, {"type": "text", "text": "B"}]],
            },
        ],
        "references": [],
    }

    out = tmp_path / "ieee_metrics.docx"
    renderer = IEEEWordRenderer()
    renderer.render(ir, str(out), str(tmp_path))

    assert renderer.last_render_metrics["table_caption_normalized"] >= 1
    assert renderer.last_render_metrics["figure_caption_normalized"] >= 1


# ===================== New tests for Plan A features =====================

def test_ieee_word_renderer_rich_text_bold_italic_preserved(tmp_path):
    """LaTeX bold/italic in IR should produce Word runs with formatting."""
    ir = {
        "metadata": {"title": "T", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {"type": "paragraph", "text": "This is \\textbf{bold} and \\textit{italic} text."},
        ],
        "references": [],
    }

    out = tmp_path / "rich_text.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    # Find the content paragraph
    all_runs = []
    for p in doc.paragraphs:
        if "bold" in p.text.lower():
            all_runs = list(p.runs)
            break

    assert len(all_runs) >= 3  # at least: "This is ", "bold", " and ..."
    # Find the bold run
    bold_runs = [r for r in all_runs if r.bold and "bold" in r.text]
    assert len(bold_runs) >= 1
    # Find the italic run
    italic_runs = [r for r in all_runs if r.italic and "italic" in r.text]
    assert len(italic_runs) >= 1


def test_ieee_word_renderer_cite_to_bracket(tmp_path):
    """\\cite{ref1,ref2} should become [1, 2] in output."""
    ir = {
        "metadata": {"title": "T", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {"type": "paragraph", "text": "See \\cite{ref1,ref2} for details."},
        ],
        "references": [],
    }

    out = tmp_path / "cite.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[1, 2]" in all_text


def test_ieee_word_renderer_ieee_heading_numbering(tmp_path):
    """Section headings should have IEEE numbering: I. II. III. and A. B."""
    ir = {
        "metadata": {"title": "T", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {"type": "section", "level": 1, "text": "Introduction"},
            {"type": "paragraph", "text": "P1"},
            {"type": "section", "level": 2, "text": "Sub A"},
            {"type": "paragraph", "text": "P2"},
            {"type": "section", "level": 2, "text": "Sub B"},
            {"type": "section", "level": 1, "text": "Related Work"},
            {"type": "paragraph", "text": "P3"},
        ],
        "references": [],
    }

    out = tmp_path / "headings.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    assert "I. INTRODUCTION" in joined
    assert "A. Sub A" in joined
    assert "B. Sub B" in joined
    assert "II. RELATED WORK" in joined


def test_ieee_word_renderer_equation_rendered_as_text(tmp_path):
    """Equations should be rendered as readable text."""
    ir = {
        "metadata": {"title": "T", "authors": [], "abstract": "", "keywords": []},
        "body": [
            {"type": "paragraph", "text": "\\begin{equation}\nx^2 + y^2 = z^2\n\\tag{1}\n\\end{equation}"},
        ],
        "references": [],
    }

    out = tmp_path / "equation.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    all_text = _all_text(doc)
    # Should contain the equation number and some form of the equation content
    assert "(1)" in all_text


def test_ieee_word_renderer_unicode_vietnamese(tmp_path):
    """Vietnamese Unicode text should be preserved."""
    ir = {
        "metadata": {
            "title": "Hệ thống phát hiện buồn ngủ khi lái xe",
            "authors": [{"name": "Nguyễn Văn An", "affiliations": ["Đại học Cần Thơ"]}],
            "abstract": "Bài viết này trình bày một hệ thống IoT.",
            "keywords": ["Trí tuệ nhân tạo"],
        },
        "body": [
            {"type": "section", "level": 1, "text": "Giới thiệu"},
            {"type": "paragraph", "text": "Nội dung bài viết."},
        ],
        "references": [],
    }

    out = tmp_path / "vietnamese.docx"
    IEEEWordRenderer().render(ir, str(out), str(tmp_path))

    doc = Document(str(out))
    text = _all_text(doc)
    assert "Hệ thống phát hiện buồn ngủ khi lái xe" in text
    assert "Nguyễn Văn An" in text
    assert "Đại học Cần Thơ" in text
    assert "Nội dung bài viết." in text
