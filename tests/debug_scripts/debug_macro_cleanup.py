"""Debug: Check if sua_docx_co_macro destroys w:br elements."""
import shutil
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from backend.core_engine.utils import sua_docx_co_macro

OUTPUT_DIR = Path(r"c:\221761_TIEN_PHONG_TT_VL_2026\debug_output5")
OUTPUT_DIR.mkdir(exist_ok=True)

# Step 1: Create fresh IEEE Word
from backend.core_engine.ast_parser import WordASTParser
from backend.core_engine.word_ieee_renderer import IEEEWordRenderer

INPUT_FILE = r"c:\221761_TIEN_PHONG_TT_VL_2026\input_data\Template_word\-TUAN_Magazine_01-12- Customer Churn Prediction in Vietnam's Enterprise Market Using Machine Learning Methods in a Streaming Data (1).docx"
IEEE_TEMPLATE = r"c:\221761_TIEN_PHONG_TT_VL_2026\input_data\Template_word\conference-template-a4 (ieee).docx"

(OUTPUT_DIR / "images1").mkdir(exist_ok=True)
parser1 = WordASTParser(INPUT_FILE, thu_muc_anh=str(OUTPUT_DIR / "images1"), mode="latex")
ir1 = parser1.parse()

ieee_file = str(OUTPUT_DIR / "test_ieee.docx")
renderer = IEEEWordRenderer()
renderer.render(ir_data=ir1, output_path=ieee_file, image_root_dir=str(OUTPUT_DIR), ieee_template_path=IEEE_TEMPLATE)

# Copy to have a backup
backup_file = str(OUTPUT_DIR / "test_ieee_backup.docx")
shutil.copy2(ieee_file, backup_file)

# Check cell content BEFORE sua_docx_co_macro
print("=== BEFORE sua_docx_co_macro ===")
doc_before = Document(backup_file)
body = doc_before.element.body
for child in body:
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag == "tbl":
        t = Table(child, doc_before)
        if len(t.rows) == 1:
            for ci, cell in enumerate(t.rows[0].cells):
                print(f"  Cell [{ci}]:")
                print(f"    cell.text repr: {repr(cell.text.strip()[:80])}")
                for pi, p in enumerate(cell.paragraphs):
                    print(f"    p[{pi}].text repr: {repr(p.text.strip()[:80])}")
                    br_count = len(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'))
                    print(f"    p[{pi}] <w:br> count: {br_count}")
            break

# Run sua_docx_co_macro on the test file
print("\n--- Running sua_docx_co_macro ---")
sua_docx_co_macro(ieee_file)

# Check cell content AFTER sua_docx_co_macro  
print("\n=== AFTER sua_docx_co_macro ===")
doc_after = Document(ieee_file)
body2 = doc_after.element.body
for child in body2:
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag == "tbl":
        t = Table(child, doc_after)
        if len(t.rows) == 1:
            for ci, cell in enumerate(t.rows[0].cells):
                print(f"  Cell [{ci}]:")
                print(f"    cell.text repr: {repr(cell.text.strip()[:80])}")
                for pi, p in enumerate(cell.paragraphs):
                    print(f"    p[{pi}].text repr: {repr(p.text.strip()[:80])}")
                    br_count = len(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'))
                    print(f"    p[{pi}] <w:br> count: {br_count}")
            break
