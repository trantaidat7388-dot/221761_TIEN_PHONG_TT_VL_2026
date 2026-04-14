"""Debug: Compare cell.text vs cell.paragraphs[i].text."""
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

ieee_file = str(Path(r"c:\221761_TIEN_PHONG_TT_VL_2026\debug_output3\step1_ieee.docx"))
doc = Document(ieee_file)
body = doc.element.body

for child in body:
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag == "tbl":
        t = Table(child, doc)
        if len(t.rows) == 1:
            print("=== Author table found ===")
            for ci, cell in enumerate(t.rows[0].cells):
                print(f"\n  Cell [{ci}]:")
                print(f"    cell.text repr: {repr(cell.text[:100])}")
                print(f"    cell.paragraphs count: {len(list(cell.paragraphs))}")
                for pi, p in enumerate(cell.paragraphs):
                    print(f"    p[{pi}].text repr: {repr(p.text[:100])}")
                    # Check for br elements
                    from lxml import etree
                    br_elements = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                    print(f"    p[{pi}] <w:br> count: {len(br_elements)}")
            break
