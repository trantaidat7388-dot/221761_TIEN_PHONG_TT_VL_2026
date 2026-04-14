"""Debug: Run full 2-step pipeline and inspect intermediate/final Word docs."""
import sys
import json
from pathlib import Path

# Step 1: Original Springer Word -> IEEE Word
from backend.core_engine.ast_parser import WordASTParser
from backend.core_engine.word_ieee_renderer import IEEEWordRenderer
from backend.core_engine.word_springer_renderer import SpringerWordRenderer

INPUT_FILE = r"c:\221761_TIEN_PHONG_TT_VL_2026\input_data\Template_word\-TUAN_Magazine_01-12- Customer Churn Prediction in Vietnam's Enterprise Market Using Machine Learning Methods in a Streaming Data (1).docx"
IEEE_TEMPLATE = r"c:\221761_TIEN_PHONG_TT_VL_2026\input_data\Template_word\conference-template-a4 (ieee).docx"
SPRINGER_TEMPLATE = r"c:\221761_TIEN_PHONG_TT_VL_2026\input_data\Template_word\splnproc2510.docm"
OUTPUT_DIR = Path(r"c:\221761_TIEN_PHONG_TT_VL_2026\debug_output")
OUTPUT_DIR.mkdir(exist_ok=True)

print("=" * 60)
print("STEP 1: Parse original Springer Word -> IR")
print("=" * 60)

parser1 = WordASTParser(INPUT_FILE, thu_muc_anh=str(OUTPUT_DIR / "images1"), mode="word2word")
ir1 = parser1.parse()
print(f"\nTitle: {ir1['metadata']['title'][:100]}")
print(f"Authors: {json.dumps(ir1['metadata']['authors'][:3], indent=2, ensure_ascii=False)[:500]}")
print(f"Abstract: {ir1['metadata']['abstract'][:100]}...")

print("\n" + "=" * 60)
print("STEP 2: Render to IEEE Word")
print("=" * 60)

ieee_output = str(OUTPUT_DIR / "step2_ieee.docx")
renderer_ieee = IEEEWordRenderer()
renderer_ieee.render(
    ir_data=ir1,
    output_path=ieee_output,
    image_root_dir=str(OUTPUT_DIR),
    ieee_template_path=IEEE_TEMPLATE,
)
print(f"IEEE Word saved to: {ieee_output}")

# Inspect the IEEE Word document
from docx import Document
doc_ieee = Document(ieee_output)
print("\n--- IEEE Word first 15 paragraphs ---")
for i, p in enumerate(doc_ieee.paragraphs[:15]):
    style = p.style.name if p.style else "None"
    text = p.text.strip()[:120]
    print(f"  P{i:3d}: style='{style}' => '{text}'")

print("\n" + "=" * 60)
print("STEP 3: Parse IEEE Word -> IR (2nd pass)")
print("=" * 60)

parser2 = WordASTParser(ieee_output, thu_muc_anh=str(OUTPUT_DIR / "images2"), mode="word2word")
ir2 = parser2.parse()
print(f"\nTitle: '{ir2['metadata']['title'][:200]}'")
print(f"Authors: {json.dumps(ir2['metadata']['authors'][:3], indent=2, ensure_ascii=False)[:500]}")
print(f"Abstract: {ir2['metadata']['abstract'][:100]}...")

print("\n" + "=" * 60)
print("STEP 4: Render to Springer Word")
print("=" * 60)

springer_output = str(OUTPUT_DIR / "step4_springer.docx")
renderer_springer = SpringerWordRenderer()
renderer_springer.render(
    ir_data=ir2,
    output_path=springer_output,
    image_root_dir=str(OUTPUT_DIR),
    springer_template_path=SPRINGER_TEMPLATE,
)
print(f"Springer Word saved to: {springer_output}")

# Inspect the Springer Word document
doc_springer = Document(springer_output)
print("\n--- Springer Word first 15 paragraphs ---")
for i, p in enumerate(doc_springer.paragraphs[:15]):
    style = p.style.name if p.style else "None"
    text = p.text.strip()[:150]
    print(f"  P{i:3d}: style='{style}' => '{text}'")

print("\n" + "=" * 60)
print("STEP 5: Parse Springer Word -> IR (3rd pass, for LaTeX)")
print("=" * 60)

parser3 = WordASTParser(springer_output, thu_muc_anh=str(OUTPUT_DIR / "images3"), mode="latex")
ir3 = parser3.parse()
print(f"\nTitle: '{ir3['metadata']['title'][:300]}'")
print(f"Authors: {json.dumps(ir3['metadata']['authors'][:5], indent=2, ensure_ascii=False)[:600]}")
print(f"Abstract: {ir3['metadata']['abstract'][:100]}...")

print("\nDONE!")
