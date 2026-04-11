from docx import Document

from backend.core_engine.word_springer_renderer import SpringerWordRenderer


def _collect_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(c.text)
    return "\n".join(parts)


def test_springer_heading_deduplicates_existing_numbers(tmp_path):
    ir = {
        "metadata": {
            "title": "Roundtrip Heading Test",
            "authors": [],
            "abstract": "",
            "keywords": [],
        },
        "body": [
            {"type": "section", "level": 1, "text": "3 3 EXPERIMENT AND DISCUSSION"},
        ],
        "references": [],
    }

    out = tmp_path / "springer_heading.docx"
    SpringerWordRenderer().render(ir, str(out))

    doc = Document(str(out))
    text = _collect_text(doc)
    assert "3 3 EXPERIMENT" not in text
    assert "1 Experiment And Discussion" in text


def test_springer_table_caption_and_reference_number_cleanup(tmp_path):
    ir = {
        "metadata": {
            "title": "Roundtrip Caption Ref Test",
            "authors": [],
            "abstract": "",
            "keywords": [],
        },
        "body": [
            {
                "type": "table",
                "cols": 2,
                "caption": "TABLE II AVERAGE OF BALANCED ACCURACY",
                "data": [
                    [{"type": "text", "text": "H1"}, {"type": "text", "text": "H2"}],
                    [{"type": "text", "text": "1"}, {"type": "text", "text": "2"}],
                ],
            }
        ],
        "references": [
            {"text": "1. 1. Juliane Waack (2025). https://example.com/ref"},
        ],
    }

    out = tmp_path / "springer_caption_ref.docx"
    SpringerWordRenderer().render(ir, str(out))

    doc = Document(str(out))
    text = _collect_text(doc)

    assert "TABLE II" not in text
    assert "Table 1. AVERAGE OF BALANCED ACCURACY" in text
    assert "1. 1. Juliane" not in text
    assert "1. Juliane Waack" in text


def test_springer_syncs_header_short_title():
    renderer = SpringerWordRenderer()
    doc = Document()
    header = doc.sections[0].header
    if not header.paragraphs:
        header.add_paragraph("")
    header.paragraphs[0].text = "Contribution Title (shortened if too long)\t7"

    renderer._sync_template_header_title(doc, {"title": "Customer Churn Prediction in Vietnam"})

    assert "Contribution Title" not in header.paragraphs[0].text
    assert "Customer Churn Prediction in Vietnam" in header.paragraphs[0].text


def test_springer_skips_duplicate_table_title_paragraph(tmp_path):
    ir = {
        "metadata": {
            "title": "Roundtrip Duplicate Caption Test",
            "authors": [],
            "abstract": "",
            "keywords": [],
        },
        "body": [
            {"type": "paragraph", "text": "DATASET FEATURES AND DESCRIPTIONS"},
            {
                "type": "table",
                "cols": 2,
                "caption": "DATASET FEATURES AND DESCRIPTIONS",
                "data": [
                    [{"type": "text", "text": "Feature"}, {"type": "text", "text": "Type"}],
                    [{"type": "text", "text": "Age"}, {"type": "text", "text": "Numeric"}],
                ],
            },
        ],
        "references": [],
    }

    out = tmp_path / "springer_duplicate_table_title.docx"
    SpringerWordRenderer().render(ir, str(out))

    doc = Document(str(out))
    combined = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    assert combined.count("DATASET FEATURES AND DESCRIPTIONS") == 1


def test_springer_detects_plain_numbered_equation(tmp_path):
    ir = {
        "metadata": {
            "title": "Roundtrip Equation Test",
            "authors": [],
            "abstract": "",
            "keywords": [],
        },
        "body": [
            {"type": "paragraph", "text": "Balanced Accuracy (BA) = 1 2 (Specificity + Sensitivity) (1)"},
        ],
        "references": [],
    }

    out = tmp_path / "springer_equation.docx"
    SpringerWordRenderer().render(ir, str(out))

    doc = Document(str(out))
    eq_paras = [p for p in doc.paragraphs if "Balanced Accuracy" in p.text]
    assert eq_paras, "Expected equation-like paragraph to be rendered"
    assert "(1)" in eq_paras[0].text
    assert eq_paras[0].alignment == 1  # WD_ALIGN_PARAGRAPH.CENTER


def test_springer_drops_orphan_table_label_paragraph(tmp_path):
    ir = {
        "metadata": {
            "title": "Roundtrip Orphan Table Label",
            "authors": [],
            "abstract": "",
            "keywords": [],
        },
        "body": [
            {"type": "paragraph", "text": "TABLE II"},
            {
                "type": "table",
                "cols": 2,
                "caption": "AVERAGE OF BALANCED ACCURACY",
                "data": [
                    [{"type": "text", "text": "H1"}, {"type": "text", "text": "H2"}],
                    [{"type": "text", "text": "1"}, {"type": "text", "text": "2"}],
                ],
            },
        ],
        "references": [],
    }

    out = tmp_path / "springer_orphan_table_label.docx"
    SpringerWordRenderer().render(ir, str(out))

    doc = Document(str(out))
    text = _collect_text(doc)
    assert "TABLE II" not in text
    assert "Table 1. AVERAGE OF BALANCED ACCURACY" in text
